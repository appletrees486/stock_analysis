#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
대량 분석 처리 클래스
비동기적으로 다수 종목을 분석하는 기능
"""

import os
import json
import threading
import zipfile
import logging
from datetime import datetime
from typing import List, Dict, Any
import time
import sys

# 프로젝트 루트를 Python 경로에 추가
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

logger = logging.getLogger(__name__)

class BatchAnalyzer:
    _instance = None
    _lock = threading.Lock()
    
    def __new__(cls):
        if cls._instance is None:
            with cls._lock:
                if cls._instance is None:
                    cls._instance = super(BatchAnalyzer, cls).__new__(cls)
                    cls._instance.batch_status = {}
                    cls._instance.batch_results = {}
                    cls._instance.batch_threads = {}
        return cls._instance
    
    def __init__(self):
        # 싱글톤이므로 초기화는 한 번만
        pass
    
    def start_batch_analysis(self, stock_list_path: str, chart_type: str, batch_id: str, trading_type: str = ''):
        """대량 분석 시작"""
        try:
            logger.info(f"대량 분석 시작: batch_id={batch_id}, trading_type={trading_type}")
            
            # 초기 상태 설정
            self.batch_status[batch_id] = {
                'status': 'running',
                'total': 0,
                'completed': 0,
                'failed': 0,
                'start_time': datetime.now().isoformat(),
                'progress': 0,
                'chart_type': chart_type,
                'trading_type': trading_type
            }
            
            # 동기적으로 분석 실행 (스레드 문제 해결)
            self._run_batch_analysis(stock_list_path, chart_type, batch_id, trading_type)
            
            logger.info(f"대량 분석 완료: batch_id={batch_id}")
            
        except Exception as e:
            logger.error(f"대량 분석 시작 오류: {e}")
            self.batch_status[batch_id] = {
                'status': 'failed',
                'error': str(e),
                'start_time': datetime.now().isoformat()
            }
    
    def _run_batch_analysis(self, stock_list_path: str, chart_type: str, batch_id: str, trading_type: str = ''):
        """백그라운드에서 대량 분석 실행"""
        try:
            # 종목 리스트 읽기
            from .utils import get_stock_list_from_file
            stock_codes = get_stock_list_from_file(stock_list_path)
            
            if not stock_codes:
                raise Exception("종목 리스트가 비어있습니다.")
            
            self.batch_status[batch_id]['total'] = len(stock_codes)
            self.batch_results[batch_id] = []
            
            logger.info(f"대량 분석 시작: {len(stock_codes)}개 종목, 차트타입={chart_type}, 거래타입={trading_type}")
            
            # 각 종목별 분석
            for i, stock_code in enumerate(stock_codes):
                try:
                    logger.info(f"종목 분석 중: {stock_code} ({i+1}/{len(stock_codes)})")
                    
                    # 기존 배치 분석 모듈 활용
                    from batch_stock_analyzer_optimized import analyze_single_stock_fast
                    
                    # 차트 타입 매핑
                    chart_type_mapping = {
                        "일봉": "daily",
                        "주봉": "weekly", 
                        "월봉": "monthly"
                    }
                    chart_type_en = chart_type_mapping.get(chart_type, "daily")
                    
                    logger.info(f"배치 분석 시작: {stock_code}, 차트타입={chart_type}, 차트타입_en={chart_type_en}")
                    
                    # FastProgressTracker 인스턴스 생성
                    from batch_stock_analyzer_optimized import FastProgressTracker
                    tracker = FastProgressTracker(1)
                    
                    batch_result = analyze_single_stock_fast(
                        stock_code, chart_type, chart_type_en, tracker, batch_id, trading_type
                    )
                    
                    logger.info(f"배치 분석 결과: {batch_result}")
                    logger.info(f"배치 분석 결과 타입: {type(batch_result)}")
                    logger.info(f"배치 분석 success 값: {batch_result.get('success', 'NOT_FOUND')}")
                    logger.info(f"배치 분석 error 값: {batch_result.get('error', 'NOT_FOUND')}")
                    
                    # 배치 결과를 AI 분석 결과 형식으로 변환
                    if batch_result.get('success', False):
                        logger.info(f"종목 {stock_code} 분석 성공")
                        # 성공한 경우 실제 AI 분석 결과 파일 찾기
                        ai_results_dir = "ai_analysis_results"
                        if os.path.exists(ai_results_dir):
                            # 해당 종목의 최신 분석 결과 파일 찾기
                            pattern = f"analysis_{chart_type_en}_{stock_code}_*.json"
                            json_files = []
                            for file in os.listdir(ai_results_dir):
                                if file.startswith(f"analysis_{chart_type_en}_{stock_code}_") and file.endswith('.json'):
                                    json_files.append(file)
                            
                            logger.info(f"찾은 JSON 파일들: {json_files}")
                            
                            if json_files:
                                # 가장 최근 JSON 파일 로드
                                latest_json = sorted(json_files)[-1]
                                json_path = os.path.join(ai_results_dir, latest_json)
                                
                                try:
                                    with open(json_path, 'r', encoding='utf-8') as f:
                                        ai_analysis_result = json.load(f)
                                    
                                    # AI 분석 결과에서 정보 추출
                                    stock_info = ai_analysis_result.get("종목정보", {})
                                    analysis_score = ai_analysis_result.get("종합분석점수", {})
                                    
                                    result = {
                                        'stock_name': stock_info.get('종목명', stock_code),
                                        'stock_code': stock_code,
                                        'chart_type': chart_type,
                                        'analysis_score': analysis_score.get('점수', 75),
                                        'summary': analysis_score.get('요약', f'{stock_code} 종목 분석이 완료되었습니다.'),
                                        'detailed_analysis': f'{stock_code} 종목의 {chart_type} 차트 분석이 성공적으로 완료되었습니다.',
                                        'timestamp': datetime.now().isoformat(),
                                        'ai_analysis_file': latest_json,
                                        'ai_analysis_data': ai_analysis_result,
                                        'success': True,
                                        'ai_analysis_done': True
                                    }
                                    logger.info(f"AI 분석 결과 로드 성공: {latest_json}")
                                except Exception as e:
                                    logger.warning(f"AI 분석 결과 파일 로드 실패: {e}")
                                    # 기본 결과 생성
                                    result = {
                                        'stock_name': stock_code,
                                        'chart_type': chart_type,
                                        'analysis_score': 75,
                                        'summary': f'{stock_code} 종목 분석이 완료되었습니다.',
                                        'detailed_analysis': f'{stock_code} 종목의 {chart_type} 차트 분석이 성공적으로 완료되었습니다.',
                                        'timestamp': datetime.now().isoformat(),
                                        'success': True,
                                        'ai_analysis_done': False
                                    }
                            else:
                                logger.warning(f"종목 {stock_code}의 AI 분석 결과 파일을 찾을 수 없습니다")
                                # AI 분석 결과 파일이 없는 경우 기본 결과 생성
                                result = {
                                    'stock_name': stock_code,
                                    'chart_type': chart_type,
                                    'analysis_score': 75,
                                    'summary': f'{stock_code} 종목 분석이 완료되었습니다.',
                                    'detailed_analysis': f'{stock_code} 종목의 {chart_type} 차트 분석이 성공적으로 완료되었습니다.',
                                    'timestamp': datetime.now().isoformat(),
                                    'success': True,
                                    'ai_analysis_done': False
                                }
                        else:
                            logger.warning(f"AI 분석 결과 폴더가 존재하지 않습니다: {ai_results_dir}")
                            # AI 분석 결과 폴더가 없는 경우 기본 결과 생성
                            result = {
                                'stock_name': stock_code,
                                'chart_type': chart_type,
                                'analysis_score': 75,
                                'summary': f'{stock_code} 종목 분석이 완료되었습니다.',
                                'detailed_analysis': f'{stock_code} 종목의 {chart_type} 차트 분석이 성공적으로 완료되었습니다.',
                                'timestamp': datetime.now().isoformat(),
                                'success': True,
                                'ai_analysis_done': False
                            }
                    else:
                        logger.error(f"종목 {stock_code} 분석 실패: {batch_result.get('error', '알 수 없는 오류')}")
                        logger.error(f"배치 결과 상세: {batch_result}")
                        # 실패한 경우 에러 결과 생성
                        result = {
                            'stock_name': stock_code,
                            'chart_type': chart_type,
                            'analysis_score': 0,
                            'summary': f'분석 실패: {batch_result.get("error", "알 수 없는 오류")}',
                            'detailed_analysis': f'종목 {stock_code} 분석 중 오류가 발생했습니다. 오류: {batch_result.get("error", "알 수 없는 오류")}',
                            'timestamp': datetime.now().isoformat(),
                            'success': False,
                            'ai_analysis_done': False
                        }
                    
                    # 결과에 메타데이터 추가
                    result['batch_id'] = batch_id
                    result['stock_code'] = stock_code
                    result['chart_type'] = chart_type
                    result['trading_type'] = trading_type  # 거래타입 추가
                    result['processed_at'] = datetime.now().isoformat()
                    
                    self.batch_results[batch_id].append(result)
                    self.batch_status[batch_id]['completed'] += 1
                    
                    logger.info(f"종목 분석 완료: {stock_code}")
                    
                except Exception as e:
                    logger.error(f"종목 {stock_code} 분석 실패: {e}")
                    self.batch_status[batch_id]['failed'] += 1
                    
                    # 실패한 종목에 대한 에러 결과 추가
                    error_result = {
                        'batch_id': batch_id,
                        'stock_code': stock_code,
                        'chart_type': chart_type,
                        'trading_type': trading_type,  # 거래타입 추가
                        'processed_at': datetime.now().isoformat(),
                        'error': str(e),
                        'analysis_score': 0,
                        'summary': f'분석 실패: {str(e)}',
                        'detailed_analysis': f'종목 {stock_code} 분석 중 오류가 발생했습니다.'
                    }
                    self.batch_results[batch_id].append(error_result)
                
                # 진행률 업데이트
                progress = (i + 1) / len(stock_codes) * 100
                self.batch_status[batch_id]['progress'] = progress
            
            # 분석 완료
            self.batch_status[batch_id]['status'] = 'completed'
            self.batch_status[batch_id]['end_time'] = datetime.now().isoformat()
            
            logger.info(f"배치 분석 완료: batch_id={batch_id}, 총 {len(self.batch_results[batch_id])}개 결과")
            
            # 요약 분석 실행 (ZIP에 추가하기 위해 결과 저장)
            summary_files = None
            try:
                logger.info(f"요약 분석 시작: batch_id={batch_id}, chart_type={chart_type}")
                # 배치 결과에서 성공한 종목들의 파일 경로 추출
                successful_stocks = []
                logger.info(f"배치 ID: {batch_id}")
                logger.info(f"배치 결과 키들: {list(self.batch_results.keys())}")
                
                if batch_id in self.batch_results:
                    batch_results = self.batch_results[batch_id]
                    logger.info(f"배치 결과에서 파일 경로 추출 중: {len(batch_results)}개 결과")
                    logger.info(f"배치 결과 전체: {batch_results}")
                    
                    for i, result in enumerate(batch_results):
                        logger.info(f"배치 결과 {i+1} 검토: {result}")
                        logger.info(f"  - success: {result.get('success', 'NOT_FOUND')}")
                        logger.info(f"  - ai_analysis_done: {result.get('ai_analysis_done', 'NOT_FOUND')}")
                        logger.info(f"  - stock_code: {result.get('stock_code', 'NOT_FOUND')}")
                        
                        if result.get('success', False) and result.get('ai_analysis_done', False):
                            stock_code = result.get('stock_code')
                            logger.info(f"성공한 종목 발견: {stock_code}")
                            if stock_code:
                                # 배치 ID를 포함한 파일명으로 검색 (더 정확한 방법)
                                # 차트 유형을 영문으로 변환
                                chart_type_en = {
                                    "일봉": "daily",
                                    "주봉": "weekly", 
                                    "월봉": "monthly"
                                }.get(chart_type, chart_type.lower())
                                
                                batch_pattern = f"ai_analysis_results/analysis_{chart_type_en}_{stock_code}_*_{batch_id}.json"
                                import glob
                                batch_files = glob.glob(batch_pattern)
                                logger.info(f"배치 ID 패턴 검색: {batch_pattern}")
                                logger.info(f"찾은 배치 파일들: {batch_files}")
                                
                                if batch_files:
                                    # 배치 ID가 포함된 파일이 있으면 사용
                                    latest_file = max(batch_files, key=os.path.getctime)
                                    successful_stocks.append(latest_file)
                                    logger.info(f"배치 ID 파일 선택: {latest_file}")
                                else:
                                    # 배치 ID가 포함된 파일이 없으면 시간 기반으로 검색
                                    json_file = f"ai_analysis_results/analysis_{chart_type}_{stock_code}_*.json"
                                    matching_files = glob.glob(json_file)
                                    if matching_files:
                                        # 배치 시작 시간 이후에 생성된 파일만 선택
                                        batch_start_time = self.batch_status.get(batch_id, {}).get('start_time')
                                        if batch_start_time:
                                            batch_start_dt = datetime.fromisoformat(batch_start_time.replace('Z', '+00:00'))
                                            recent_files = []
                                            for file_path in matching_files:
                                                file_time = datetime.fromtimestamp(os.path.getctime(file_path))
                                                if file_time >= batch_start_dt:
                                                    recent_files.append(file_path)
                                            
                                            if recent_files:
                                                # 가장 최근 파일 선택
                                                latest_file = max(recent_files, key=os.path.getctime)
                                                successful_stocks.append(latest_file)
                                                logger.info(f"시간 기반 파일 선택: {latest_file}")
                                            else:
                                                logger.warning(f"배치 시작 시간 이후 파일 없음: {stock_code}")
                                        else:
                                            # 배치 시작 시간이 없으면 가장 최근 파일 선택 (기존 방식)
                                            latest_file = max(matching_files, key=os.path.getctime)
                                            successful_stocks.append(latest_file)
                                            logger.info(f"최근 파일 선택: {latest_file}")
                
                logger.info(f"추출된 성공 파일 수: {len(successful_stocks)}")
                for file_path in successful_stocks:
                    logger.info(f"  - {file_path}")
                
                summary_files = self._create_summary_analysis(chart_type, successful_stocks)
                logger.info(f"요약 분석 완료: batch_id={batch_id}")
            except Exception as e:
                logger.error(f"요약 분석 중 오류: {e}")
                logger.info("개별 분석 결과는 정상적으로 저장되었습니다.")
            
            # 요약 분석 결과를 배치 결과에 저장 (ZIP 생성 시 사용)
            if summary_files and batch_id in self.batch_results:
                self.batch_results[batch_id].append({
                    'type': 'summary_info',
                    'chart_type': chart_type,
                    'summary_files': summary_files
                })
            
            # 통합 파일 생성 (기존 방식 - 호환성 유지)
            try:
                logger.info(f"통합 분석 파일 생성 시작: batch_id={batch_id}")
                self._create_consolidated_files(batch_id, chart_type)
                logger.info(f"통합 분석 파일 생성 완료: batch_id={batch_id}")
            except Exception as e:
                logger.error(f"통합 파일 생성 중 오류: {e}")
                logger.info("개별 파일들은 정상적으로 생성되었습니다.")
            
            # 결과 저장
            try:
                logger.info(f"배치 결과 저장 시작: batch_id={batch_id}")
                self._save_batch_results(batch_id)
                logger.info(f"배치 결과 저장 완료: batch_id={batch_id}")
            except Exception as e:
                logger.error(f"배치 결과 저장 실패: batch_id={batch_id}, 오류: {e}")
                self.batch_status[batch_id]['status'] = 'failed'
                self.batch_status[batch_id]['error'] = f"결과 저장 실패: {str(e)}"
            
            logger.info(f"대량 분석 완료: batch_id={batch_id}")
            
        except Exception as e:
            logger.error(f"대량 분석 실행 오류: {e}")
            import traceback
            logger.error(f"상세 오류: {traceback.format_exc()}")
            self.batch_status[batch_id]['status'] = 'failed'
            self.batch_status[batch_id]['error'] = str(e)
            self.batch_status[batch_id]['end_time'] = datetime.now().isoformat()
    
    def get_batch_status(self, batch_id: str) -> Dict[str, Any]:
        """배치 상태 조회"""
        if batch_id in self.batch_status:
            return self.batch_status[batch_id]
        else:
            return {'error': '배치를 찾을 수 없습니다'}
    
    def get_batch_results(self, batch_id: str) -> Dict[str, Any]:
        """배치 결과 조회"""
        if batch_id in self.batch_results:
            return {
                'batch_id': batch_id,
                'results': self.batch_results[batch_id],
                'status': self.batch_status.get(batch_id, {}),
                'total_results': len(self.batch_results[batch_id])
            }
        else:
            return {'error': '결과를 찾을 수 없습니다'}
    
    def _generate_zip_filename(self, batch_id: str) -> str:
        """summary_meta 정보를 기반으로 zip 파일명 생성"""
        try:
            # summary JSON 파일에서 메타데이터 읽기
            summary_meta = self._get_summary_meta(batch_id)
            if summary_meta:
                chart_type = summary_meta.get('chart_type', '')
                trading_date = summary_meta.get('trading_date', '')
                trading_type = summary_meta.get('trading_type', '')
                total_stocks = summary_meta.get('total_stocks', 0)
                generated_at = summary_meta.get('generated_at', '').replace(' ', '_').replace(':', '').replace('-', '')
                
                # 파일명 생성: {chart_type}_{trading_date}_{trading_type}_{total_stocks}개_{generated_at}_{batch_id}.zip
                zip_filename = f"{chart_type}_{trading_date}_{trading_type}_{total_stocks}개_{generated_at}_{batch_id}.zip"
                return f"results/{zip_filename}"
            else:
                # summary_meta를 찾을 수 없는 경우 기본 파일명 사용
                return f"results/{batch_id}_results.zip"
        except Exception as e:
            logger.warning(f"zip 파일명 생성 중 오류: {e}, 기본 파일명 사용")
            return f"results/{batch_id}_results.zip"
    
    def _get_summary_meta(self, batch_id: str) -> dict:
        """summary JSON 파일에서 메타데이터 추출"""
        try:
            # ai_analysis_results 폴더에서 summary 파일 찾기
            ai_results_dir = "ai_analysis_results"
            if not os.path.exists(ai_results_dir):
                return None
            
            # summary 파일 패턴으로 검색
            import glob
            summary_pattern = f"{ai_results_dir}/summary_*.json"
            summary_files = glob.glob(summary_pattern)
            
            # 최신 summary 파일 선택 (batch_id와 가장 가까운 시간)
            latest_summary = None
            latest_time = None
            
            for summary_file in summary_files:
                try:
                    with open(summary_file, 'r', encoding='utf-8') as f:
                        summary_data = json.load(f)
                        if 'summary_meta' in summary_data:
                            # 파일 생성 시간과 batch_id 시간 비교
                            file_time = os.path.getmtime(summary_file)
                            if latest_time is None or file_time > latest_time:
                                latest_time = file_time
                                latest_summary = summary_data['summary_meta']
                except Exception as e:
                    logger.warning(f"summary 파일 읽기 실패: {summary_file}, {e}")
                    continue
            
            return latest_summary
        except Exception as e:
            logger.warning(f"summary_meta 추출 중 오류: {e}")
            return None

    def _save_batch_results(self, batch_id: str):
        """배치 결과를 파일로 저장"""
        try:
            logger.info(f"배치 결과 저장 시작: {batch_id}")
            
            # results 폴더 생성
            results_dir = "results"
            if not os.path.exists(results_dir):
                os.makedirs(results_dir)
                logger.info(f"results 폴더 생성: {results_dir}")
            
            # ZIP 파일명 생성 (summary_meta 기반)
            zip_file = self._generate_zip_filename(batch_id)
            logger.info(f"ZIP 파일 생성: {zip_file}")
            
            with zipfile.ZipFile(zip_file, 'w', zipfile.ZIP_DEFLATED) as zipf:
                # ai_analysis_results 폴더의 개별 종목별 분석 결과 파일들 추가 (DOCX만)
                logger.info(f"AI 분석 결과 파일들 ZIP에 추가 시작")
                self._add_analysis_files_to_zip(zipf, batch_id)
                logger.info(f"AI 분석 결과 파일들 ZIP에 추가 완료")
                
                # 통합 파일들 ZIP에 추가 (DOCX만)
                logger.info(f"통합 분석 파일들 ZIP에 추가 시작")
                self._add_consolidated_files_to_zip(zipf, batch_id)
                logger.info(f"통합 분석 파일들 ZIP에 추가 완료")
            
            logger.info(f"ZIP 파일 생성 완료: {zip_file}")
            logger.info(f"배치 결과 저장 완료: {batch_id}")
            
            # 메모리 정리
            self.cleanup_batch_memory(batch_id)
            
        except Exception as e:
            logger.error(f"배치 결과 저장 오류: {e}")
            import traceback
            logger.error(f"상세 오류: {traceback.format_exc()}")
            # 오류 발생 시에도 메모리 정리
            self.cleanup_batch_memory(batch_id)
    
    def _add_analysis_files_to_zip(self, zipf, batch_id: str):
        """ai_analysis_results 폴더의 분석 결과 파일들을 ZIP에 추가"""
        try:
            ai_results_dir = "ai_analysis_results"
            if not os.path.exists(ai_results_dir):
                logger.warning(f"AI 분석 결과 폴더가 존재하지 않습니다: {ai_results_dir}")
                return
            
            # 배치에 포함된 종목 코드들 가져오기
            if batch_id not in self.batch_results:
                return
            
            batch_results = self.batch_results[batch_id]
            chart_type = self.batch_status.get(batch_id, {}).get('chart_type', '일봉')
            
            # 차트 타입 매핑
            chart_type_mapping = {
                "일봉": "daily",
                "주봉": "weekly", 
                "월봉": "monthly"
            }
            chart_type_en = chart_type_mapping.get(chart_type, "daily")
            
            # 각 종목별로 분석 결과 파일 찾기
            for result in batch_results:
                stock_code = result.get('stock_code')
                if not stock_code:
                    continue
                
                # 배치 결과에서 AI 분석 파일 정보 확인
                ai_analysis_file = result.get('ai_analysis_file')
                
                if ai_analysis_file:
                    # 배치 결과에 AI 분석 파일 정보가 있는 경우
                    json_path = os.path.join(ai_results_dir, ai_analysis_file)
                    if os.path.exists(json_path):
                        # 해당하는 DOCX 파일 찾기
                        docx_file = ai_analysis_file.replace('.json', '.docx')
                        docx_path = os.path.join(ai_results_dir, docx_file)
                        if os.path.exists(docx_path):
                            docx_zip_path = f"analysis_results/{docx_file}"
                            zipf.write(docx_path, docx_zip_path)
                            logger.info(f"분석 결과 파일 추가: {stock_code} - {docx_file}")
                        else:
                            logger.info(f"분석 결과 파일 추가: {stock_code} - {ai_analysis_file} (DOCX 없음)")
                    else:
                        logger.warning(f"AI 분석 JSON 파일을 찾을 수 없습니다: {json_path}")
                else:
                    # 기존 방식으로 파일 찾기 (하위 호환성) - DOCX만 추가
                    pattern = f"analysis_{chart_type_en}_{stock_code}_*.docx"
                    docx_files = []
                    for file in os.listdir(ai_results_dir):
                        if file.startswith(f"analysis_{chart_type_en}_{stock_code}_") and file.endswith('.docx'):
                            docx_files.append(file)
                    
                    # 가장 최근 파일 선택 (타임스탬프가 가장 큰 파일)
                    if docx_files:
                        latest_docx = sorted(docx_files)[-1]
                        docx_path = os.path.join(ai_results_dir, latest_docx)
                        
                        # ZIP에 DOCX 파일 추가
                        docx_zip_path = f"analysis_results/{latest_docx}"
                        zipf.write(docx_path, docx_zip_path)
                        logger.info(f"분석 결과 파일 추가: {stock_code} - {latest_docx}")
                    else:
                        logger.warning(f"종목 {stock_code}의 분석 결과 DOCX 파일을 찾을 수 없습니다")
            
        except Exception as e:
            logger.error(f"분석 결과 파일 ZIP 추가 오류: {e}")
    
    def _create_summary_content(self, batch_id: str) -> str:
        """분석 결과 요약 내용을 메모리에서 생성"""
        try:
            if batch_id not in self.batch_results:
                return ""
            
            results = self.batch_results[batch_id]
            status = self.batch_status.get(batch_id, {})
            
            summary_lines = []
            summary_lines.append("=" * 50)
            summary_lines.append("AI 주식 차트 분석 결과 요약")
            summary_lines.append("=" * 50)
            summary_lines.append("")
            
            summary_lines.append(f"배치 ID: {batch_id}")
            summary_lines.append(f"차트 유형: {status.get('chart_type', 'N/A')}")
            summary_lines.append(f"시작 시간: {status.get('start_time', 'N/A')}")
            summary_lines.append(f"완료 시간: {status.get('end_time', 'N/A')}")
            summary_lines.append(f"총 종목 수: {status.get('total', 0)}")
            summary_lines.append(f"성공: {status.get('completed', 0)}")
            summary_lines.append(f"실패: {status.get('failed', 0)}")
            summary_lines.append(f"진행률: {status.get('progress', 0):.1f}%")
            summary_lines.append("")
            
            summary_lines.append("=" * 50)
            summary_lines.append("종목별 분석 결과")
            summary_lines.append("=" * 50)
            summary_lines.append("")
            
            for i, result in enumerate(results, 1):
                summary_lines.append(f"{i}. 종목코드: {result.get('stock_code', 'N/A')}")
                summary_lines.append(f"   분석 점수: {result.get('analysis_score', 0)}")
                summary_lines.append(f"   요약: {result.get('summary', 'N/A')}")
                if 'error' in result:
                    summary_lines.append(f"   오류: {result['error']}")
                summary_lines.append("")
            
            return "\n".join(summary_lines)
            
        except Exception as e:
            logger.error(f"요약 내용 생성 오류: {e}")
            return f"요약 생성 중 오류 발생: {str(e)}"
    
    def _create_summary_file(self, batch_id: str, summary_file: str):
        """분석 결과 요약 파일 생성"""
        try:
            if batch_id not in self.batch_results:
                return
            
            results = self.batch_results[batch_id]
            status = self.batch_status.get(batch_id, {})
            
            with open(summary_file, 'w', encoding='utf-8') as f:
                f.write("=" * 50 + "\n")
                f.write("AI 주식 차트 분석 결과 요약\n")
                f.write("=" * 50 + "\n\n")
                
                f.write(f"배치 ID: {batch_id}\n")
                f.write(f"차트 유형: {status.get('chart_type', 'N/A')}\n")
                f.write(f"시작 시간: {status.get('start_time', 'N/A')}\n")
                f.write(f"완료 시간: {status.get('end_time', 'N/A')}\n")
                f.write(f"총 종목 수: {status.get('total', 0)}\n")
                f.write(f"성공: {status.get('completed', 0)}\n")
                f.write(f"실패: {status.get('failed', 0)}\n")
                f.write(f"진행률: {status.get('progress', 0):.1f}%\n\n")
                
                f.write("=" * 50 + "\n")
                f.write("종목별 분석 결과\n")
                f.write("=" * 50 + "\n\n")
                
                for i, result in enumerate(results, 1):
                    f.write(f"{i}. 종목코드: {result.get('stock_code', 'N/A')}\n")
                    f.write(f"   분석 점수: {result.get('analysis_score', 0)}\n")
                    f.write(f"   요약: {result.get('summary', 'N/A')}\n")
                    if 'error' in result:
                        f.write(f"   오류: {result['error']}\n")
                    f.write("\n")
            
        except Exception as e:
            logger.error(f"요약 파일 생성 오류: {e}")
    
    def cleanup_batch_memory(self, batch_id: str):
        """배치 완료 후 메모리 강제 정리"""
        try:
            if batch_id in self.batch_results:
                del self.batch_results[batch_id]
                logger.info(f"배치 결과 메모리 해제: {batch_id}")
            if batch_id in self.batch_status:
                del self.batch_status[batch_id]
                logger.info(f"배치 상태 메모리 해제: {batch_id}")
            if batch_id in self.batch_threads:
                del self.batch_threads[batch_id]
                logger.info(f"배치 스레드 메모리 해제: {batch_id}")
            
            logger.info(f"배치 메모리 정리 완료: {batch_id}")
            
        except Exception as e:
            logger.error(f"배치 메모리 정리 오류: {e}")

    def cleanup_batch(self, batch_id: str):
        """배치 데이터 정리 (기존 호환성 유지)"""
        self.cleanup_batch_memory(batch_id)
    
    def get_all_batches(self) -> List[str]:
        """모든 배치 ID 목록 반환"""
        return list(self.batch_status.keys())
    
    def _add_consolidated_files_to_zip(self, zipf, batch_id: str):
        """통합 분석 파일들을 ZIP에 추가 (메모리에서 생성)"""
        try:
            # 배치 상태에서 차트 타입 가져오기
            chart_type = self.batch_status.get(batch_id, {}).get('chart_type', '월봉')
            
            # 차트 타입 매핑
            chart_type_mapping = {
                "일봉": "daily",
                "주봉": "weekly", 
                "월봉": "monthly"
            }
            chart_type_en = chart_type_mapping.get(chart_type, "monthly")
            
            # 배치 결과에서 통합 파일 정보 찾기
            if batch_id not in self.batch_results:
                return
            
            batch_results = self.batch_results[batch_id]
            consolidated_info = None
            
            # 통합 파일 정보 찾기
            for result in batch_results:
                if result.get('type') == 'consolidated_info':
                    consolidated_info = result
                    break
            
            if not consolidated_info:
                logger.warning(f"배치 {batch_id}의 통합 파일 정보를 찾을 수 없습니다.")
                return
            
            timestamp = consolidated_info.get('timestamp', '')
            consolidated_result = consolidated_info.get('consolidated_result', {})
            analysis_results = consolidated_info.get('analysis_results', [])
            
            # consolidated_analysis 폴더 파일들 제거 (사용자 요청)
            
            # 4. Total Analysis JSON을 메모리에서 생성 (ZIP에는 추가하지 않음)
            try:
                total_analysis_result = self._create_total_analysis_json(batch_id, chart_type)
                if not total_analysis_result:
                    logger.error(f"Total Analysis JSON 생성 실패")
            except Exception as e:
                logger.error(f"Total Analysis JSON 생성 실패: {e}")
                total_analysis_result = None
            
            # 5. Total Analysis DOCX를 메모리에서 생성하여 ZIP에 추가
            try:
                import tempfile
                if total_analysis_result:
                    total_analysis_doc_filename = f"total_analysis/total_analysis_{chart_type_en}_{timestamp}.docx"
                    
                    # 임시 파일로 생성 후 ZIP에 추가
                    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                        temp_path = temp_file.name
                    
                    doc_success = self._create_total_analysis_docx(
                        total_analysis_result, chart_type, temp_path, batch_id
                    )
                    
                    if doc_success:
                        zipf.write(temp_path, total_analysis_doc_filename)
                        logger.info(f"Total Analysis DOCX 추가: {total_analysis_doc_filename}")
                        
                        # 임시 파일 삭제
                        os.unlink(temp_path)
                    else:
                        logger.error(f"Total Analysis DOCX 생성 실패")
                        if os.path.exists(temp_path):
                            os.unlink(temp_path)
                else:
                    logger.error(f"Total Analysis DOCX 생성 실패 (JSON 데이터 없음)")
            except Exception as e:
                logger.error(f"Total Analysis DOCX ZIP 추가 실패: {e}")
            
            # 6. 요약 분석 파일들을 ZIP에 추가
            try:
                summary_info = None
                for result in batch_results:
                    if result.get('type') == 'summary_info':
                        summary_info = result
                        break
                
                if summary_info:
                    summary_files = summary_info.get('summary_files', {})
                    if summary_files:
                        # DOCX 파일 추가
                        docx_path = summary_files.get('docx_path')
                        if docx_path and os.path.exists(docx_path):
                            summary_docx_filename = f"summary_analysis/summary_{chart_type_en}_{timestamp}.docx"
                            zipf.write(docx_path, summary_docx_filename)
                            logger.info(f"요약 분석 DOCX 추가: {summary_docx_filename}")
                    else:
                        logger.warning(f"요약 분석 파일 정보가 없습니다")
                else:
                    logger.warning(f"요약 분석 정보를 찾을 수 없습니다")
            except Exception as e:
                logger.error(f"요약 분석 파일 ZIP 추가 실패: {e}")
            
            logger.info(f"통합 파일들 ZIP 추가 완료: {chart_type_en}")
            
        except Exception as e:
            logger.error(f"통합 파일 ZIP 추가 오류: {e}")
            import traceback
            logger.error(f"상세 오류: {traceback.format_exc()}")
    
    def _create_consolidated_files(self, batch_id: str, chart_type: str):
        """배치 분석 결과를 통합하여 통합 파일들 생성 (메모리에서만 처리, ZIP에만 추가)"""
        try:
            logger.info(f"🔗 {chart_type} 통합 파일 생성 중...")
            
            # ai_chart_analysis 모듈에서 통합 함수들 import
            sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
            try:
                from ai_chart_analysis import (
                    create_consolidated_analysis,
                    create_consolidated_word_document,
                    create_summary_document
                )
            except ImportError as ie:
                logger.error(f"ai_chart_analysis 모듈 import 실패: {ie}")
                return
            
            # 배치 결과를 통합 분석용 형식으로 변환
            batch_results = self.batch_results.get(batch_id, [])
            if not batch_results:
                logger.warning(f"배치 {batch_id}의 결과가 없습니다.")
                return
            
            # 통합 분석용 결과 형식으로 변환 (통합 예시.json과 동일한 구조)
            analysis_results = []
            for result in batch_results:
                analysis_result = {
                    "stock_code": result.get("stock_code"),
                    "stock_name": result.get("stock_name", ""),
                    "success": result.get("analysis_score", 0) > 0,
                    "error": result.get("error") if result.get("analysis_score", 0) == 0 else None,
                    "timestamp": result.get("timestamp", result.get("processed_at", ""))
                }
                
                # AI 분석 데이터가 있는 경우 전체 데이터 추가 (통합 예시.json과 동일한 구조)
                if "ai_analysis_data" in result:
                    ai_data = result["ai_analysis_data"]
                    # 통합 예시.json과 동일한 구조로 데이터 매핑
                    analysis_result.update({
                        "ai_analysis_data": ai_data  # 전체 AI 분석 데이터 포함
                    })
                
                analysis_results.append(analysis_result)
            
            # 통합 분석 결과 생성
            consolidated_result = create_consolidated_analysis(analysis_results, chart_type)
            if not consolidated_result:
                logger.error("통합 분석 결과 생성 실패")
                return
            
            # 차트 타입 매핑
            chart_type_mapping = {
                "일봉": "daily",
                "주봉": "weekly", 
                "월봉": "monthly"
            }
            chart_type_en = chart_type_mapping.get(chart_type, "monthly")
            
            # 메타데이터만 생성하여 ZIP에 추가할 준비
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            
            # 통합 파일 정보를 배치 결과에 저장 (ZIP 생성 시 사용)
            if batch_id in self.batch_results:
                self.batch_results[batch_id].append({
                    'type': 'consolidated_info',
                    'chart_type': chart_type,
                    'chart_type_en': chart_type_en,
                    'timestamp': timestamp,
                    'consolidated_result': consolidated_result,
                    'analysis_results': analysis_results
                })
            
            logger.info(f"🔗 {chart_type} 통합 파일 정보 생성 완료!")
            logger.info(f"   📊 통합 JSON: consolidated_{chart_type_en}_{timestamp}.json")
            logger.info(f"   📄 통합 Word: consolidated_{chart_type_en}_{timestamp}.docx")
            logger.info(f"   📋 요약본: summary_{chart_type_en}_{timestamp}.docx")
            logger.info(f"   📊 Total Analysis: total_analysis_{chart_type_en}_{timestamp}.json")
            logger.info(f"   📄 Total Analysis DOCX: total_analysis_{chart_type_en}_{timestamp}.docx")
            
        except Exception as e:
            logger.error(f"❌ 통합 파일 생성 중 치명적 오류: {e}")
            import traceback
            logger.error(f"상세 오류: {traceback.format_exc()}")
    
    def _create_total_analysis_json(self, batch_id: str, chart_type: str) -> dict:
        """
        모든 개별 분석 결과를 순차적으로 통합하여 total_analysis.json 생성
        개별 종목의 전체 구조를 그대로 유지
        
        Args:
            batch_id (str): 배치 ID
            chart_type (str): 차트 유형
            
        Returns:
            dict: 통합된 분석 결과
        """
        try:
            if batch_id not in self.batch_results:
                return None
            
            results = self.batch_results[batch_id]
            
            # 거래정보 통계 수집
            trading_stats = {
                "거래대금_통계": {"총합": 0, "평균": 0, "최대": 0, "최소": float('inf'), "단위": "원"},
                "거래률_통계": {"총합": 0, "평균": 0, "최대": 0, "최소": float('inf'), "단위": "%"},
                "순위_통계": {"1위": 0, "10위이내": 0, "50위이내": 0, "100위이내": 0},
                "거래타입_분포": {"거래량": 0, "거래률": 0},
                "유통주식수_통계": {"총합": 0, "평균": 0, "최대": 0, "최소": float('inf'), "단위": "주"},
                "거래량_통계": {"총합": 0, "평균": 0, "최대": 0, "최소": float('inf'), "단위": "주"}
            }
            
            valid_trading_data = 0
            
            # 메타데이터
            total_analysis = {
                "metadata": {
                    "chart_type": chart_type,
                    "total_stocks": len(results),
                    "created_at": datetime.now().isoformat(),
                    "analysis_version": "1.0",
                    "file_type": "total_analysis"
                },
                "consolidated_analysis": {},
                "trading_statistics": trading_stats
            }
            
            # 각 종목별 분석 결과를 개별 종목과 동일한 구조로 추가
            for result in results:
                if "ai_analysis_data" in result:
                    ai_data = result["ai_analysis_data"]
                    stock_code = result.get("stock_code", "000000")
                    stock_name = result.get("stock_name", "알수없음")
                    
                    # 거래정보 통계 수집
                    if "종목정보" in ai_data:
                        trading_info = ai_data["종목정보"]
                        
                        # 거래대금 파싱 및 통계 수집
                        trading_amount_str = trading_info.get("거래대금", "0")
                        trading_amount = self._parse_trading_amount(trading_amount_str)
                        if trading_amount > 0:
                            trading_stats["거래대금_통계"]["총합"] += trading_amount
                            trading_stats["거래대금_통계"]["최대"] = max(trading_stats["거래대금_통계"]["최대"], trading_amount)
                            trading_stats["거래대금_통계"]["최소"] = min(trading_stats["거래대금_통계"]["최소"], trading_amount)
                        
                        # 거래률 파싱 및 통계 수집
                        turnover_rate_str = trading_info.get("거래률", "0%")
                        turnover_rate = self._parse_percentage(turnover_rate_str)
                        if turnover_rate > 0:
                            trading_stats["거래률_통계"]["총합"] += turnover_rate
                            trading_stats["거래률_통계"]["최대"] = max(trading_stats["거래률_통계"]["최대"], turnover_rate)
                            trading_stats["거래률_통계"]["최소"] = min(trading_stats["거래률_통계"]["최소"], turnover_rate)
                        
                        # 순위 통계 수집
                        ranking_str = trading_info.get("순위", "999위")
                        ranking = self._parse_ranking(ranking_str)
                        if ranking == 1:
                            trading_stats["순위_통계"]["1위"] += 1
                        elif ranking <= 10:
                            trading_stats["순위_통계"]["10위이내"] += 1
                        elif ranking <= 50:
                            trading_stats["순위_통계"]["50위이내"] += 1
                        elif ranking <= 100:
                            trading_stats["순위_통계"]["100위이내"] += 1
                        
                        # 거래타입 분포 수집
                        trading_type = trading_info.get("거래타입", "거래량")
                        if trading_type in trading_stats["거래타입_분포"]:
                            trading_stats["거래타입_분포"][trading_type] += 1
                        
                        # 유통주식수 파싱 및 통계 수집
                        shares_str = trading_info.get("유통주식수", "0주")
                        shares = self._parse_shares(shares_str)
                        if shares > 0:
                            trading_stats["유통주식수_통계"]["총합"] += shares
                            trading_stats["유통주식수_통계"]["최대"] = max(trading_stats["유통주식수_통계"]["최대"], shares)
                            trading_stats["유통주식수_통계"]["최소"] = min(trading_stats["유통주식수_통계"]["최소"], shares)
                        
                        # 거래량 파싱 및 통계 수집
                        volume_str = trading_info.get("거래량", "0주")
                        volume = self._parse_shares(volume_str)
                        if volume > 0:
                            trading_stats["거래량_통계"]["총합"] += volume
                            trading_stats["거래량_통계"]["최대"] = max(trading_stats["거래량_통계"]["최대"], volume)
                            trading_stats["거래량_통계"]["최소"] = min(trading_stats["거래량_통계"]["최소"], volume)
                        
                        valid_trading_data += 1
                    
                    # 개별 종목의 전체 구조를 그대로 유지하면서 키-값 형태로 저장
                    # 키: 종목코드, 값: 전체 AI 분석 데이터
                    total_analysis["consolidated_analysis"][stock_code] = ai_data
                    
                    # 거래타입 확인 로그
                    if "종목정보" in ai_data:
                        trading_type = ai_data["종목정보"].get("거래타입", "N/A")
                        logger.info(f"종목 {stock_code} 거래타입 확인: {trading_type}")
                    
                    # ai_analysis_file 정보 추가 (차트 이미지 찾기용)
                    if "ai_analysis_file" in result:
                        total_analysis["consolidated_analysis"][stock_code]["ai_analysis_file"] = result["ai_analysis_file"]
                    
                    logger.info(f"종목 {stock_code} ({stock_name}) 분석 데이터 추가 완료")
                elif result.get("success", False):
                    # AI 분석 데이터가 없는 경우 기본 정보라도 추가
                    stock_code = result.get("stock_code", "000000")
                    stock_name = result.get("stock_name", "알수없음")
                    
                    basic_result = {
                        "종목정보": {
                            "종목번호": stock_code,
                            "종목명": stock_name,
                            "분석일시": result.get("timestamp", result.get("processed_at", "")),
                            "차트유형": chart_type
                        },
                        "분석상태": "성공",
                        "분석점수": result.get("analysis_score", 0)
                    }
                    
                    total_analysis["consolidated_analysis"][stock_code] = basic_result
                    logger.info(f"종목 {stock_code} ({stock_name}) 기본 정보 추가 완료")
            
            # 통계 계산
            if valid_trading_data > 0:
                # 평균 계산
                trading_stats["거래대금_통계"]["평균"] = trading_stats["거래대금_통계"]["총합"] / valid_trading_data
                trading_stats["거래률_통계"]["평균"] = trading_stats["거래률_통계"]["총합"] / valid_trading_data
                trading_stats["유통주식수_통계"]["평균"] = trading_stats["유통주식수_통계"]["총합"] / valid_trading_data
                trading_stats["거래량_통계"]["평균"] = trading_stats["거래량_통계"]["총합"] / valid_trading_data
                
                # 최소값이 무한대인 경우 0으로 설정
                if trading_stats["거래대금_통계"]["최소"] == float('inf'):
                    trading_stats["거래대금_통계"]["최소"] = 0
                if trading_stats["거래률_통계"]["최소"] == float('inf'):
                    trading_stats["거래률_통계"]["최소"] = 0
                if trading_stats["유통주식수_통계"]["최소"] == float('inf'):
                    trading_stats["유통주식수_통계"]["최소"] = 0
                if trading_stats["거래량_통계"]["최소"] == float('inf'):
                    trading_stats["거래량_통계"]["최소"] = 0
            
            # 통계 정보를 메타데이터에 추가
            total_analysis["metadata"]["trading_data_count"] = valid_trading_data
            total_analysis["metadata"]["trading_statistics"] = trading_stats
            
            logger.info(f"Total Analysis JSON 생성 완료: {len(total_analysis['consolidated_analysis'])}개 종목 (거래정보: {valid_trading_data}개)")
            
            # JSON 생성 후 배치 결과 메모리 해제 (여러 차례 분석 가능하도록)
            if batch_id in self.batch_results:
                del self.batch_results[batch_id]
                logger.info(f"Total Analysis JSON 생성 후 배치 결과 메모리 해제: {batch_id}")
            
            return total_analysis
            
        except Exception as e:
            logger.error(f"Total Analysis JSON 생성 중 오류: {e}")
            # 오류 발생 시에도 메모리 해제
            if batch_id in self.batch_results:
                del self.batch_results[batch_id]
            return None
    
    def _create_total_analysis_docx(self, total_analysis_result: dict, chart_type: str, output_path: str, batch_id: str) -> bool:
        """
        Total Analysis JSON을 기반으로 DOCX 문서 생성
        개별 분석 결과와 동일한 레이아웃과 구조 적용
        
        Args:
            total_analysis_result (dict): Total Analysis JSON 데이터
            chart_type (str): 차트 유형
            output_path (str): 저장할 Word 파일 경로
            
        Returns:
            bool: 생성 성공 여부
        """
        try:
            if not total_analysis_result:
                return False
            
            # 출력 디렉토리 생성 (파일이 현재 디렉토리에 있는 경우 처리)
            output_dir = os.path.dirname(output_path)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir, exist_ok=True)
            
            # Word 문서 생성
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # 한글 폰트 설정을 위한 스타일 설정
            from docx.oxml.ns import qn
            
            # 제목 설정 (개별 분석 결과와 동일한 스타일)
            title = doc.add_heading(f'{chart_type} Total Analysis 보고서', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 제목에 한글 폰트 적용
            for run in title.runs:
                run.font.name = '맑은 고딕'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                run.font.size = Pt(14)
            
            # 분석 개요 테이블 생성
            doc.add_heading('📊 분석 개요', level=1)
            
            # 분석 개요 테이블 생성
            table = doc.add_table(rows=2, cols=3)
            table.style = 'Table Grid'
            
            # 헤더 행 설정
            header_cells = table.rows[0].cells
            header_cells[0].text = '차트 유형'
            header_cells[1].text = '분석 종목 수'
            header_cells[2].text = '생성일시'
            
            # 데이터 행 설정
            data_cells = table.rows[1].cells
            metadata = total_analysis_result.get("metadata", {})
            data_cells[0].text = metadata.get("chart_type", "N/A")
            data_cells[1].text = f'{metadata.get("total_stocks", 0)}개'
            # 날짜 형식을 YY-MM-DD (ddd) HH 형태로 변경
            from datetime import datetime
            now = datetime.now()
            weekday_korean = ['월', '화', '수', '목', '금', '토', '일']
            formatted_date = f"{now.strftime('%y.%m.%d')}({weekday_korean[now.weekday()]}) {now.strftime('%H')}시"
            data_cells[2].text = formatted_date
            
            # 테이블 전체에 한글 폰트 적용
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = '맑은 고딕'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            
            # 종목별 분석 결과 (개별 분석 결과와 동일한 구조)
            doc.add_heading('📈 종목별 상세 분석 결과', level=1)
            
            consolidated_analysis = total_analysis_result.get("consolidated_analysis", {})
            
            for i, (stock_code, stock_data) in enumerate(consolidated_analysis.items(), 1):
                # 종목 정보 (total_analysis JSON 구조에 맞게 수정)
                stock_info = stock_data.get("종목정보", {})
                stock_name = stock_info.get("종목명", f"종목{i}")
                analysis_time = stock_info.get("분석일시", "N/A")
                
                # 종목 구분선 (첫 번째 종목이 아닌 경우)
                if i > 1:
                    doc.add_paragraph("─" * 50)
                
                # 종목 제목 (개별 분석 결과와 동일한 레벨)
                heading_stock = doc.add_heading(f'{i}. {stock_name} ({stock_code})', level=1)
                for run in heading_stock.runs:
                    run.font.name = '맑은 고딕'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                
                # 종목 정보 테이블 생성 (개별 분석 결과와 동일한 구조)
                table = doc.add_table(rows=2, cols=4)
                table.style = 'Table Grid'
                
                # 헤더 행 설정
                header_cells = table.rows[0].cells
                header_cells[0].text = '종목명'
                header_cells[1].text = '종목번호'
                header_cells[2].text = '분석일시'
                header_cells[3].text = '차트유형'
                
                # 데이터 행 설정
                data_cells = table.rows[1].cells
                data_cells[0].text = stock_name
                data_cells[1].text = stock_code
                # 분석일시 사용 (없으면 현재 시간)
                if analysis_time and analysis_time != "N/A":
                    data_cells[2].text = analysis_time
                else:
                    from datetime import datetime
                    now = datetime.now()
                    weekday_korean = ['월', '화', '수', '목', '금', '토', '일']
                    formatted_date = f"{now.strftime('%y.%m.%d')}({weekday_korean[now.weekday()]}) {now.strftime('%H')}시"
                    data_cells[2].text = formatted_date
                data_cells[3].text = chart_type
                
                # 테이블 전체에 한글 폰트 적용
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = '맑은 고딕'
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                
                # 거래정보 섹션 추가
                if "거래일" in stock_info or "거래대금" in stock_info or "거래률" in stock_info:
                    doc.add_heading('💰 거래정보', level=2)
                    
                    # 거래정보 테이블 생성
                    trading_table = doc.add_table(rows=2, cols=3)
                    trading_table.style = 'Table Grid'
                    
                    # 거래정보 헤더
                    trading_header = trading_table.rows[0].cells
                    trading_header[0].text = '거래일'
                    trading_header[1].text = '거래대금'
                    trading_header[2].text = '거래률'
                    
                    # 거래정보 데이터
                    trading_data = trading_table.rows[1].cells
                    trading_data[0].text = stock_info.get("거래일", "N/A")
                    trading_data[1].text = stock_info.get("거래대금", "N/A")
                    trading_data[2].text = stock_info.get("거래률", "N/A")
                    
                    # 순위 및 기타 정보 테이블
                    if "순위" in stock_info or "유통주식수" in stock_info or "거래량" in stock_info:
                        trading_table2 = doc.add_table(rows=2, cols=3)
                        trading_table2.style = 'Table Grid'
                        
                        trading_header2 = trading_table2.rows[0].cells
                        trading_header2[0].text = '순위'
                        trading_header2[1].text = '유통주식수'
                        trading_header2[2].text = '거래량'
                        
                        trading_data2 = trading_table2.rows[1].cells
                        trading_data2[0].text = stock_info.get("순위", "N/A")
                        trading_data2[1].text = stock_info.get("유통주식수", "N/A")
                        trading_data2[2].text = stock_info.get("거래량", "N/A")
                        
                        # 거래정보 테이블들에 한글 폰트 적용
                        for table in [trading_table, trading_table2]:
                            for row in table.rows:
                                for cell in row.cells:
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.font.name = '맑은 고딕'
                                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                
                # 차트 이미지 추가
                chart_image_path = None
                
                # AI 분석 파일 정보 확인
                ai_analysis_file = stock_data.get("ai_analysis_file")
                if ai_analysis_file:
                    # 차트 타입에 따른 이미지 폴더 및 파일명 패턴
                    chart_type_mapping = {
                        "일봉": "daily",
                        "주봉": "weekly", 
                        "월봉": "monthly"
                    }
                    chart_type_en = chart_type_mapping.get(chart_type, "daily")
                    
                    # 차트 이미지 파일 찾기 (여러 폴더에서 검색)
                    chart_folders = [
                        chart_type_en + "_charts",  # daily_charts, weekly_charts, monthly_charts
                        "charts"                    # 일반 charts 폴더
                    ]
                    
                    logger.info(f"차트 이미지 검색 시작: 종목코드={stock_code}, 차트타입={chart_type}, 차트타입_en={chart_type_en}")
                    
                    for folder in chart_folders:
                        if os.path.exists(folder):
                            logger.info(f"폴더 검색 중: {folder}")
                            # 종목코드가 포함된 이미지 파일 찾기 (여러 패턴 시도)
                            for file in os.listdir(folder):
                                # 패턴 1: chart_type_en_stock_code_*.png
                                if file.startswith(f"{chart_type_en}_{stock_code}_") and file.endswith('.png'):
                                    chart_image_path = os.path.join(folder, file)
                                    logger.info(f"패턴 1으로 찾음: {file}")
                                    break
                                # 패턴 2: chart_type_en_*stock_code*.png (종목명_종목코드_날짜 형태)
                                elif f"{chart_type_en}_" in file and f"_{stock_code}_" in file and file.endswith('.png'):
                                    chart_image_path = os.path.join(folder, file)
                                    logger.info(f"패턴 2로 찾음: {file}")
                                    break
                            if chart_image_path:
                                break
                        else:
                            logger.warning(f"폴더가 존재하지 않음: {folder}")
                    
                    if not chart_image_path:
                        logger.warning(f"종목 {stock_code}의 차트 이미지를 찾을 수 없음. 검색한 폴더: {chart_folders}")
                else:
                    logger.warning(f"종목 {stock_code}의 AI 분석 파일 정보가 없음")
                
                if chart_image_path and os.path.exists(chart_image_path):
                    heading_chart = doc.add_heading('차트 이미지', level=1)
                    for run in heading_chart.runs:
                        run.font.name = '맑은 고딕'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                    
                    doc.add_picture(chart_image_path, width=Inches(6))
                    doc.add_paragraph()
                    logger.info(f"차트 이미지 추가: {stock_code} - {chart_image_path}")
                    
                    # 거래 정보 추가 (거래타입에 따라 구분)
                    if "종목정보" in stock_data:
                        trading_info = stock_data["종목정보"]
                        
                        # 거래타입 정보 가져오기 (total_analysis JSON에서 직접)
                        trading_type = trading_info.get("거래타입", "거래량")
                        
                        # 디버깅을 위한 로그
                        logger.info(f"종목 {stock_code} 거래타입: {trading_type}")
                        
                        # 차트 타입별 기간 텍스트
                        period_text = {
                            "일봉": "일일",
                            "주봉": "주간", 
                            "월봉": "월간"
                        }.get(chart_type, "일일")
                        
                        # 거래 타입별 문구 생성
                        if trading_type == "거래률" and "거래률" in trading_info:
                            turnover_rate = trading_info.get("거래률")
                            volume = trading_info.get("거래량", "N/A")
                            outstanding_shares = trading_info.get("유통주식수", "N/A")
                            ranking = trading_info.get("순위", "N/A")
                            
                            if turnover_rate and turnover_rate != "N/A":
                                # 거래률 기준 문구 - ai_chart_analysis.py 로직 참조
                                rate_value = turnover_rate.replace("%", "").strip()
                                if rate_value.replace(".", "").isdigit() and volume != "N/A" and outstanding_shares != "N/A":
                                    # 거래량과 유통주식수에서 숫자만 추출
                                    volume_clean = volume.replace(",", "").replace("주", "").strip()
                                    shares_clean = outstanding_shares.replace(",", "").replace("주", "").strip()
                                    
                                    if volume_clean.isdigit() and shares_clean.isdigit():
                                        trading_info_text = f"위 주식의 {period_text} 거래량은 {volume}로 유통주식수 {outstanding_shares} 대비 {turnover_rate}로 전체 종목 중 상위 {ranking}를 차지하여 분석 대상에 포함되었습니다."
                                    else:
                                        trading_info_text = f"위 주식의 {period_text} 거래률은 {turnover_rate}로 전체 종목 중 상위 {ranking}를 차지하여 분석 대상에 포함되었습니다."
                                else:
                                    trading_info_text = f"위 주식의 {period_text} 거래률 정보를 확인할 수 없어 분석 대상에 포함되었습니다."
                                
                                para_trading = doc.add_paragraph(trading_info_text)
                                for run in para_trading.runs:
                                    run.font.name = '맑은 고딕'
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                                doc.add_paragraph()
                        
                        else:
                            # 거래량 기준 문구 (기본값) - ai_chart_analysis.py 로직 참조
                            total_trading_amount = trading_info.get("거래대금")
                            trading_rank = trading_info.get("순위", "N/A")
                            
                            if total_trading_amount and total_trading_amount != "N/A":
                                # 거래대금 파싱 함수 사용
                                amount_numeric = self._parse_trading_amount(total_trading_amount)
                                
                                if amount_numeric > 0:
                                    amount_billion = amount_numeric / 100_000_000  # 억원 단위로 변환
                                    amount_text = f"{amount_billion:.1f}억원"
                                else:
                                    amount_text = total_trading_amount  # 원본 문자열 사용
                                
                                # 순위에서 숫자만 추출
                                if trading_rank and trading_rank != "N/A" and "위" in str(trading_rank):
                                    rank_number = str(trading_rank).replace("위", "").strip()
                                    if rank_number.isdigit():
                                        rank_text = f"상위 {rank_number}위"
                                    else:
                                        rank_text = "순위 정보 없음"
                                else:
                                    rank_text = "순위 정보 없음"
                                
                                trading_info_text = f"위 주식의 {period_text} 거래량은 {amount_text}으로 전체 종목 중 {rank_text}를 차지하여 분석 대상에 포함되었습니다."
                            else:
                                trading_info_text = f"위 주식의 {period_text} 거래대금 정보를 확인할 수 없어 분석 대상에 포함되었습니다."
                            
                            para_trading = doc.add_paragraph(trading_info_text)
                            for run in para_trading.runs:
                                run.font.name = '맑은 고딕'
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                            doc.add_paragraph()
                else:
                    # 차트 이미지가 없는 경우 안내 메시지 추가
                    heading_chart = doc.add_heading('차트 이미지', level=1)
                    for run in heading_chart.runs:
                        run.font.name = '맑은 고딕'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                    
                    # 더 구체적인 안내 메시지
                    if ai_analysis_file:
                        no_image_msg = doc.add_paragraph(f"차트 이미지를 찾을 수 없습니다. (종목코드: {stock_code}, 차트타입: {chart_type})")
                        no_image_msg2 = doc.add_paragraph(f"검색한 폴더: {chart_type_en}_charts, charts")
                        no_image_msg3 = doc.add_paragraph(f"파일명 패턴: {chart_type_en}_{stock_code}_*.png 또는 {chart_type_en}_*{stock_code}*.png")
                        
                        for p in [no_image_msg, no_image_msg2, no_image_msg3]:
                            for run in p.runs:
                                run.font.name = '맑은 고딕'
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                    else:
                        no_image_msg = doc.add_paragraph(f"차트 이미지를 찾을 수 없습니다. (AI 분석 파일 정보 없음)")
                        for run in no_image_msg.runs:
                            run.font.name = '맑은 고딕'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                    
                    logger.warning(f"차트 이미지를 찾을 수 없습니다: {stock_code}")
                
                # 차트 유형별 봉 요약 (개별 분석 결과와 동일한 구조)
                if chart_type == "일봉" and "오늘의일봉" in stock_data:
                    heading_candle = doc.add_heading('오늘의 일봉 요약', level=1)
                    for run in heading_candle.runs:
                        run.font.name = '맑은 고딕'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                    
                    candle = stock_data["오늘의일봉"]
                    p_candle1 = doc.add_paragraph(f"종가: {candle.get('종가', 'N/A')}원")
                    p_candle2 = doc.add_paragraph(f"등락률: {candle.get('등락률', 'N/A')}%")
                    p_candle3 = doc.add_paragraph(f"거래량: {candle.get('거래량', 'N/A')}주")
                    p_candle4 = doc.add_paragraph(f"주요 특징: {candle.get('주요특징', 'N/A')}")
                    
                    for p in [p_candle1, p_candle2, p_candle3, p_candle4]:
                        for run in p.runs:
                            run.font.name = '맑은 고딕'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                
                elif chart_type == "주봉" and "이번주봉" in stock_data:
                    heading_candle = doc.add_heading('이번 주 봉 요약', level=1)
                    for run in heading_candle.runs:
                        run.font.name = '맑은 고딕'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                    
                    candle = stock_data["이번주봉"]
                    p_candle1 = doc.add_paragraph(f"종가: {candle.get('종가', 'N/A')}원")
                    p_candle2 = doc.add_paragraph(f"등락률: {candle.get('등락률', 'N/A')}%")
                    p_candle3 = doc.add_paragraph(f"거래량: {candle.get('거래량', 'N/A')}주")
                    p_candle4 = doc.add_paragraph(f"주요 특징: {candle.get('주요특징', 'N/A')}")
                    
                    for p in [p_candle1, p_candle2, p_candle3, p_candle4]:
                        for run in p.runs:
                            run.font.name = '맑은 고딕'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                
                elif chart_type == "월봉" and "이번월봉" in stock_data:
                    heading_candle = doc.add_heading('이번 월봉 요약', level=1)
                    for run in heading_candle.runs:
                        run.font.name = '맑은 고딕'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                    
                    candle = stock_data["이번월봉"]
                    p_candle1 = doc.add_paragraph(f"종가: {candle.get('종가', 'N/A')}원")
                    p_candle2 = doc.add_paragraph(f"등락률: {candle.get('등락률', 'N/A')}%")
                    p_candle3 = doc.add_paragraph(f"거래량: {candle.get('거래량', 'N/A')}주")
                    p_candle4 = doc.add_paragraph(f"주요 특징: {candle.get('주요특징', 'N/A')}")
                    
                    for p in [p_candle1, p_candle2, p_candle3, p_candle4]:
                        for run in p.runs:
                            run.font.name = '맑은 고딕'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                
                # 종합 분석 점수 (개별 분석 결과와 동일한 구조)
                if "종합분석점수" in stock_data:
                    heading_score = doc.add_heading('종합 분석 점수', level=1)
                    for run in heading_score.runs:
                        run.font.name = '맑은 고딕'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                    
                    score = stock_data["종합분석점수"]
                    p_score1 = doc.add_paragraph(f"점수: {score.get('점수', 'N/A')}/100")
                    p_score2 = doc.add_paragraph(f"요약: {score.get('요약', 'N/A')}")
                    
                    # 점수 강조 (개별 분석 결과와 동일한 스타일)
                    p_score1.runs[0].bold = True
                    p_score1.runs[0].font.size = Pt(14)
                    
                    for p in [p_score1, p_score2]:
                        for run in p.runs:
                            run.font.name = '맑은 고딕'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                
                # 핵심 기술적 분석 지표 (개별 분석 결과와 동일한 구조)
                if "핵심기술적지표" in stock_data:
                    heading_tech = doc.add_heading('핵심 기술적 분석 지표', level=1)
                    for run in heading_tech.runs:
                        run.font.name = '맑은 고딕'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                    
                    tech = stock_data["핵심기술적지표"]
                    
                    # 일봉 차트 지표들
                    if chart_type == "일봉":
                        if "이동평균선정배열" in tech:
                            p_tech1 = doc.add_paragraph(f"이동평균선 정배열: {tech.get('이동평균선정배열', 'N/A')}")
                        if "골든데드크로스" in tech:
                            p_tech2 = doc.add_paragraph(f"골든/데드 크로스: {tech.get('골든데드크로스', 'N/A')}")
                        if "MACD상태" in tech:
                            p_tech3 = doc.add_paragraph(f"MACD 상태: {tech.get('MACD상태', 'N/A')}")
                        if "RSI상태" in tech:
                            p_tech4 = doc.add_paragraph(f"RSI 상태: {tech.get('RSI상태', 'N/A')}")
                        if "볼린저밴드" in tech:
                            p_tech5 = doc.add_paragraph(f"볼린저밴드: {tech.get('볼린저밴드', 'N/A')}")
                    
                    # 주봉 차트 지표들
                    elif chart_type == "주봉":
                        if "이동평균선정배열" in tech:
                            p_tech1 = doc.add_paragraph(f"이동평균선 정배열: {tech.get('이동평균선정배열', 'N/A')}")
                        if "골든데드크로스" in tech:
                            p_tech2 = doc.add_paragraph(f"골든/데드 크로스: {tech.get('골든데드크로스', 'N/A')}")
                        if "Stochastic상태" in tech:
                            p_tech3 = doc.add_paragraph(f"Stochastic 상태: {tech.get('Stochastic상태', 'N/A')}")
                        if "볼린저밴드" in tech:
                            p_tech4 = doc.add_paragraph(f"볼린저밴드: {tech.get('볼린저밴드', 'N/A')}")
                    
                    # 월봉 차트 지표들
                    elif chart_type == "월봉":
                        if "장기정배열" in tech:
                            p_tech1 = doc.add_paragraph(f"장기 정배열: {tech.get('장기정배열', 'N/A')}")
                        if "CCI상태" in tech:
                            p_tech2 = doc.add_paragraph(f"CCI 상태: {tech.get('CCI상태', 'N/A')}")
                        if "ADX상태" in tech:
                            p_tech3 = doc.add_paragraph(f"ADX 상태: {tech.get('ADX상태', 'N/A')}")
                        if "주요이동평균선" in tech:
                            p_tech4 = doc.add_paragraph(f"주요 이동평균선: {tech.get('주요이동평균선', 'N/A')}")
                    
                    # 한글 폰트 적용
                    tech_paragraphs = []
                    for j in range(1, 6):
                        if f'p_tech{j}' in locals():
                            tech_paragraphs.append(locals()[f'p_tech{j}'])
                    
                    for p in tech_paragraphs:
                        for run in p.runs:
                            run.font.name = '맑은 고딕'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                
                # 세부 분석 (개별 분석 결과와 동일한 구조)
                if "세부분석" in stock_data:
                    heading_detail = doc.add_heading('세부 분석', level=1)
                    for run in heading_detail.runs:
                        run.font.name = '맑은 고딕'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                    
                    detail = stock_data["세부분석"]
                    
                    # 가격 및 거래량 분석
                    if "가격및거래량" in detail:
                        sub_heading1 = doc.add_heading('가격 및 거래량', level=2)
                        for run in sub_heading1.runs:
                            run.font.name = '맑은 고딕'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                        
                        price_vol = detail["가격및거래량"]
                        if "거래량비교" in price_vol:
                            p_detail1 = doc.add_paragraph(f"거래량 비교: {price_vol.get('거래량비교', 'N/A')}")
                        if "주요가격대" in price_vol:
                            p_detail2 = doc.add_paragraph(f"주요 가격대: {price_vol.get('주요가격대', 'N/A')}")
                        if "박스권분석" in price_vol:
                            p_detail3 = doc.add_paragraph(f"박스권 분석: {price_vol.get('박스권분석', 'N/A')}")
                        if "역사적고점저점" in price_vol:
                            p_detail4 = doc.add_paragraph(f"역사적 고점/저점: {price_vol.get('역사적고점저점', 'N/A')}")
                    
                    # 이동평균선 분석
                    if "이동평균선" in detail:
                        sub_heading2 = doc.add_heading('이동평균선', level=2)
                        for run in sub_heading2.runs:
                            run.font.name = '맑은 고딕'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                        
                        ma = detail["이동평균선"]
                        if "현재가위치" in ma:
                            p_detail5 = doc.add_paragraph(f"현재가 위치: {ma.get('현재가위치', 'N/A')}")
                        if "밀집도" in ma:
                            p_detail6 = doc.add_paragraph(f"밀집도: {ma.get('밀집도', 'N/A')}")
                        if "20주선역할" in ma:
                            p_detail7 = doc.add_paragraph(f"20주선 역할: {ma.get('20주선역할', 'N/A')}")
                        if "20개월선역할" in ma:
                            p_detail8 = doc.add_paragraph(f"20개월선 역할: {ma.get('20개월선역할', 'N/A')}")
                    
                    # 모멘텀 분석
                    if "모멘텀" in detail:
                        sub_heading3 = doc.add_heading('모멘텀 및 강도', level=2)
                        for run in sub_heading3.runs:
                            run.font.name = '맑은 고딕'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                        
                        momentum = detail["모멘텀"]
                        if "MACD분석" in momentum:
                            p_detail9 = doc.add_paragraph(f"MACD 분석: {momentum.get('MACD분석', 'N/A')}")
                        if "RSI분석" in momentum:
                            p_detail10 = doc.add_paragraph(f"RSI 분석: {momentum.get('RSI분석', 'N/A')}")
                        if "Stochastic분석" in momentum:
                            p_detail11 = doc.add_paragraph(f"Stochastic 분석: {momentum.get('Stochastic분석', 'N/A')}")
                        if "볼린저밴드분석" in momentum:
                            p_detail12 = doc.add_paragraph(f"볼린저밴드 분석: {momentum.get('볼린저밴드분석', 'N/A')}")
                        if "CCI분석" in momentum:
                            p_detail13 = doc.add_paragraph(f"CCI 분석: {momentum.get('CCI분석', 'N/A')}")
                        if "ADX분석" in momentum:
                            p_detail14 = doc.add_paragraph(f"ADX 분석: {momentum.get('ADX분석', 'N/A')}")
                    
                    # 세부 분석 한글 폰트 적용
                    detail_paragraphs = []
                    for j in range(1, 15):
                        if f'p_detail{j}' in locals():
                            detail_paragraphs.append(locals()[f'p_detail{j}'])
                    
                    for p in detail_paragraphs:
                        for run in p.runs:
                            run.font.name = '맑은 고딕'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                
                # 투자 아이디어 (total_analysis JSON 구조에 맞게 수정)
                if chart_type == "일봉" and "단기투자아이디어" in stock_data:
                    heading_idea = doc.add_heading('단기 투자 아이디어', level=1)
                    for run in heading_idea.runs:
                        run.font.name = '맑은 고딕'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                    
                    idea = stock_data["단기투자아이디어"]
                    if "추세요약" in idea:
                        p_idea1 = doc.add_paragraph(f"추세 요약: {idea.get('추세요약', 'N/A')}")
                    if "매매시그널" in idea:
                        p_idea2 = doc.add_paragraph(f"매매 시그널: {idea.get('매매시그널', 'N/A')}")
                        p_idea2.runs[0].bold = True
                        p_idea2.runs[0].font.size = Pt(14)
                    
                    for p in [p_idea1, p_idea2]:
                        for run in p.runs:
                            run.font.name = '맑은 고딕'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                
                elif chart_type == "주봉" and "중기투자아이디어" in stock_data:
                    heading_idea = doc.add_heading('중기 투자 아이디어', level=1)
                    for run in heading_idea.runs:
                        run.font.name = '맑은 고딕'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                    
                    idea = stock_data["중기투자아이디어"]
                    if "추세요약" in idea:
                        p_idea1 = doc.add_paragraph(f"추세 요약: {idea.get('추세요약', 'N/A')}")
                    if "매매시그널" in idea:
                        p_idea2 = doc.add_paragraph(f"매매 시그널: {idea.get('매매시그널', 'N/A')}")
                        p_idea2.runs[0].bold = True
                        p_idea2.runs[0].font.size = Pt(14)
                    
                    for p in [p_idea1, p_idea2]:
                        for run in p.runs:
                            run.font.name = '맑은 고딕'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                
                elif chart_type == "월봉" and "장기투자아이디어" in stock_data:
                    heading_idea = doc.add_heading('장기 투자 아이디어', level=1)
                    for run in heading_idea.runs:
                        run.font.name = '맑은 고딕'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                    
                    idea = stock_data["장기투자아이디어"]
                    if "사이클요약" in idea:
                        p_idea1 = doc.add_paragraph(f"사이클 요약: {idea.get('사이클요약', 'N/A')}")
                    if "투자전략" in idea:
                        p_idea2 = doc.add_paragraph(f"투자 전략: {idea.get('투자전략', 'N/A')}")
                        p_idea2.runs[0].bold = True
                        p_idea2.runs[0].font.size = Pt(14)
                    
                    for p in [p_idea1, p_idea2]:
                        for run in p.runs:
                            run.font.name = '맑은 고딕'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                
                # AI 분석 결과 추가 (total_analysis JSON 구조에 맞게 수정)
                if "AI분석결과" in stock_data:
                    # AI 분석 결과 제목 (종목명 포함)
                    heading_analysis = doc.add_heading(f'{stock_name} 차트 분석 결과', level=1)
                    for run in heading_analysis.runs:
                        run.font.name = '맑은 고딕'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                    
                    # AI 분석 결과를 마크다운 파싱하여 추가
                    ai_analysis_result = stock_data.get("AI분석결과", "")
                    if ai_analysis_result:
                        self._parse_markdown_to_word(ai_analysis_result, doc)
                    else:
                        # AI 분석 결과가 없는 경우
                        para = doc.add_paragraph("AI 분석 결과가 없습니다.")
                        for run in para.runs:
                            run.font.name = '맑은 고딕'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            
            # 거래정보 통계 섹션 추가
            metadata = total_analysis_result.get("metadata", {})
            trading_stats = metadata.get("trading_statistics", {})
            
            if trading_stats and any(trading_stats.values()):
                doc.add_heading('📊 거래정보 통계', level=1)
                
                # 거래대금 통계
                if trading_stats.get("거래대금_통계", {}).get("총합", 0) > 0:
                    doc.add_heading('💰 거래대금 통계', level=2)
                    
                    amount_stats = trading_stats["거래대금_통계"]
                    p_amount1 = doc.add_paragraph(f"총 거래대금: {amount_stats.get('총합', 0):,.0f}원")
                    p_amount2 = doc.add_paragraph(f"평균 거래대금: {amount_stats.get('평균', 0):,.0f}원")
                    p_amount3 = doc.add_paragraph(f"최대 거래대금: {amount_stats.get('최대', 0):,.0f}원")
                    p_amount4 = doc.add_paragraph(f"최소 거래대금: {amount_stats.get('최소', 0):,.0f}원")
                    
                    for p in [p_amount1, p_amount2, p_amount3, p_amount4]:
                        for run in p.runs:
                            run.font.name = '맑은 고딕'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                
                # 거래률 통계
                if trading_stats.get("거래률_통계", {}).get("총합", 0) > 0:
                    doc.add_heading('📈 거래률 통계', level=2)
                    
                    rate_stats = trading_stats["거래률_통계"]
                    p_rate1 = doc.add_paragraph(f"평균 거래률: {rate_stats.get('평균', 0):.2f}%")
                    p_rate2 = doc.add_paragraph(f"최대 거래률: {rate_stats.get('최대', 0):.2f}%")
                    p_rate3 = doc.add_paragraph(f"최소 거래률: {rate_stats.get('최소', 0):.2f}%")
                    
                    for p in [p_rate1, p_rate2, p_rate3]:
                        for run in p.runs:
                            run.font.name = '맑은 고딕'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                
                # 순위 통계
                if trading_stats.get("순위_통계", {}):
                    doc.add_heading('🏆 순위 통계', level=2)
                    
                    ranking_stats = trading_stats["순위_통계"]
                    p_rank1 = doc.add_paragraph(f"1위: {ranking_stats.get('1위', 0)}개 종목")
                    p_rank2 = doc.add_paragraph(f"10위 이내: {ranking_stats.get('10위이내', 0)}개 종목")
                    p_rank3 = doc.add_paragraph(f"50위 이내: {ranking_stats.get('50위이내', 0)}개 종목")
                    p_rank4 = doc.add_paragraph(f"100위 이내: {ranking_stats.get('100위이내', 0)}개 종목")
                    
                    for p in [p_rank1, p_rank2, p_rank3, p_rank4]:
                        for run in p.runs:
                            run.font.name = '맑은 고딕'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                
                # 거래타입 분포
                if trading_stats.get("거래타입_분포", {}):
                    doc.add_heading('📋 거래타입 분포', level=2)
                    
                    type_stats = trading_stats["거래타입_분포"]
                    p_type1 = doc.add_paragraph(f"거래량 기준: {type_stats.get('거래량', 0)}개 종목")
                    p_type2 = doc.add_paragraph(f"거래률 기준: {type_stats.get('거래률', 0)}개 종목")
                    
                    for p in [p_type1, p_type2]:
                        for run in p.runs:
                            run.font.name = '맑은 고딕'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            
            # 문서 저장
            doc.save(output_path)
            logger.info(f"Total Analysis DOCX 생성 완료: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"Total Analysis DOCX 생성 중 오류: {e}")
            return False
    
    def _parse_markdown_to_word(self, text: str, doc) -> None:
        """
        마크다운 텍스트를 Word 문서 형식으로 파싱하여 추가
        
        Args:
            text (str): 마크다운 형태의 텍스트
            doc: Word 문서 객체
        """
        try:
            import re
            from docx.oxml.ns import qn
            
            # **내용** 형태를 볼드체로 변환하는 정규식
            bold_pattern = r'\*\*(.*?)\*\*'
            
            # 줄바꿈으로 분리
            lines = text.split('\n')
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # **내용** 패턴을 찾아서 볼드체로 변환
                if re.search(bold_pattern, line):
                    # 볼드체가 포함된 줄 처리
                    para = doc.add_paragraph()
                    
                    # **내용** 패턴으로 분할
                    parts = re.split(bold_pattern, line)
                    
                    for i, part in enumerate(parts):
                        if i % 2 == 0:
                            # 일반 텍스트
                            if part:
                                run = para.add_run(part)
                                run.font.name = '맑은 고딕'
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                        else:
                            # 볼드체 텍스트
                            if part:
                                run = para.add_run(part)
                                run.font.name = '맑은 고딕'
                                run.bold = True
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                else:
                    # 볼드체가 없는 일반 줄 처리
                    if line.startswith('###'):
                        # ### 헤더 처리 (### 제거하고 제목으로 변환)
                        heading_text = line.replace('###', '').strip()
                        if heading_text:
                            para = doc.add_heading(heading_text, level=3)
                    elif line.startswith('##'):
                        # ## 헤더 처리 (## 제거하고 제목으로 변환)
                        heading_text = line.replace('##', '').strip()
                        if heading_text:
                            para = doc.add_heading(heading_text, level=2)
                    elif line.startswith('#'):
                        # # 헤더 처리 (# 제거하고 제목으로 변환)
                        heading_text = line.replace('#', '').strip()
                        if heading_text:
                            para = doc.add_heading(heading_text, level=1)
                    elif line.startswith('📊') or line.startswith('📈') or line.startswith('📉') or line.startswith('💡') or line.startswith('🔍') or line.startswith('⚠️') or line.startswith('✅'):
                        para = doc.add_heading(line, level=2)
                    elif line.startswith('•') or line.startswith('-') or line.startswith('*'):
                        para = doc.add_paragraph(line)
                    else:
                        para = doc.add_paragraph(line)
                    
                    # 한글 폰트 적용
                    for run in para.runs:
                        run.font.name = '맑은 고딕'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                        
        except Exception as e:
            logger.error(f"⚠️ 마크다운 파싱 중 오류: {e}")
            # 오류 발생 시 일반 텍스트로 처리
            para = doc.add_paragraph(text)
            for run in para.runs:
                run.font.name = '맑은 고딕'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    
    def _create_summary_analysis(self, chart_type: str, specific_files: list = None) -> dict:
        """
        배치 분석 완료 후 요약 분석 실행
        
        Args:
            chart_type (str): 차트 유형 (한글)
            specific_files (list): 특정 파일 경로들 (None이면 전체 폴더 스캔)
            
        Returns:
            dict: 생성된 요약 파일 정보 {"json_path": "...", "docx_path": "..."}
        """
        try:
            logger.info(f"요약 분석 시작: chart_type={chart_type}")
            
            # ai_chart_analysis 모듈에서 SummaryFileGenerator import
            sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
            from ai_chart_analysis import SummaryFileGenerator
            from database_config import get_db_config
            
            # 차트 유형을 영문으로 변환
            chart_type_en = {
                "일봉": "daily",
                "주봉": "weekly", 
                "월봉": "monthly"
            }.get(chart_type, chart_type.lower())
            
            # 데이터베이스 설정 로드
            db_config = get_db_config()
            
            # 요약 파일 생성기 초기화
            summary_generator = SummaryFileGenerator(db_config)
            
            # 통합 요약 파일 생성 (차트 유형별)
            if specific_files:
                # 특정 파일들만 사용하여 요약 생성
                result = summary_generator.create_consolidated_summary_from_files(chart_type_en, specific_files)
            else:
                # 전체 폴더 스캔하여 요약 생성 (기존 방식)
                result = summary_generator.create_consolidated_summary_by_type(chart_type_en)
            
            if result:
                logger.info(f"요약 분석 완료: chart_type={chart_type}")
                logger.info(f"JSON 파일: {result.get('json_path', 'N/A')}")
                logger.info(f"DOCX 파일: {result.get('docx_path', 'N/A')}")
                return result
            else:
                logger.warning(f"요약 분석 실패 또는 분석할 데이터 없음: chart_type={chart_type}")
                return None
            
        except Exception as e:
            logger.error(f"요약 분석 중 오류 발생: {e}")
            logger.info("개별 분석 결과는 정상적으로 저장되었습니다.")
            return None 
    
    def _parse_trading_amount(self, amount_str: str) -> float:
        """거래대금 문자열을 숫자로 변환 (예: '5263억원' -> 526300000000)"""
        try:
            if not amount_str or amount_str == "N/A":
                return 0
            
            # 숫자와 단위 추출
            import re
            # 쉼표 제거 후 숫자와 단위 추출
            clean_str = amount_str.replace(',', '')
            match = re.search(r'([\d.]+)(억원|만원|원)?', clean_str)
            if not match:
                return 0
            
            number = float(match.group(1))
            unit = match.group(2) or "원"
            
            if unit == "억원":
                return number * 100000000
            elif unit == "만원":
                return number * 10000
            else:  # "원"
                return number
        except:
            return 0
    
    def _parse_percentage(self, percent_str: str) -> float:
        """퍼센트 문자열을 숫자로 변환 (예: '1250.18%' -> 1250.18)"""
        try:
            if not percent_str or percent_str == "N/A":
                return 0
            
            import re
            match = re.search(r'([\d.]+)%', percent_str)
            if match:
                return float(match.group(1))
            return 0
        except:
            return 0
    
    def _parse_ranking(self, ranking_str: str) -> int:
        """순위 문자열을 숫자로 변환 (예: '1위' -> 1)"""
        try:
            if not ranking_str or ranking_str == "N/A":
                return 999
            
            import re
            match = re.search(r'(\d+)위', ranking_str)
            if match:
                return int(match.group(1))
            return 999
        except:
            return 999
    
    def _parse_shares(self, shares_str: str) -> int:
        """주식수 문자열을 숫자로 변환 (예: '16,366,428주' -> 16366428)"""
        try:
            if not shares_str or shares_str == "N/A":
                return 0
            
            import re
            match = re.search(r'([\d,]+)주', shares_str.replace(',', ''))
            if match:
                return int(match.group(1))
            return 0
        except:
            return 0