#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
기획자 관점 바이브코딩 포트폴리오 생성기
AI 주식 차트 분석 플랫폼 프로젝트
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os
from datetime import datetime

def create_portfolio_document():
    """기획자 관점 포트폴리오 DOCX 문서 생성"""
    
    # 새 문서 생성
    doc = Document()
    
    # 제목 설정
    title = doc.add_heading('AI 기반 주식 차트 분석 플랫폼', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 부제목
    subtitle = doc.add_heading('서비스 기획자 관점의 바이브코딩 포트폴리오', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 작성자 정보
    author_info = doc.add_paragraph()
    author_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_info.add_run(f"작성일: {datetime.now().strftime('%Y년 %m월 %d일')}").bold = True
    
    doc.add_page_break()
    
    # 1. 프로젝트 개요
    doc.add_heading('1. 프로젝트 개요', level=1)
    
    overview_content = """
본 프로젝트는 서비스 기획자로서 Cursor AI를 활용한 바이브코딩을 통해 개발한 
AI 기반 주식 차트 분석 플랫폼입니다. 

개발자 지식 없이도 AI와의 협업을 통해 복잡한 금융 데이터 처리, 
웹 애플리케이션 개발, 데이터베이스 설계까지 완성도 높은 서비스를 구현했습니다.
"""
    
    doc.add_paragraph(overview_content)
    
    # 2. 서비스 기획 배경
    doc.add_heading('2. 서비스 기획 배경', level=1)
    
    background_content = """
### 2.1 문제 인식
- 개인 투자자들이 주식 차트를 읽고 기술적 분석을 수행하는 데 어려움
- 전문적인 분석 도구는 고가이고 복잡함
- AI 기술을 활용한 접근성 높은 분석 도구 필요

### 2.2 솔루션 기획
- Google Gemini AI를 활용한 차트 이미지 분석
- 웹 기반 사용자 친화적 인터페이스
- 일봉/주봉/월봉 다중 시간대 분석 지원
- 배치 처리로 대량 종목 분석 가능
- 실시간 데이터 수집 및 자동화 시스템
"""
    
    doc.add_paragraph(background_content)
    
    # 3. 서비스 아키텍처
    doc.add_heading('3. 서비스 아키텍처', level=1)
    
    architecture_content = """
### 3.1 전체 시스템 구조
```
┌─────────────────┐    ┌─────────────────┐    ┌─────────────────┐
│   웹 인터페이스   │    │   API 서버      │    │   데이터베이스   │
│   (Flask)       │◄──►│   (REST API)    │◄──►│   (MySQL)       │
└─────────────────┘    └─────────────────┘    └─────────────────┘
         │                       │                       │
         │              ┌─────────────────┐              │
         │              │   AI 분석 엔진   │              │
         │              │   (Gemini)      │              │
         │              └─────────────────┘              │
         │                       │                       │
         │              ┌─────────────────┐              │
         └─────────────►│   데이터 수집기   │◄─────────────┘
                        │   (Yahoo Finance)│
                        └─────────────────┘
```

### 3.2 핵심 모듈 구조
- **웹 인터페이스**: Flask 기반 사용자 인터페이스
- **API 서버**: RESTful API를 통한 백엔드 서비스
- **AI 분석 엔진**: Google Gemini를 활용한 차트 분석
- **데이터 수집기**: Yahoo Finance API 기반 실시간 데이터 수집
- **데이터베이스**: MySQL 기반 구조화된 데이터 저장
- **배치 처리기**: 대량 데이터 처리 및 스케줄링
"""
    
    doc.add_paragraph(architecture_content)
    
    # 4. 주요 기능
    doc.add_heading('4. 주요 기능', level=1)
    
    features_content = """
### 4.1 차트 분석 기능
- **다중 시간대 분석**: 일봉(240일), 주봉(240주), 월봉(120개월)
- **기술적 지표**: 이동평균선, RSI, MACD, 볼린저밴드, ADX 등
- **AI 기반 해석**: 차트 패턴 인식 및 투자 아이디어 제안

### 4.2 데이터 관리 기능
- **자동 데이터 수집**: 10년간 일봉 데이터 자동 수집
- **데이터 품질 관리**: 장중/장마감 데이터 구분 및 검증
- **증분 업데이트**: 효율적인 데이터 업데이트 시스템

### 4.3 웹 인터페이스 기능
- **단일 분석**: 개별 종목 차트 업로드 및 분석
- **대량 분석**: 다수 종목 일괄 처리
- **거래량 랭킹**: 일/주/월별 거래량 상위 50개 종목 조회
- **프롬프트 관리**: AI 분석 프롬프트 동적 관리

### 4.4 자동화 기능
- **배치 스케줄러**: 정기적인 분석 작업 자동 실행
- **일일 데이터 수집**: 장 마감 후 자동 데이터 수집
- **진행률 추적**: 실시간 작업 상태 모니터링
"""
    
    doc.add_paragraph(features_content)
    
    # 5. 기술 스택
    doc.add_heading('5. 기술 스택', level=1)
    
    tech_stack_content = """
### 5.1 백엔드 기술
- **Python 3.7+**: 메인 개발 언어
- **Flask**: 웹 프레임워크
- **MySQL**: 데이터베이스
- **pandas, numpy**: 데이터 처리
- **yfinance**: 주식 데이터 수집

### 5.2 AI 및 분석
- **Google Gemini AI**: 차트 이미지 분석
- **matplotlib, mplfinance**: 차트 생성
- **TA-Lib**: 기술적 지표 계산

### 5.3 프론트엔드
- **HTML5, CSS3, JavaScript**: 웹 인터페이스
- **반응형 디자인**: 모바일/데스크톱 지원

### 5.4 인프라 및 도구
- **Git**: 버전 관리
- **REST API**: 서비스 간 통신
- **APScheduler**: 배치 작업 스케줄링
"""
    
    doc.add_paragraph(tech_stack_content)
    
    # 6. 데이터베이스 설계
    doc.add_heading('6. 데이터베이스 설계', level=1)
    
    db_design_content = """
### 6.1 핵심 테이블 구조
- **stocks**: 종목 정보 (종목코드, 종목명, 시장구분, 시가총액)
- **daily_data**: 일봉 데이터 (OHLCV, 거래일)
- **weekly_data**: 주봉 데이터 (보조지표 포함)
- **monthly_data**: 월봉 데이터 (보조지표 포함)
- **technical_indicators**: 기술적 지표 데이터

### 6.2 관리 테이블
- **data_collection_log**: 데이터 수집 이력
- **stock_collection_status**: 종목별 수집 상태
- **prompts**: AI 분석 프롬프트 관리
- **secure_configs**: 보안 설정 (API 키 등)

### 6.3 데이터 품질 관리
- 장중/장마감 데이터 구분
- 자동 데이터 검증 및 재수집
- 변경 이력 추적 시스템
"""
    
    doc.add_paragraph(db_design_content)
    
    # 7. 서비스 플로우
    doc.add_heading('7. 서비스 플로우', level=1)
    
    service_flow_content = """
### 7.1 단일 분석 플로우
```
사용자 차트 업로드 → 종목 정보 입력 → AI 분석 요청 → 
차트 이미지 분석 → 기술적 지표 계산 → 분석 결과 생성 → 
JSON/Word 문서 저장 → 사용자에게 결과 제공
```

### 7.2 대량 분석 플로우
```
종목 리스트 업로드 → 배치 작업 큐 등록 → 병렬 처리 시작 → 
각 종목별 차트 생성 → AI 분석 수행 → 진행률 추적 → 
결과 파일 생성 → 완료 알림 및 다운로드 제공
```

### 7.3 데이터 수집 플로우
```
스케줄러 트리거 → 시장 상태 확인 → 종목별 데이터 수집 → 
데이터 품질 검증 → 데이터베이스 저장 → 수집 상태 업데이트 → 
로그 기록
```
"""
    
    doc.add_paragraph(service_flow_content)
    
    # 8. 바이브코딩 경험
    doc.add_heading('8. 바이브코딩 경험', level=1)
    
    vibecoding_content = """
### 8.1 Cursor AI 활용 방식
- **코드 생성**: 복잡한 알고리즘과 데이터 처리 로직 자동 생성
- **디버깅 지원**: 오류 발생 시 빠른 문제 해결
- **코드 리팩토링**: 기존 코드 최적화 및 개선
- **문서화**: README, API 문서 자동 생성

### 8.2 기획자 관점의 개발 접근
- **요구사항 정의**: 사용자 관점에서 기능 명세
- **시스템 설계**: 모듈간 의존성과 데이터 흐름 설계
- **사용자 경험**: 직관적인 웹 인터페이스 기획
- **품질 관리**: 데이터 검증 및 에러 처리 로직 설계

### 8.3 학습 및 성장
- **Python 기초**: 프로그래밍 언어 기초 학습
- **웹 개발**: Flask 프레임워크 활용법 습득
- **데이터베이스**: MySQL 스키마 설계 및 쿼리 작성
- **AI API**: Google Gemini API 연동 방법 학습
"""
    
    doc.add_paragraph(vibecoding_content)
    
    # 9. 프로젝트 성과
    doc.add_heading('9. 프로젝트 성과', level=1)
    
    achievements_content = """
### 9.1 기술적 성과
- **완전 자동화**: 데이터 수집부터 분석까지 전 과정 자동화
- **확장성**: 대량 데이터 처리 가능한 아키텍처 구축
- **안정성**: 에러 처리 및 데이터 검증 시스템 구축
- **사용성**: 직관적인 웹 인터페이스 제공

### 9.2 비즈니스 가치
- **접근성**: 개인 투자자도 전문적 분석 도구 활용 가능
- **효율성**: 수동 분석 시간 대폭 단축
- **정확성**: AI 기반 객관적 분석 결과 제공
- **확장성**: 추가 기능 개발 및 서비스 확장 가능

### 9.3 학습 성과
- **기술 스택**: 풀스택 개발 경험 습득
- **문제 해결**: 복잡한 비즈니스 로직 구현 능력 향상
- **시스템 설계**: 확장 가능한 아키텍처 설계 경험
- **AI 활용**: AI 기술을 활용한 서비스 개발 경험
"""
    
    doc.add_paragraph(achievements_content)
    
    # 10. 향후 계획
    doc.add_heading('10. 향후 계획', level=1)
    
    future_plans_content = """
### 10.1 기능 확장
- **실시간 알림**: 매매 신호 발생 시 알림 서비스
- **포트폴리오 분석**: 종합적인 포트폴리오 관리 기능
- **백테스팅**: 과거 데이터를 활용한 전략 검증
- **모바일 앱**: iOS/Android 네이티브 앱 개발

### 10.2 기술 개선
- **성능 최적화**: 대량 데이터 처리 속도 향상
- **보안 강화**: API 키 암호화 및 접근 제어
- **모니터링**: 시스템 상태 실시간 모니터링
- **CI/CD**: 자동 배포 및 테스트 파이프라인 구축

### 10.3 서비스 확장
- **다국가 지원**: 해외 주식 시장 데이터 추가
- **API 서비스**: 외부 개발자용 API 제공
- **구독 모델**: 프리미엄 기능 유료화
- **파트너십**: 증권사와의 데이터 연동
"""
    
    doc.add_paragraph(future_plans_content)
    
    # 11. 결론
    doc.add_heading('11. 결론', level=1)
    
    conclusion_content = """
본 프로젝트를 통해 서비스 기획자로서 Cursor AI를 활용한 바이브코딩의 
가능성을 확인할 수 있었습니다. 

개발 지식이 부족한 상황에서도 AI와의 협업을 통해 복잡한 금융 데이터 처리, 
웹 애플리케이션 개발, 데이터베이스 설계까지 완성도 높은 서비스를 구현했습니다.

이 경험을 통해 얻은 인사이트:
- **AI 활용**: 개발 도구로서 AI의 강력한 능력 확인
- **기획 역량**: 기술적 구현보다는 서비스 기획의 중요성 인식
- **학습 능력**: 새로운 기술 스택 빠른 습득 및 적용
- **문제 해결**: 복잡한 비즈니스 요구사항을 기술적 솔루션으로 구현

앞으로도 AI 기술을 활용하여 더욱 혁신적인 서비스를 기획하고 구현해나가겠습니다.
"""
    
    doc.add_paragraph(conclusion_content)
    
    # 문서 저장
    filename = f"AI_주식차트분석_바이브코딩_포트폴리오_{datetime.now().strftime('%Y%m%d')}.docx"
    doc.save(filename)
    
    print(f"✅ 포트폴리오 문서가 생성되었습니다: {filename}")
    return filename

if __name__ == "__main__":
    create_portfolio_document()
