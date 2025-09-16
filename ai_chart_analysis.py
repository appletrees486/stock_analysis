#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
AI 제미나이를 활용한 주식 차트 분석 스크립트 (일봉/주봉/월봉 지원)
"""

import os
import json
import base64
import time
import google.generativeai as genai
from PIL import Image
import requests
from typing import Dict, Any, Optional
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import pandas as pd
from datetime import datetime, timedelta
# from api.volume_ranking_utils import VolumeRankingDataManager  # 제거: 사이드 임팩트 방지
import re
from week_calculator import WeekCalculator, get_week_number, get_week_number_string, parse_week_string, get_week_start_date
from ranking_calculator import RankingCalculator

class StockNameMapper:
    """종목번호와 종목명 매핑 클래스 (DB 기반)"""
    
    def __init__(self, db_config: dict = None):
        """
        종목명 매퍼 초기화 (DB 기반)
        
        Args:
            db_config (dict): 데이터베이스 설정
        """
        self.db_config = db_config
        self.stock_mapping = {}
        self.use_db = False
        
        if db_config:
            try:
                self._load_stock_mapping_from_db()
                self.use_db = True
                print(f"✅ DB 기반 종목 매핑 초기화 완료: {len(self.stock_mapping)}개 종목")
            except Exception as e:
                print(f"⚠️ DB 기반 종목 매핑 초기화 실패: {e}")
                print("   fallback 매핑을 사용합니다.")
                self._create_fallback_mapping()
        else:
            print("⚠️ DB 설정이 없어 fallback 매핑을 사용합니다.")
            self._create_fallback_mapping()
    
    def _load_stock_mapping_from_db(self):
        """DB의 stocks 테이블에서 종목번호와 종목명 매핑 로드"""
        try:
            from database_config import DatabaseManager
            db_manager = DatabaseManager()
            
            if not db_manager.connect():
                raise Exception("데이터베이스 연결 실패")
            
            # stocks 테이블에서 종목코드와 종목명 조회
            query = "SELECT stock_code, stock_name FROM stocks WHERE stock_code IS NOT NULL AND stock_name IS NOT NULL"
            result = db_manager.fetch_all(query)
            
            if result:
                for row in result:
                    stock_code = str(row['stock_code']).zfill(6)  # 6자리로 패딩
                    stock_name = str(row['stock_name'])
                    self.stock_mapping[stock_code] = stock_name
                
                print(f"✅ DB에서 종목 매핑 로드 완료: {len(self.stock_mapping)}개 종목")
            else:
                print("⚠️ DB에서 종목 정보를 찾을 수 없습니다.")
                print("   stocks 테이블에 종목 데이터가 없거나 조회 조건에 맞는 데이터가 없습니다.")
                self._create_fallback_mapping()
            
            db_manager.disconnect()
                
        except ImportError as e:
            print(f"❌ database_config 모듈을 찾을 수 없습니다: {e}")
            self._create_fallback_mapping()
        except Exception as e:
            print(f"❌ DB에서 종목 매핑 로드 중 오류: {e}")
            print("   DB 연결 상태와 stocks 테이블 존재 여부를 확인해주세요.")
            self._create_fallback_mapping()
    
    def _create_fallback_mapping(self):
        """기본 종목 매핑 데이터 생성 (fallback용)"""
        self.stock_mapping = {
            "019210": "YG-1",
            "023410": "유진기업", 
            "145720": "덴티움",
            "005930": "삼성전자",
            "014280": "금강철강"
        }
        print(f"⚠️ fallback 종목 매핑 생성: {len(self.stock_mapping)}개 종목")
    
    def get_stock_name(self, stock_code: str) -> str:
        """
        종목번호로 종목명 조회 (DB 우선, fallback 매핑 차선)
        
        Args:
            stock_code (str): 종목번호
            
        Returns:
            str: 종목명 (찾지 못한 경우 종목번호 반환)
        """
        if not stock_code:
            return "알 수 없음"
        
        # 종목번호 정리 (앞의 0 제거 후 6자리로 패딩)
        clean_code = stock_code.lstrip('0')
        if not clean_code:
            clean_code = '0'
        
        padded_code = clean_code.zfill(6)
        
        # DB 매핑에서 찾기
        if padded_code in self.stock_mapping:
            return self.stock_mapping[padded_code]
        
        # 원본 코드로도 시도
        if stock_code in self.stock_mapping:
            return self.stock_mapping[stock_code]
        
        # DB에서 실시간 조회 시도 (DB 연결이 있는 경우)
        if self.use_db and self.db_config:
            try:
                from database_config import DatabaseManager
                db_manager = DatabaseManager()
                
                if not db_manager.connect():
                    raise Exception("데이터베이스 연결 실패")
                
                # 종목코드로 정확한 종목명 조회
                query = "SELECT stock_name FROM stocks WHERE stock_code = %s OR stock_code = %s"
                result = db_manager.fetch_all(query, (padded_code, stock_code))
                
                if result and len(result) > 0:
                    stock_name = str(result[0]['stock_name'])
                    # 캐시에 추가 (원본 코드와 패딩된 코드 모두)
                    self.stock_mapping[padded_code] = stock_name
                    self.stock_mapping[stock_code] = stock_name
                    print(f"✅ DB에서 실시간 조회 성공: {stock_code} -> {stock_name}")
                    db_manager.disconnect()
                    return stock_name
                
                db_manager.disconnect()
                    
            except Exception as e:
                print(f"⚠️ DB 실시간 조회 실패: {e}")
                # DB 연결 상태 재확인
                self.use_db = False
        
        # 찾지 못한 경우 종목번호 반환
        print(f"⚠️ 종목명 매핑 실패: {stock_code} (DB 조회 실패 또는 데이터 없음)")
        return stock_code
    
    def extract_stock_info_from_filename(self, filename: str) -> tuple:
        """
        파일명에서 종목정보 추출 (DB 기반 종목명 조회)
        
        Args:
            filename (str): 파일명
            
        Returns:
            tuple: (종목명, 종목번호)
        """
        try:
            # 파일명에서 확장자 제거
            name_without_ext = os.path.splitext(filename)[0]
            
            # 언더스코어로 분리
            parts = name_without_ext.split('_')
            
            stock_name = ""
            stock_code = ""
            
            # 파일명 패턴 분석
            if len(parts) >= 3:
                # weekly_Samsung_Electronics_Co.,_Ltd._005930_20250804 형태
                # 또는 daily_380550_380550_20250804 형태
                
                # 마지막에서 두 번째 부분이 종목번호일 가능성이 높음
                for i, part in enumerate(parts):
                    # 6자리 숫자인 경우 종목번호로 간주
                    if len(part) == 6 and part.isdigit():
                        stock_code = part
                        # 종목번호 앞의 부분들을 종목명으로 조합
                        stock_name_parts = parts[1:i] if i > 1 else parts[1:]
                        stock_name = "_".join(stock_name_parts)
                        break
                
                # 종목번호를 찾지 못한 경우, 파일명에서 직접 추출 시도
                if not stock_code:
                    for part in parts:
                        if len(part) == 6 and part.isdigit():
                            stock_code = part
                            break
            
            # 종목명이 비어있거나 종목번호가 없는 경우
            if not stock_name or not stock_code:
                # 파일명에서 6자리 숫자 찾기
                import re
                code_match = re.search(r'(\d{6})', filename)
                if code_match:
                    stock_code = code_match.group(1)
                    # 종목번호로 DB에서 종목명 조회
                    stock_name = self.get_stock_name(stock_code)
                else:
                    # 파일명 그대로 사용
                    stock_name = name_without_ext
                    stock_code = "000000"
            
            # 종목번호가 있는 경우 무조건 DB에서 한글 종목명 조회 (최우선)
            if stock_code and stock_code != "000000":
                mapped_name = self.get_stock_name(stock_code)
                if mapped_name != stock_code:  # DB에서 한글 종목명을 찾은 경우
                    stock_name = mapped_name
                    print(f"✅ DB에서 한글 종목명으로 교체: {stock_code} -> {stock_name}")
                else:
                    # DB에서도 찾지 못한 경우 파일명에서 추출한 종목명 사용
                    print(f"⚠️ DB에서 종목명을 찾지 못해 파일명 종목명 사용: {stock_code} -> {stock_name}")
            else:
                print(f"⚠️ 종목번호가 없어 파일명 종목명 사용: {stock_name}")
            
            return stock_name, stock_code
            
        except Exception as e:
            print(f"❌ 파일명에서 종목정보 추출 중 오류: {e}")
            return filename, "000000"

class ChartAnalysisPrompts:
    """차트 분석 프롬프트 관리 클래스 (DB 기반)"""
    
    def __init__(self, db_config: dict):
        """
        프롬프트 매니저 초기화
        
        Args:
            db_config (dict): 데이터베이스 설정
        """
        try:
            from prompt_manager import PromptManager
            self.prompt_manager = PromptManager(db_config)
            self.use_db = True
            print("✅ DB 기반 프롬프트 관리자 초기화 완료")
        except ImportError:
            print("⚠️ prompt_manager를 찾을 수 없어 하드코딩된 프롬프트를 사용합니다.")
            self.use_db = False
    
    def get_prompt(self, chart_type: str) -> str:
        """차트 유형에 따른 프롬프트 반환 (DB에서 조회)"""
        if self.use_db:
            prompt = self.prompt_manager.get_prompt(chart_type)
            if prompt:
                return prompt
            else:
                print(f"⚠️ DB에서 {chart_type} 프롬프트를 찾을 수 없습니다.")
                return self._get_fallback_prompt(chart_type)
        else:
            print("⚠️ DB 연결이 불가능하여 하드코딩된 프롬프트를 사용합니다.")
            return self._get_fallback_prompt(chart_type)
    
    def _get_fallback_prompt(self, chart_type: str) -> str:
        """하드코딩된 기본 프롬프트 반환 (fallback용)"""
        chart_type = chart_type.lower()
        if chart_type in ['daily', '일봉', 'day']:
            return self.get_daily_prompt()
        elif chart_type in ['weekly', '주봉', 'week']:
            return self.get_weekly_prompt()
        elif chart_type in ['monthly', '월봉', 'month']:
            return self.get_monthly_prompt()
        else:
            # 기본값은 일봉
            return self.get_daily_prompt()

    def get_summary_prompt(self, chart_type: str) -> str:
        """차트 유형에 따른 요약 프롬프트 반환 (DB에서 조회)"""
        if self.use_db:
            summary_prompt_type = f"{chart_type} 요약"
            prompt = self.prompt_manager.get_prompt(summary_prompt_type)
            if prompt:
                return prompt
            else:
                print(f"⚠️ DB에서 {summary_prompt_type} 프롬프트를 찾을 수 없습니다.")
                return self._get_fallback_summary_prompt(chart_type)
        else:
            print("⚠️ DB 연결이 불가능하여 하드코딩된 요약 프롬프트를 사용합니다.")
            return self._get_fallback_summary_prompt(chart_type)

    def _get_fallback_summary_prompt(self, chart_type: str) -> str:
        """하드코딩된 기본 요약 프롬프트 반환 (fallback용)"""
        chart_type = chart_type.lower()
        if chart_type in ['daily', '일봉', 'day']:
            return self.get_daily_summary_prompt()
        elif chart_type in ['weekly', '주봉', 'week']:
            return self.get_weekly_summary_prompt()
        elif chart_type in ['monthly', '월봉', 'month']:
            return self.get_monthly_summary_prompt()
        else:
            # 기본값은 일봉 요약
            return self.get_daily_summary_prompt()

    @staticmethod
    def get_daily_prompt() -> str:
        """일봉 차트 분석 프롬프트 (fallback용) - DB에서 로드 실패 시 사용"""
        return "일봉 차트 분석 프롬프트 (DB에서 로드 실패 시 기본값)"

    @staticmethod
    def get_weekly_prompt() -> str:
        """주봉 차트 분석 프롬프트 (fallback용) - DB에서 로드 실패 시 사용"""
        return "주봉 차트 분석 프롬프트 (DB에서 로드 실패 시 기본값)"

    @staticmethod
    def get_monthly_prompt() -> str:
        """월봉 차트 분석 프롬프트 (fallback용) - DB에서 로드 실패 ㄷㅏ정시 사용"""
        return "월봉 차트 분석 프롬프트 (DB에서 로드 실패 시 기본값)"

    @staticmethod
    def get_daily_summary_prompt() -> str:
        """일봉 요약 프롬프트 (fallback용) - DB에서 로드 실패 시 사용"""
        return "일봉 차트 분석 결과를 요약하여 핵심 포인트를 정리해주세요."

    @staticmethod
    def get_weekly_summary_prompt() -> str:
        """주봉 요약 프롬프트 (fallback용) - DB에서 로드 실패 시 사용"""
        return "주봉 차트 분석 결과를 요약하여 핵심 포인트를 정리해주세요."

    @staticmethod
    def get_monthly_summary_prompt() -> str:
        """월봉 요약 프롬프트 (fallback용) - DB에서 로드 실패 시 사용"""
        return "월봉 차트 분석 결과를 요약하여 핵심 포인트를 정리해주세요."

class AIChartAnalyzer:
    def __init__(self, api_key: str = None, db_config: dict = None):
        """
        AI 차트 분석기 초기화 (DB 기반 설정 지원)
        
        Args:
            api_key (str): Google AI API 키 (None이면 DB에서 로드 시도)
            db_config (dict): 데이터베이스 설정 (None이면 config.py에서 로드)
        """
        # 데이터베이스 설정 로드
        if db_config is None:
            try:
                from config import config
                db_config = config.get_database_config()
            except ImportError:
                print("⚠️ config.py를 찾을 수 없어 DB 기능을 사용할 수 없습니다.")
                db_config = None
        
        # API 키 설정
        if api_key is None and db_config:
            try:
                from prompt_manager import SecureConfigManager
                secure_manager = SecureConfigManager(db_config)
                api_key = secure_manager.get_api_key()
                if api_key:
                    print("✅ DB에서 API 키를 로드했습니다.")
                else:
                    print("⚠️ DB에서 API 키를 찾을 수 없습니다.")
                    api_key = self._get_api_key_from_user()
            except ImportError:
                print("⚠️ prompt_manager를 찾을 수 없어 DB에서 API 키를 로드할 수 없습니다.")
                api_key = self._get_api_key_from_user()
        elif api_key is None:
            api_key = self._get_api_key_from_user()
        
        if not api_key:
            raise ValueError("유효한 API 키가 필요합니다.")
        
        # API 키 유효성 검증
        if not self._validate_api_key(api_key):
            raise ValueError("유효하지 않은 API 키입니다.")
        
        self.api_key = api_key
        genai.configure(api_key=api_key)
        
        # 제미나이 모델 설정
        self.model = genai.GenerativeModel('gemini-1.5-flash')
        
        # 종목명 매퍼 초기화 (DB 기반)
        self.stock_mapper = StockNameMapper(db_config)
        
        # 순위 계산기 초기화
        self.ranking_calculator = RankingCalculator(db_config)
        
        # 프롬프트 매니저 초기화 (DB 기반)
        if db_config:
            try:
                self.prompts = ChartAnalysisPrompts(db_config)
                print("✅ DB 기반 프롬프트 관리자 초기화 완료")
            except Exception as e:
                print(f"⚠️ DB 기반 프롬프트 관리자 초기화 실패: {e}")
                self.prompts = ChartAnalysisPrompts(None)  # fallback
        else:
            self.prompts = ChartAnalysisPrompts(None)  # fallback
        
        # 주봉 거래일 캐시 (추출 실패 시 성공한 종목의 거래일 사용)
        self._weekly_trading_date_cache = None
        
        # 시스템 설정 로드 (DB에서)
        if db_config:
            try:
                from prompt_manager import UnifiedConfigManager
                self.config_manager = UnifiedConfigManager(db_config)
                self.max_retries = self.config_manager.get_system_config('max_retry_count', 3)
                self.timeout = self.config_manager.get_system_config('analysis_timeout', 300)
                self.max_image_size = self.config_manager.get_system_config('chart_image_max_size', 1500)
                print("✅ DB에서 시스템 설정을 로드했습니다.")
            except Exception as e:
                print(f"⚠️ DB에서 시스템 설정 로드 실패: {e}")
                self._set_default_config()
        else:
            self._set_default_config()
    
    def _get_api_key_from_user(self) -> str:
        """사용자로부터 API 키 입력 받기"""
        print("🔑 Google AI API 키를 입력해주세요:")
        print("(https://makersuite.google.com/app/apikey 에서 발급 가능)")
        
        api_key = input("Google AI API 키: ").strip()
        if not api_key:
            raise ValueError("API 키가 필요합니다.")
        
        # DB에 저장 시도
        try:
            from config import config
            db_config = config.get_database_config()
            from prompt_manager import SecureConfigManager
            secure_manager = SecureConfigManager(db_config)
            if secure_manager.save_api_key(api_key):
                print("✅ API 키가 DB에 암호화되어 저장되었습니다.")
        except Exception as e:
            print(f"⚠️ API 키 DB 저장 실패: {e}")
        
        return api_key
    
    def _validate_api_key(self, api_key: str) -> bool:
        """API 키 유효성 검증"""
        try:
            # API 키를 먼저 설정
            genai.configure(api_key=api_key)
            test_model = genai.GenerativeModel('gemini-1.5-flash')
            response = test_model.generate_content("테스트")
            print("✅ API 키 유효성 검증 완료")
            return True
        except Exception as e:
            print(f"❌ API 키 유효성 검증 실패: {e}")
            return False
    
    def _set_default_config(self):
        """기본 설정값 설정"""
        self.max_retries = 3
        self.timeout = 300
        self.max_image_size = 1500
        print("⚠️ 기본 설정값을 사용합니다.")

    def encode_image_to_base64(self, image_path: str) -> str:
        """
        이미지를 base64로 인코딩
        
        Args:
            image_path (str): 이미지 파일 경로
            
        Returns:
            str: base64 인코딩된 이미지
        """
        try:
            with open(image_path, "rb") as image_file:
                return base64.b64encode(image_file.read()).decode('utf-8')
        except Exception as e:
            print(f"❌ 이미지 인코딩 오류: {e}")
            return ""

    def analyze_chart_image(self, image_path: str, stock_name: str = "", chart_type: str = "일봉", chart_data: Optional[pd.DataFrame] = None, 
                           json_data_path: str = "", csv_data_path: str = "", text_summary_path: str = "", 
                           enable_summary_analysis: bool = False, additional_info: Optional[Dict[str, Any]] = None, trading_type: str = "거래대금") -> Optional[Dict[str, Any]]:
        """
        차트 이미지를 AI로 분석 (하이브리드 버전 - 개별 + 요약 분석 동시 지원)
        
        Args:
            image_path (str): 차트 이미지 파일 경로
            stock_name (str): 종목명
            chart_type (str): 차트 유형 (일봉/주봉/월봉)
            chart_data (pd.DataFrame): 차트 데이터 (Open, High, Low, Close, Volume)
            json_data_path (str): JSON 데이터 파일 경로
            csv_data_path (str): CSV 데이터 파일 경로
            text_summary_path (str): 텍스트 요약 파일 경로
            enable_summary_analysis (bool): 요약 분석 활성화 여부
            additional_info (Dict[str, Any]): 추가 정보
            trading_type (str): 거래 타입 (거래대금, 거래율)
            
        Returns:
            Dict[str, Any]: 분석 결과 JSON (요약 분석 포함 시 summary_analysis 키 추가)
        """
        try:
            print(f"🔍 AI 차트 분석 시작: {image_path}")
            print(f"📊 차트 유형: {chart_type}")
            
            # 분석 시작 시간 기록
            start_time = time.time()
            
            # additional_info에서 거래타입 우선 사용
            if additional_info and "trading_type" in additional_info:
                trading_type = additional_info["trading_type"]
                print(f"✅ additional_info에서 거래타입 읽기: {trading_type}")
            else:
                print(f"⚠️ additional_info에 거래타입 없음, 기본값 사용: {trading_type}")
            
            # JSON 파일에서 거래 타입 읽기 (additional_info가 없을 때만)
            if json_data_path and os.path.exists(json_data_path) and not (additional_info and "trading_type" in additional_info):
                try:
                    with open(json_data_path, 'r', encoding='utf-8') as f:
                        json_data = json.load(f)
                        if '거래 타입' in json_data:
                            trading_type = json_data['거래 타입']
                            print(f"✅ JSON에서 거래 타입 읽기: {trading_type}")
                        else:
                            print(f"⚠️ JSON에 '거래 타입' 키가 없음, 기본값 사용: {trading_type}")
                except Exception as e:
                    print(f"⚠️ JSON 파일 읽기 실패: {e}, 기본값 사용: {trading_type}")
            
            # additional_info에서 종목명 우선 사용
            if additional_info and "stock_name" in additional_info:
                stock_name = additional_info["stock_name"]
                print(f"✅ additional_info에서 종목명 사용: {stock_name}")
            
            # 파일명에서 종목정보 추출
            filename = os.path.basename(image_path)
            extracted_stock_name, extracted_stock_code = self.stock_mapper.extract_stock_info_from_filename(filename)
            
            # 종목번호가 있는 경우 무조건 DB에서 한글 종목명 우선 조회
            if extracted_stock_code and extracted_stock_code != "000000":
                mapped_name = self.stock_mapper.get_stock_name(extracted_stock_code)
                if mapped_name != extracted_stock_code:  # DB에서 한글 종목명을 찾은 경우
                    if not stock_name:  # additional_info에 종목명이 없는 경우에만 DB에서 조회한 이름 사용
                        stock_name = mapped_name
                    print(f"✅ DB에서 한글 종목명 조회 성공: {extracted_stock_code} -> {mapped_name}")
                else:
                    # DB에서 찾지 못한 경우 처리
                    if not stock_name:  # 기존 종목명이 없는 경우에만 파일명 사용
                        stock_name = extracted_stock_name
                    print(f"⚠️ DB에서 종목명을 찾지 못함: {extracted_stock_code}, 사용할 종목명: {stock_name}")
            elif not stock_name:  # 종목번호도 없고 종목명도 없는 경우
                stock_name = extracted_stock_name
                print(f"⚠️ 종목번호가 없어 파일명 종목명 사용: {stock_name}")
            
            print(f"📈 종목명: {stock_name}")
            print(f"📈 종목번호: {extracted_stock_code}")
            
            # 1. 이미지 파일 존재 및 무결성 검증
            if not os.path.exists(image_path):
                print(f"❌ 이미지 파일을 찾을 수 없습니다: {image_path}")
                return None
            
            # 이미지 파일 크기 확인
            file_size = os.path.getsize(image_path)
            print(f"📊 이미지 파일 크기: {file_size:,} bytes")
            
            if file_size == 0:
                print(f"❌ 이미지 파일이 비어있습니다: {image_path}")
                return None
            
            # 2. 이미지 형식 및 크기 검증
            try:
                with Image.open(image_path) as img:
                    print(f"📊 이미지 크기: {img.size}")
                    print(f"📊 이미지 형식: {img.format}")
                    print(f"📊 이미지 모드: {img.mode}")
                    
                    # 이미지 크기가 너무 작으면 경고
                    if img.size[0] < 100 or img.size[1] < 100:
                        print(f"⚠️ 이미지 크기가 너무 작습니다: {img.size}")
                    
                    # 이미지가 너무 크면 리사이즈 고려
                    if img.size[0] > 4000 or img.size[1] > 4000:
                        print(f"⚠️ 이미지 크기가 너무 큽니다: {img.size}")
                        print(f"🔄 이미지를 2000x2000으로 리사이즈합니다...")
                        
                        # 이미지 리사이즈
                        max_size = 2000
                        if img.size[0] > img.size[1]:
                            new_width = max_size
                            new_height = int(img.size[1] * max_size / img.size[0])
                        else:
                            new_height = max_size
                            new_width = int(img.size[0] * max_size / img.size[1])
                        
                        img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
                        print(f"✅ 이미지 리사이즈 완료: {img.size}")
                        
            except Exception as e:
                print(f"❌ 이미지 파일 검증 실패: {e}")
                return None
            
            # 3. 차트 유형에 따른 프롬프트 선택
            prompt = self.prompts.get_prompt(chart_type)
            
            # 4. 추가 데이터 파일들 로드 및 프롬프트에 추가
            additional_data_info = self._load_additional_data_files(json_data_path, csv_data_path, text_summary_path)
            if additional_data_info:
                prompt += f"\n\n{additional_data_info}"
                print(f"✅ 추가 데이터 파일 정보가 프롬프트에 추가되었습니다.")
            
            # 5. 차트 데이터 정보를 프롬프트에 추가 (기존 방식)
            if chart_data is not None and hasattr(chart_data, 'empty') and not chart_data.empty:
                print(f"📊 차트 데이터 정보 추가: {len(chart_data)}개 데이터 포인트")
                
                # 최근 데이터 요약 정보 생성
                recent_data = chart_data.tail(10)  # 최근 10개 데이터
                
                data_summary = f"""
**차트 데이터 정보:**
- 데이터 기간: {chart_data.index[0].strftime('%Y-%m-%d')} ~ {chart_data.index[-1].strftime('%Y-%m-%d')}
- 총 데이터 수: {len(chart_data)}개
- 최근 10개 데이터:
"""
                
                for i, (date, row) in enumerate(recent_data.iterrows()):
                    # 안전한 포맷팅 적용
                    open_val = f"{row['Open']:,.0f}" if pd.notna(row['Open']) else "N/A"
                    high_val = f"{row['High']:,.0f}" if pd.notna(row['High']) else "N/A"
                    low_val = f"{row['Low']:,.0f}" if pd.notna(row['Low']) else "N/A"
                    close_val = f"{row['Close']:,.0f}" if pd.notna(row['Close']) else "N/A"
                    volume_val = f"{row['Volume']:,.0f}" if pd.notna(row['Volume']) else "N/A"
                    
                    data_summary += f"- {date.strftime('%Y-%m-%d')}: 시가 {open_val}, 고가 {high_val}, 저가 {low_val}, 종가 {close_val}, 거래대금 {volume_val}\n"
                
                # 기술적 지표 정보 추가 (있는 경우)
                if 'MA5' in chart_data.columns and pd.notna(chart_data['MA5'].iloc[-1]):
                    data_summary += f"- 5기간 이동평균: {chart_data['MA5'].iloc[-1]:,.0f}\n"
                if 'MA20' in chart_data.columns and pd.notna(chart_data['MA20'].iloc[-1]):
                    data_summary += f"- 20기간 이동평균: {chart_data['MA20'].iloc[-1]:,.0f}\n"
                if 'RSI' in chart_data.columns and pd.notna(chart_data['RSI'].iloc[-1]):
                    data_summary += f"- RSI: {chart_data['RSI'].iloc[-1]:.1f}\n"
                if 'MACD' in chart_data.columns and pd.notna(chart_data['MACD'].iloc[-1]):
                    data_summary += f"- MACD: {chart_data['MACD'].iloc[-1]:.2f}\n"
                
                # 가격 변동 정보 (안전한 계산)
                try:
                    if pd.notna(chart_data['Close'].iloc[-1]) and pd.notna(chart_data['Open'].iloc[0]):
                        price_change = chart_data['Close'].iloc[-1] - chart_data['Open'].iloc[0]
                        price_change_pct = (price_change / chart_data['Open'].iloc[0]) * 100
                        data_summary += f"- 전체 기간 가격 변동: {price_change:+,.0f}원 ({price_change_pct:+.2f}%)\n"
                except (ValueError, TypeError, ZeroDivisionError):
                    data_summary += f"- 전체 기간 가격 변동: N/A\n"
                
                # 최근 변동 정보 (안전한 계산)
                try:
                    if len(chart_data) > 1 and pd.notna(chart_data['Close'].iloc[-1]) and pd.notna(chart_data['Close'].iloc[-2]):
                        recent_change = chart_data['Close'].iloc[-1] - chart_data['Close'].iloc[-2]
                        recent_change_pct = (recent_change / chart_data['Close'].iloc[-2]) * 100
                        data_summary += f"- 최근 변동: {recent_change:+,.0f}원 ({recent_change_pct:+.2f}%)\n"
                    else:
                        data_summary += f"- 최근 변동: N/A\n"
                except (ValueError, TypeError, ZeroDivisionError):
                    data_summary += f"- 최근 변동: N/A\n"
                
                prompt += data_summary
                
                print(f"✅ 차트 데이터 정보가 프롬프트에 추가되었습니다.")
            elif chart_data is not None and isinstance(chart_data, dict):
                print(f"📊 딕셔너리 형태의 차트 데이터 정보 추가")
                
                # 딕셔너리 형태의 데이터에서 정보 추출
                data_summary = f"""
**차트 데이터 정보 (딕셔너리):**
- 종목명: {chart_data.get('metadata', {}).get('stock_name', 'N/A')}
- 종목코드: {chart_data.get('metadata', {}).get('stock_code', 'N/A')}
- 차트 유형: {chart_data.get('metadata', {}).get('chart_type', 'N/A')}
- 데이터 기간: {chart_data.get('metadata', {}).get('data_period', {}).get('start', 'N/A')} ~ {chart_data.get('metadata', {}).get('data_period', {}).get('end', 'N/A')}
- 총 데이터 수: {chart_data.get('metadata', {}).get('total_records', 'N/A')}일
"""
                
                # 가격 요약 정보 추가
                if 'price_summary' in chart_data:
                    price_info = chart_data['price_summary']
                    
                    # 안전한 포맷팅 함수
                    def safe_format_number(value, format_str=":,.0f", default="N/A"):
                        try:
                            if value is None or value == 'N/A':
                                return default
                            return f"{value:{format_str}}"
                        except (ValueError, TypeError):
                            return default
                    
                    def safe_format_percent(value, format_str=":+.2f", default="N/A"):
                        try:
                            if value is None or value == 'N/A':
                                return default
                            return f"{value:{format_str}}%"
                        except (ValueError, TypeError):
                            return default
                    
                    data_summary += f"""
- 최근 종가: {safe_format_number(price_info.get('latest_close'))}원
- 가격 변동: {safe_format_number(price_info.get('price_change'), ':+,.0f')}원
- 가격 변동률: {safe_format_percent(price_info.get('price_change_pct'))}
- 최고가: {safe_format_number(price_info.get('highest_price'))}원
- 최저가: {safe_format_number(price_info.get('lowest_price'))}원
- 평균 거래량: {safe_format_number(price_info.get('avg_volume'))}주
"""
                
                # 보조지표 정보 추가
                if 'technical_indicators' in chart_data:
                    ti = chart_data['technical_indicators']
                    if 'latest_values' in ti:
                        latest_values = ti['latest_values']
                        data_summary += f"""
- 5일 이동평균: {safe_format_number(latest_values.get('ma5'))}원
- 20일 이동평균: {safe_format_number(latest_values.get('ma20'))}원
- RSI: {safe_format_number(latest_values.get('rsi'), ':.1f')}
- MACD: {safe_format_number(latest_values.get('macd'), ':.2f')}
"""
                
                prompt += data_summary
                print(f"✅ 딕셔너리 형태의 차트 데이터 정보가 프롬프트에 추가되었습니다.")
                
            else:
                print(f"⚠️ 차트 데이터가 제공되지 않았습니다. 이미지와 추가 데이터 파일만으로 분석을 진행합니다.")
            
            # 종목명 정보를 프롬프트에 추가
            prompt += f"\n\n**중요: 분석할 종목은 '{stock_name}' (종목번호: {extracted_stock_code})입니다.**"
            prompt = prompt.replace("[종목명]", stock_name)
            
            # 6. AI 분석 재시도 메커니즘
            max_retries = self.max_retries
            for attempt in range(max_retries):
                try:
                    print(f"🔄 AI 분석 시도 {attempt + 1}/{max_retries}")
                    
                    # 이미지 로드 및 리사이즈
                    image = Image.open(image_path)
                    
                    # 이미지가 너무 크면 리사이즈
                    if image.size[0] > self.max_image_size or image.size[1] > self.max_image_size:
                        max_size = self.max_image_size
                        if image.size[0] > image.size[1]:
                            new_width = max_size
                            new_height = int(image.size[1] * max_size / image.size[0])
                        else:
                            new_height = max_size
                            new_width = int(image.size[0] * max_size / image.size[1])
                        
                        image = image.resize((new_width, new_height), Image.Resampling.LANCZOS)
                        print(f"🔄 AI 분석용 이미지 리사이즈: {image.size}")
                    
                    # AI 분석 요청 (이미지를 base64로 인코딩하여 전송)
                    try:
                        # 이미지를 base64로 인코딩
                        import base64
                        import io
                        
                        # 이미지를 메모리에 저장
                        img_buffer = io.BytesIO()
                        image.save(img_buffer, format='PNG')
                        img_buffer.seek(0)
                        
                        # base64로 인코딩
                        img_base64 = base64.b64encode(img_buffer.getvalue()).decode('utf-8')
                        print(f"📊 이미지 base64 인코딩 완료: {len(img_base64)} 문자")
                        
                        # AI 분석 요청 (base64 이미지 포함)
                        response = self.model.generate_content([
                            prompt,
                            {
                                "mime_type": "image/png",
                                "data": img_base64
                            }
                        ])
                        
                        # 응답이 이미지를 인식하지 못하는 경우, 다른 방식 시도
                        if "이미지가 제공되지 않았으므로" in response.text or "이미지가 없으므로" in response.text:
                            print("⚠️ AI가 이미지를 인식하지 못함. 다른 방식으로 시도...")
                            
                            # 파일 경로를 직접 전달하는 방식 시도
                            response = self.model.generate_content([
                                prompt + "\n\n이미지 파일 경로: " + image_path,
                                image
                            ])
                        
                    except Exception as e:
                        print(f"⚠️ base64 인코딩 실패, 기본 방식으로 시도: {e}")
                        # 기본 방식으로 시도
                        response = self.model.generate_content([
                            prompt,
                            image
                        ])
                    
                    if response.text:
                        print(f"✅ AI 분석 완료 (시도 {attempt + 1})")
                        print(f"📝 AI 응답 길이: {len(response.text)}")
                        print(f"📝 AI 응답 시작: {response.text[:100]}...")
                        
                        # 7. 응답 검증 및 JSON 파싱
                        if self._is_valid_json_response(response.text):
                            try:
                                analysis_result = self._parse_json_response(response.text)
                                
                                # 분석 일시 및 차트 유형 추가
                                if "종목정보" in analysis_result:
                                    analysis_result["종목정보"]["분석일시"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                                    analysis_result["종목정보"]["차트유형"] = chart_type
                                    # 거래타입 추가
                                    if additional_info and "trading_type" in additional_info:
                                        analysis_result["종목정보"]["거래타입"] = additional_info["trading_type"]
                                    # 종목정보가 없는 경우 추가
                                    if "종목명" not in analysis_result["종목정보"]:
                                        analysis_result["종목정보"]["종목명"] = stock_name
                                    if "종목번호" not in analysis_result["종목정보"]:
                                        analysis_result["종목정보"]["종목번호"] = extracted_stock_code
                                    
                                else:
                                    # 종목정보 섹션이 없는 경우 생성
                                    analysis_result["종목정보"] = {
                                        "종목명": stock_name,
                                        "종목번호": extracted_stock_code,
                                        "분석일시": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                                        "차트유형": chart_type,
                                        "거래타입": additional_info.get("trading_type", "") if additional_info else ""
                                    }
                                
                                # 거래정보 추가 (JSON 파싱 성공)
                                # AI 응답이 JSON 형식으로 정상 파싱된 경우에만 실행됨
                                # 현재는 AI 응답이 JSON 형식이 아니어서 이 경로는 거의 실행되지 않음
                                # 향후 AI 응답 개선 시 자동으로 작동하도록 유지
                                analysis_result = self._add_trading_info_to_result(analysis_result, extracted_stock_code, "json파싱성공", chart_type, response.text, trading_type)
                                
                                # 하이브리드 방식: 원본 AI 응답 저장
                                analysis_result["original_ai_response"] = response.text
                                
                                print(f"✅ JSON 파싱 성공 (시도 {attempt + 1})")
                                
                                # 투 트랙 분석: 요약 분석 실행 (요청된 경우)
                                if enable_summary_analysis:
                                    print(f"🔄 요약 분석 시작...")
                                    summary_result = self._run_summary_analysis(
                                        image_path, stock_name, chart_type, chart_data, 
                                        json_data_path, csv_data_path, text_summary_path,
                                        analysis_result
                                    )
                                    if summary_result:
                                        analysis_result["summary_analysis"] = summary_result
                                        print(f"✅ 요약 분석 완료")
                                    else:
                                        print(f"⚠️ 요약 분석 실패")
                                
                                return analysis_result
                                
                            except json.JSONDecodeError as e:
                                print(f"⚠️ JSON 파싱 오류 (시도 {attempt + 1}): {e}")
                                if attempt < max_retries - 1:
                                    print(f"🔄 재시도 중... ({attempt + 2}/{max_retries})")
                                    time.sleep(2)
                                    continue
                                else:
                                    print(f"❌ 모든 시도 실패. 마지막 응답: {response.text}")
                                    fallback_result = self._create_fallback_result(stock_name, chart_type, response.text, "JSON 파싱 실패", extracted_stock_code, chart_data, additional_info, trading_type)
                                    fallback_result["original_ai_response"] = response.text
                                    return fallback_result
                        else:
                            print(f"⚠️ AI 응답이 JSON 형식이 아닙니다 (시도 {attempt + 1})")
                            if attempt < max_retries - 1:
                                print(f"🔄 재시도 중... ({attempt + 2}/{max_retries})")
                                time.sleep(2)
                                continue
                            else:
                                print(f"❌ 모든 시도 실패. 마지막 응답: {response.text}")
                                fallback_result = self._create_fallback_result(stock_name, chart_type, response.text, "JSON 형식 아님", extracted_stock_code, chart_data, additional_info, trading_type)
                                fallback_result["original_ai_response"] = response.text
                                return fallback_result
                    else:
                        print(f"❌ AI 분석 응답이 없습니다 (시도 {attempt + 1})")
                        if attempt < max_retries - 1:
                            print(f"🔄 재시도 중... ({attempt + 2}/{max_retries})")
                            time.sleep(2)
                            continue
                        else:
                            return None
                            
                except Exception as e:
                    print(f"❌ AI 분석 중 오류 발생 (시도 {attempt + 1}): {e}")
                    if attempt < max_retries - 1:
                        print(f"🔄 재시도 중... ({attempt + 2}/{max_retries})")
                        time.sleep(2)
                        continue
                    else:
                        return None
            
            return None
                
        except Exception as e:
            print(f"❌ AI 분석 중 예상치 못한 오류 발생: {e}")
            return None
    
    def _run_summary_analysis(self, image_path: str, stock_name: str, chart_type: str, 
                             chart_data: Optional[pd.DataFrame], json_data_path: str, 
                             csv_data_path: str, text_summary_path: str, 
                             individual_result: Dict[str, Any]) -> Optional[Dict[str, Any]]:
        """
        요약 분석 실행 (개별 분석과 동일한 차트 이미지/데이터 사용)
        
        Args:
            image_path (str): 차트 이미지 파일 경로
            stock_name (str): 종목명
            chart_type (str): 차트 유형
            chart_data (pd.DataFrame): 차트 데이터
            json_data_path (str): JSON 데이터 파일 경로
            csv_data_path (str): CSV 데이터 파일 경로
            text_summary_path (str): 텍스트 요약 파일 경로
            individual_result (Dict): 개별 분석 결과
            
        Returns:
            Dict[str, Any]: 요약 분석 결과
        """
        try:
            # 요약 프롬프트 가져오기
            summary_prompt = self.prompts.get_summary_prompt(chart_type)
            if not summary_prompt:
                print(f"⚠️ {chart_type} 요약 프롬프트를 찾을 수 없습니다.")
                return None
            
            print(f"📝 요약 프롬프트 준비 완료")
            
            # 개별 분석 결과를 요약 프롬프트에 추가
            summary_prompt += f"\n\n**개별 분석 결과:**\n{json.dumps(individual_result, ensure_ascii=False, indent=2)}\n\n"
            summary_prompt += "위 개별 분석 결과를 바탕으로 핵심 포인트를 요약하고 투자 관점에서의 시사점을 제시해주세요."
            
            # 개별 분석 JSON 파일 경로 추가
            if "individual_analysis_file" in individual_result:
                summary_prompt += f"\n\n**개별 분석 JSON 파일:** {individual_result['individual_analysis_file']}\n"
            
            # 추가 데이터 파일들 로드 및 프롬프트에 추가 (기존 로직 재사용)
            additional_data_info = self._load_additional_data_files(json_data_path, csv_data_path, text_summary_path)
            if additional_data_info:
                summary_prompt += f"\n\n{additional_data_info}"
            
            # 차트 데이터 정보를 프롬프트에 추가 (기존 로직 재사용)
            if chart_data is not None and hasattr(chart_data, 'empty') and not chart_data.empty:
                # 최근 데이터 요약 정보 생성 (간략화)
                recent_data = chart_data.tail(5)  # 최근 5개 데이터만
                
                data_summary = f"""
**차트 데이터 요약:**
- 데이터 기간: {chart_data.index[0].strftime('%Y-%m-%d')} ~ {chart_data.index[-1].strftime('%Y-%m-%d')}
- 총 데이터 수: {len(chart_data)}개
"""
                
                # 기술적 지표 정보 추가 (있는 경우)
                if 'MA5' in chart_data.columns and pd.notna(chart_data['MA5'].iloc[-1]):
                    data_summary += f"- 5기간 이동평균: {chart_data['MA5'].iloc[-1]:,.0f}\n"
                if 'MA20' in chart_data.columns and pd.notna(chart_data['MA20'].iloc[-1]):
                    data_summary += f"- 20기간 이동평균: {chart_data['MA20'].iloc[-1]:,.0f}\n"
                if 'RSI' in chart_data.columns and pd.notna(chart_data['RSI'].iloc[-1]):
                    data_summary += f"- RSI: {chart_data['RSI'].iloc[-1]:.1f}\n"
                
                summary_prompt += data_summary
            
            # 종목명 정보를 프롬프트에 추가
            summary_prompt += f"\n\n**중요: 요약 분석할 종목은 '{stock_name}'입니다.**"
            summary_prompt = summary_prompt.replace("[종목명]", stock_name)
            
            # 이미지 로드 및 리사이즈 (기존 로직 재사용)
            image = Image.open(image_path)
            if image.size[0] > self.max_image_size or image.size[1] > self.max_image_size:
                max_size = self.max_image_size
                if image.size[0] > image.size[1]:
                    new_width = max_size
                    new_height = int(image.size[1] * max_size / image.size[0])
                else:
                    new_height = max_size
                    new_width = int(image.size[0] * max_size / image.size[1])
                
                image = image.resize((new_width, new_height), Image.Resampling.LANCZOS)
            
            # AI 요약 분석 요청
            response = self.model.generate_content([summary_prompt, image])
            
            if response and response.text:
                print(f"✅ 요약 분석 응답 수신 완료")
                
                # 요약 분석 결과 구조화
                summary_result = {
                    "summary_type": f"{chart_type}_요약",
                    "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "stock_info": {
                        "stock_name": stock_name,
                        "chart_type": chart_type
                    },
                    "summary_content": response.text,
                    "based_on_individual_analysis": True
                }
                
                # JSON 형식 응답인지 확인하고 파싱 시도
                if self._is_valid_json_response(response.text):
                    try:
                        parsed_summary = self._parse_json_response(response.text)
                        summary_result["structured_summary"] = parsed_summary
                        print(f"✅ 요약 분석 JSON 파싱 성공")
                    except json.JSONDecodeError:
                        print(f"⚠️ 요약 분석 JSON 파싱 실패, 텍스트로 저장")
                
                return summary_result
            else:
                print(f"❌ 요약 분석 응답이 없습니다")
                return None
                
        except Exception as e:
            print(f"❌ 요약 분석 중 오류: {e}")
            return None
    
    def _is_valid_json_response(self, response_text: str) -> bool:
        """AI 응답이 유효한 JSON 형식인지 확인"""
        text = response_text.strip()
        
        # JSON 코드 블록 제거 (더 강력한 방식)
        if text.startswith('```json'):
            text = text[7:]
        elif text.startswith('```'):
            text = text[3:]
        if text.endswith('```'):
            text = text[:-3]
        
        # 앞뒤 공백 제거
        text = text.strip()
        
        # JSON 형식인지 확인
        is_json = text.startswith('{') and text.endswith('}')
        
        # 디버깅을 위한 로그 추가
        print(f"🔍 JSON 검증: {is_json}")
        print(f"🔍 응답 시작: {text[:50]}...")
        print(f"🔍 응답 끝: ...{text[-50:]}")
        
        return is_json
    
    def _parse_json_response(self, response_text: str) -> Dict[str, Any]:
        """AI 응답에서 JSON 파싱"""
        json_text = response_text.strip()
        
        # JSON 코드 블록 제거 (더 강력한 방식)
        if json_text.startswith('```json'):
            json_text = json_text[7:]
        elif json_text.startswith('```'):
            json_text = json_text[3:]
        if json_text.endswith('```'):
            json_text = json_text[:-3]
        
        # 앞뒤 공백 제거
        json_text = json_text.strip()
        
        return json.loads(json_text)
    
    def _convert_structured_data_to_text(self, structured_data: Dict[str, Any]) -> str:
        """
        구조화된 JSON 데이터를 텍스트로 변환 (하이브리드 방식용)
        
        Args:
            structured_data (Dict[str, Any]): 구조화된 분석 결과
            
        Returns:
            str: 텍스트 형태의 분석 결과
        """
        try:
            text_parts = []
            
            # 종목 정보
            if "종목정보" in structured_data:
                info = structured_data["종목정보"]
                text_parts.append(f"📊 종목 정보")
                text_parts.append(f"• 종목명: {info.get('종목명', 'N/A')}")
                text_parts.append(f"• 종목번호: {info.get('종목번호', 'N/A')}")
                text_parts.append(f"• 분석일시: {info.get('분석일시', 'N/A')}")
                text_parts.append(f"• 차트유형: {info.get('차트유형', 'N/A')}")
                text_parts.append("")
            
            # 종합 분석 점수
            if "종합분석점수" in structured_data:
                score = structured_data["종합분석점수"]
                text_parts.append(f"📈 종합 분석 점수")
                text_parts.append(f"• 점수: {score.get('점수', 'N/A')}/100")
                text_parts.append(f"• 요약: {score.get('요약', 'N/A')}")
                text_parts.append("")
            
            # 차트 유형별 봉 정보
            if "오늘의일봉" in structured_data:
                candle = structured_data["오늘의일봉"]
                text_parts.append(f"📊 오늘의 일봉 요약")
                text_parts.append(f"• 종가: {candle.get('종가', 'N/A')}원")
                text_parts.append(f"• 등락률: {candle.get('등락률', 'N/A')}%")
                text_parts.append(f"• 거래량: {candle.get('거래량', 'N/A')}주")
                text_parts.append(f"• 주요 특징: {candle.get('주요특징', 'N/A')}")
                text_parts.append("")
            elif "이번주봉" in structured_data:
                candle = structured_data["이번주봉"]
                text_parts.append(f"📊 이번 주 봉 요약")
                text_parts.append(f"• 종가: {candle.get('종가', 'N/A')}원")
                text_parts.append(f"• 등락률: {candle.get('등락률', 'N/A')}%")
                text_parts.append(f"• 거래량: {candle.get('거래량', 'N/A')}주")
                text_parts.append(f"• 주요 특징: {candle.get('주요특징', 'N/A')}")
                text_parts.append("")
            elif "이번월봉" in structured_data:
                candle = structured_data["이번월봉"]
                text_parts.append(f"📊 이번 월봉 요약")
                text_parts.append(f"• 종가: {candle.get('종가', 'N/A')}원")
                text_parts.append(f"• 등락률: {candle.get('등락률', 'N/A')}%")
                text_parts.append(f"• 거래량: {candle.get('거래량', 'N/A')}주")
                text_parts.append(f"• 주요 특징: {candle.get('주요특징', 'N/A')}")
                text_parts.append("")
            
            # 핵심 기술적 지표
            if "핵심기술적지표" in structured_data:
                tech = structured_data["핵심기술적지표"]
                text_parts.append(f"🔍 핵심 기술적 분석 지표")
                for key, value in tech.items():
                    text_parts.append(f"• {key}: {value}")
                text_parts.append("")
            
            # 세부 분석
            if "세부분석" in structured_data:
                detail = structured_data["세부분석"]
                text_parts.append(f"📈 세부 분석")
                for section, content in detail.items():
                    text_parts.append(f"📊 {section}")
                    if isinstance(content, dict):
                        for key, value in content.items():
                            text_parts.append(f"  • {key}: {value}")
                    else:
                        text_parts.append(f"  • {content}")
                text_parts.append("")
            
            # 투자 아이디어
            if "단기투자아이디어" in structured_data:
                idea = structured_data["단기투자아이디어"]
                text_parts.append(f"💡 단기 투자 아이디어")
                for key, value in idea.items():
                    text_parts.append(f"• {key}: {value}")
                text_parts.append("")
            elif "중기투자아이디어" in structured_data:
                idea = structured_data["중기투자아이디어"]
                text_parts.append(f"💡 중기 투자 아이디어")
                for key, value in idea.items():
                    text_parts.append(f"• {key}: {value}")
                text_parts.append("")
            elif "장기투자아이디어" in structured_data:
                idea = structured_data["장기투자아이디어"]
                text_parts.append(f"💡 장기 투자 아이디어")
                for key, value in idea.items():
                    text_parts.append(f"• {key}: {value}")
                text_parts.append("")
            
            return "\n".join(text_parts)
            
        except Exception as e:
            print(f"⚠️ 구조화된 데이터를 텍스트로 변환 중 오류: {e}")
            return str(structured_data)

    def _parse_markdown_to_word(self, text: str, doc) -> None:
        """
        마크다운 텍스트를 Word 문서 형식으로 파싱하여 추가
        
        Args:
            text (str): 마크다운 형태의 텍스트
            doc: Word 문서 객체
        """
        try:
            import re
            
            # **내용** 형태를 볼드체로 변환하는 정규식
            bold_pattern = r'\*\*(.*?)\*\*'
            
            # 줄바꿈으로 분리
            lines = text.split('\n')
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # **내용** 패턴을 찾아서 볼드체로 변환
                if re.search(bold_pattern, line):
                    # 볼드체가 포함된 줄 처리
                    para = doc.add_paragraph()
                    
                    # **내용** 패턴으로 분할
                    parts = re.split(bold_pattern, line)
                    
                    for i, part in enumerate(parts):
                        if i % 2 == 0:
                            # 일반 텍스트
                            if part:
                                run = para.add_run(part)
                                run.font.name = '맑은 고딕'
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                        else:
                            # 볼드체 텍스트
                            if part:
                                run = para.add_run(part)
                                run.font.name = '맑은 고딕'
                                run.bold = True
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                else:
                    # 볼드체가 없는 일반 줄 처리
                    if line.startswith('📊') or line.startswith('📈') or line.startswith('📉') or line.startswith('💡') or line.startswith('🔍') or line.startswith('⚠️') or line.startswith('✅'):
                        para = doc.add_heading(line, level=2)
                    elif line.startswith('•') or line.startswith('-') or line.startswith('*'):
                        para = doc.add_paragraph(line)
                    else:
                        para = doc.add_paragraph(line)
                    
                    # 한글 폰트 적용
                    for run in para.runs:
                        run.font.name = '맑은 고딕'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                        
        except Exception as e:
            print(f"⚠️ 마크다운 파싱 중 오류: {e}")
            # 오류 발생 시 일반 텍스트로 처리
            para = doc.add_paragraph(text)
            for run in para.runs:
                run.font.name = '맑은 고딕'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    def create_word_document_hybrid(self, result: Dict[str, Any], chart_image_path: str, output_path: str, chart_type: str = "일봉") -> bool:
        """
        하이브리드 방식으로 분석 결과를 Word 문서로 생성 (개별 종목용)
        AI 피드백을 그대로 저장하여 프롬프트 변경에 유연하게 대응
        
        Args:
            result (Dict[str, Any]): 분석 결과 (JSON 파싱된 구조화 데이터 또는 원본 AI 응답)
            chart_image_path (str): 차트 이미지 경로
            output_path (str): 저장할 Word 파일 경로
            chart_type (str): 차트 유형 (일봉/주봉/월봉)
            
        Returns:
            bool: 저장 성공 여부
        """
        try:
            # 출력 디렉토리 생성
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # Word 문서 생성
            doc = Document()
            
            # 한글 폰트 설정을 위한 스타일 설정
            from docx.oxml.ns import qn
            
            # 제목 설정 (개별 분석 리포트)
            title = doc.add_heading('주식 차트 분석 리포트', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 제목에 한글 폰트 적용
            for run in title.runs:
                run.font.name = '맑은 고딕'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                run.font.size = Pt(14)
            
            # 종목 정보 (파일명에서 추출)
            heading1 = doc.add_heading('종목 정보', level=1)
            for run in heading1.runs:
                run.font.name = '맑은 고딕'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            
            # 파일명에서 종목 정보 추출
            filename = os.path.basename(output_path)
            stock_name, stock_code = self.stock_mapper.extract_stock_info_from_filename(filename)
            
            # 종목 정보 테이블 생성
            table = doc.add_table(rows=2, cols=4)
            table.style = 'Table Grid'
            
            # 헤더 행 설정
            header_cells = table.rows[0].cells
            header_cells[0].text = '종목명'
            header_cells[1].text = '종목번호'
            header_cells[2].text = '분석일시'
            header_cells[3].text = '차트유형'
            
            # 데이터 행 설정
            data_cells = table.rows[1].cells
            data_cells[0].text = stock_name
            data_cells[1].text = stock_code
            # 날짜 형식을 YY-MM-DD (ddd) HH 형태로 변경
            now = datetime.now()
            weekday_korean = ['월', '화', '수', '목', '금', '토', '일']
            formatted_date = f"{now.strftime('%y.%m.%d')}({weekday_korean[now.weekday()]}) {now.strftime('%H')}시"
            data_cells[2].text = formatted_date
            data_cells[3].text = chart_type
            
            # 테이블 전체에 한글 폰트 적용
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = '맑은 고딕'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            
            # 차트 이미지 추가
            heading2 = doc.add_heading('차트 이미지', level=1)
            for run in heading2.runs:
                run.font.name = '맑은 고딕'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            
            if os.path.exists(chart_image_path):
                doc.add_picture(chart_image_path, width=Inches(6))
                doc.add_paragraph()
            
            # 거래대금/거래대금 및 순위 정보 추가 (모든 차트 타입에 적용)
            if isinstance(result, dict) and "종목정보" in result:
                try:
                    # 거래 정보 추출
                    total_amount = result["종목정보"].get("거래대금", "N/A")
                    turnover_rate = result["종목정보"].get("거래율", "N/A")
                    ranking = result["종목정보"].get("순위", "N/A")
                    outstanding_shares = result["종목정보"].get("유통주식수", "N/A")
                    volume = result["종목정보"].get("거래대금", "N/A")
                    trading_type = result["종목정보"].get("거래타입", "거래대금")
                    
                    # 차트 타입별 기간 텍스트 설정
                    period_text = {
                        "일봉": "일일",
                        "주봉": "주간", 
                        "월봉": "월간"
                    }.get(chart_type, "일일")
                    
                    # 거래 타입별 문구 생성
                    if trading_type == "거래율" and turnover_rate != "N/A":
                        # 거래율 기준 문구 - 사용자 요구사항에 맞게 수정
                        rate_value = turnover_rate.replace("%", "").strip()
                        if rate_value.replace(".", "").isdigit() and volume != "N/A" and outstanding_shares != "N/A":
                            # 거래량과 유통주식수에서 숫자만 추출
                            volume_clean = volume.replace(",", "").replace("주", "").strip()
                            shares_clean = outstanding_shares.replace(",", "").replace("주", "").strip()
                            
                            if volume_clean.isdigit() and shares_clean.isdigit():
                                trading_info_text = f"위 주식의 {period_text} 거래량은 {volume}로 유통주식수 {outstanding_shares} 대비 {turnover_rate}로 전체 종목 중 상위 {ranking}를 차지하여 분석 대상에 포함되었습니다."
                            else:
                                trading_info_text = f"위 주식의 {period_text} 거래률은 {turnover_rate}로 전체 종목 중 상위 {ranking}를 차지하여 분석 대상에 포함되었습니다."
                        else:
                            trading_info_text = f"위 주식의 {period_text} 거래율 정보를 확인할 수 없어 분석 대상에 포함되었습니다."
                    else:
                        # 거래대금 기준 문구 - 사용자 요구사항에 맞게 수정
                        if total_amount != "N/A":
                            trading_info_text = f"위 주식의 {period_text} 거래대금은 {total_amount}으로 전체 종목 중 상위 {ranking}를 차지하여 분석 대상에 포함되었습니다."
                        else:
                            trading_info_text = f"위 주식의 {period_text} 거래대금 정보를 확인할 수 없어 분석 대상에 포함되었습니다."
                    
                    # 문단 추가
                    trading_info_para = doc.add_paragraph(trading_info_text)
                    for run in trading_info_para.runs:
                        run.font.name = '맑은 고딕'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                    
                    print(f"✅ 거래정보 추가 완료 ({chart_type}, {trading_type}): {trading_info_text}")
                    
                except Exception as e:
                    print(f"⚠️ 거래정보 추가 중 오류: {e}")
                    # 오류 발생 시에도 DOCX 생성은 계속 진행
            
            # AI 분석 결과 (하이브리드 방식: AI 피드백을 그대로 저장)
            heading_analysis = doc.add_heading(f'{stock_name} 차트 분석 결과', level=1)
            for run in heading_analysis.runs:
                run.font.name = '맑은 고딕'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            
            # AI 응답 처리 (JSON 파싱 없이 원본 텍스트 그대로 저장)
            ai_response = ""
            print(f"🔍 DOCX 생성 디버깅: result 타입 = {type(result)}")
            if isinstance(result, dict):
                print(f"🔍 DOCX 생성 디버깅: result 키들 = {list(result.keys())}")
                # JSON 파싱된 결과인 경우 원본 AI 응답 추출 시도
                if "AI분석결과" in result:
                    ai_response = result["AI분석결과"]
                    print(f"🔍 DOCX 생성 디버깅: AI분석결과 찾음, 길이 = {len(ai_response)}")
                elif "original_ai_response" in result:
                    ai_response = result["original_ai_response"]
                    print(f"🔍 DOCX 생성 디버깅: original_ai_response 찾음, 길이 = {len(ai_response)}")
                elif "ai_response" in result:
                    ai_response = result["ai_response"]
                    print(f"🔍 DOCX 생성 디버깅: ai_response 찾음, 길이 = {len(ai_response)}")
                else:
                    # 구조화된 데이터를 텍스트로 변환
                    ai_response = self._convert_structured_data_to_text(result)
                    print(f"🔍 DOCX 생성 디버깅: 구조화된 데이터 변환, 길이 = {len(ai_response)}")
            elif isinstance(result, str):
                # 문자열인 경우 그대로 사용
                ai_response = result
                print(f"🔍 DOCX 생성 디버깅: 문자열 결과, 길이 = {len(ai_response)}")
            else:
                ai_response = str(result)
                print(f"🔍 DOCX 생성 디버깅: 기타 타입 변환, 길이 = {len(ai_response)}")
            
            print(f"🔍 DOCX 생성 디버깅: 최종 ai_response 길이 = {len(ai_response)}")
            
            # AI 응답을 워드 문서에 추가 (마크다운 파싱 적용)
            if ai_response:
                # JSON 코드 블록 제거 (있는 경우)
                cleaned_response = ai_response.strip()
                if cleaned_response.startswith('```json'):
                    cleaned_response = cleaned_response[7:]
                elif cleaned_response.startswith('```'):
                    cleaned_response = cleaned_response[3:]
                if cleaned_response.endswith('```'):
                    cleaned_response = cleaned_response[:-3]
                cleaned_response = cleaned_response.strip()
                
                # 마크다운 파싱을 통해 Word 문서에 추가
                self._parse_markdown_to_word(cleaned_response, doc)
            else:
                # AI 응답이 없는 경우
                para = doc.add_paragraph("AI 분석 결과가 없습니다.")
                for run in para.runs:
                    run.font.name = '맑은 고딕'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            
            # 문서 저장
            doc.save(output_path)
            print(f"📄 하이브리드 방식 Word 문서 저장 완료: {output_path}")
            return True
            
        except Exception as e:
            print(f"❌ 하이브리드 방식 Word 문서 생성 중 오류: {e}")
            return False

    def _create_fallback_result(self, stock_name: str, chart_type: str, ai_response: str, error_type: str, stock_code: str = "000000", chart_data: Optional[pd.DataFrame] = None, additional_info: Optional[Dict[str, Any]] = None, trading_type: str = "거래대금") -> Dict[str, Any]:
        """JSON 파싱 실패 시 대체 결과 생성"""
        
        # 무조건 stocks 테이블에서 종목명 조회 (통일된 처리)
        if stock_code and stock_code != "000000":
            mapped_name = self.stock_mapper.get_stock_name(stock_code)
            if mapped_name != stock_code:  # DB에서 한글 종목명을 찾은 경우
                stock_name = mapped_name
            else:
                # DB에서도 찾지 못한 경우 종목코드를 종목명으로 사용
                stock_name = f"알 수 없음 ({stock_code})"
        else:
            stock_name = "알 수 없음"
        
        
        result = {
            "종목정보": {
                "종목명": stock_name,
                "종목번호": stock_code,
                "분석일시": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "차트유형": chart_type,
                "거래타입": trading_type,
                "파싱상태": error_type
            },
            "AI분석결과": ai_response
        }
        
        # 거래정보 추가 (JSON 파싱 실패)
        # AI 응답이 JSON 형식이 아니거나 파싱에 실패한 경우 실행됨
        # 현재 가장 많이 사용되는 경로 (AI 응답이 JSON 형식이 아니기 때문)
        # 거래일, 거래대금, 순위 정보를 DB에서 조회하여 추가
        result = self._add_trading_info_to_result(result, stock_code, "파싱실패", chart_type, ai_response, trading_type)
        
        return result

    def _parse_weekly_date(self, date_text: str) -> str:
        """
        주봉 날짜 텍스트를 파싱하여 주 시작일(월요일) 반환 (한국식 주차 기준)
        
        Args:
            date_text (str): "2025년 35주차" 형식의 날짜 텍스트
            
        Returns:
            str: "YYYY-MM-DD" 형식의 주 시작일, 파싱 실패 시 None
        """
        try:
            # 주차 문자열 파싱
            result = parse_week_string(date_text)
            if not result:
                print(f"⚠️ 주봉 날짜 형식 파싱 실패: {date_text}")
                return None
            
            year, week = result
            
            # ISO 8601 표준으로 주 시작일 계산
            week_start = get_week_start_date(year, week)
            
            print(f"📅 주봉 날짜 파싱 (ISO 8601): {date_text} → {week_start}")
            return week_start.strftime('%Y-%m-%d')
            
        except Exception as e:
            print(f"❌ 주봉 날짜 파싱 오류: {e}")
            return None
    
    def _parse_monthly_date(self, date_text: str) -> str:
        """
        월봉 날짜 텍스트를 파싱하여 년월 반환
        
        Args:
            date_text (str): "2025년 9월" 형식의 날짜 텍스트
            
        Returns:
            str: "YYYY-MM" 형식의 년월, 파싱 실패 시 None
        """
        try:
            # "2025년 9월" 형식 파싱
            pattern = r'(\d{4})년\s*(\d{1,2})월'
            match = re.search(pattern, date_text)
            
            if not match:
                print(f"⚠️ 월봉 날짜 형식 파싱 실패: {date_text}")
                return None
            
            year = int(match.group(1))
            month = int(match.group(2))
            
            # YYYY-MM 형식으로 변환
            year_month = f"{year:04d}-{month:02d}"
            
            print(f"📅 월봉 날짜 파싱: {date_text} → {year_month}")
            return year_month
            
        except Exception as e:
            print(f"❌ 월봉 날짜 파싱 오류: {e}")
            return None


    def _get_individual_stock_trading_data(self, stock_code: str, target_date: str, period_type: str) -> Optional[Dict[str, Any]]:
        """
        개별 종목의 거래량과 거래율 데이터를 직접 조회
        
        Args:
            stock_code (str): 종목코드
            target_date (str): 대상 날짜/기간
            period_type (str): 기간 타입 (daily, weekly, monthly)
            
        Returns:
            Dict[str, Any]: 종목별 거래 데이터
        """
        try:
            from database_config import DatabaseManager
            db = DatabaseManager()
            
            if not db.connect():
                print(f"⚠️ DB 연결 실패 - {stock_code} 거래정보 조회 불가")
                return None
            
            # 종목명 조회
            stock_name_query = "SELECT stock_name, market_type FROM stocks WHERE stock_code = %s"
            stock_info = db.fetch_one(stock_name_query, (stock_code,))
            
            if not stock_info:
                print(f"⚠️ {stock_code} 종목 정보를 찾을 수 없음")
                db.disconnect()
                return None
            
            stock_name = stock_info['stock_name']
            market_type = stock_info['market_type']
            
            if period_type == "daily":
                # 일봉 데이터 조회 - close 가격 및 거래대금 추가
                query = """
                SELECT 
                    SUM(volume) as total_volume,
                    SUM(volume * close) as total_amount,
                    AVG(outstanding_shares) as avg_shares,
                    AVG(close) as avg_close,
                    COUNT(trade_date) as trading_days
                FROM daily_data 
                WHERE stock_code = %s AND trade_date = %s
                """
                params = (stock_code, target_date)
                
            elif period_type == "weekly":
                # 주봉 데이터 조회 (월요일부터 금요일까지) - week_calculator 모듈 사용
                try:
                    target_datetime = datetime.strptime(target_date, '%Y-%m-%d')
                    year, week = get_week_number(target_datetime)
                    week_start = WeekCalculator.get_week_start_date(year, week)
                    week_end = WeekCalculator.get_week_end_date(year, week)
                    week_end_str = week_end.strftime('%Y-%m-%d')
                    print(f"📅 주봉 데이터 조회 기간: {week_start.strftime('%Y-%m-%d')} ~ {week_end_str}")
                except Exception as e:
                    print(f"❌ 주차 계산 실패: {e}")
                    return None
                query = """
                SELECT 
                    SUM(volume) as total_volume,
                    SUM(volume * close) as total_amount,
                    AVG(outstanding_shares) as avg_shares,
                    AVG(close) as avg_close,
                    COUNT(trade_date) as trading_days
                FROM daily_data 
                WHERE stock_code = %s 
                AND trade_date BETWEEN %s AND %s
                AND WEEKDAY(trade_date) < 5
                """
                params = (stock_code, week_start, week_end)
                
            elif period_type == "monthly":
                # 월봉 데이터 조회 - 거래대금 정확한 계산
                # target_date를 년도와 월로 분리
                year, month = target_date.split('-')
                query = """
                SELECT 
                    SUM(volume) as total_volume,
                    SUM(volume * close) as total_amount,
                    AVG(outstanding_shares) as avg_shares,
                    AVG(close) as avg_close,
                    COUNT(trade_date) as trading_days
                FROM daily_data 
                WHERE stock_code = %s 
                AND YEAR(trade_date) = %s AND MONTH(trade_date) = %s
                AND WEEKDAY(trade_date) < 5
                """
                params = (stock_code, year, month)
            
            else:
                print(f"⚠️ 잘못된 기간 타입: {period_type}")
                db.disconnect()
                return None
            
            result = db.fetch_one(query, params)
            db.disconnect()
            
            if not result or result['total_volume'] is None:
                print(f"⚠️ {stock_code} {period_type} 거래 데이터 없음")
                return None
            
            total_volume = float(result['total_volume']) if result['total_volume'] else 0
            avg_shares = float(result['avg_shares']) if result['avg_shares'] else 0
            avg_close = float(result['avg_close']) if result['avg_close'] else 0
            trading_days = int(result['trading_days']) if result['trading_days'] else 0
            
            # 거래대금 계산 (DB에서 직접 계산된 값 사용)
            total_amount = float(result['total_amount']) if result['total_amount'] else 0
            
            # 거래율 계산
            turnover_rate = 0.0
            if avg_shares > 0 and total_volume > 0:
                turnover_rate = (total_volume / avg_shares) * 100
            
            return {
                'stock_code': stock_code,
                'stock_name': stock_name,
                'market_type': market_type,
                'volume': total_volume,
                'outstanding_shares': avg_shares,
                'close_price': avg_close,
                'total_amount': total_amount,
                'turnover_rate': turnover_rate,
                'trading_days': trading_days
            }
            
        except Exception as e:
            print(f"❌ {stock_code} 개별 거래 데이터 조회 실패: {e}")
            return None



    def _calculate_week_number(self, date: datetime) -> str:
        """
        ISO 8601 표준 주차 계산 (통일된 주차 계산 방식)
        
        Args:
            date (datetime): 계산할 날짜
            
        Returns:
            str: "YYYY년 XX주차" 형식의 문자열
        """
        try:
            return get_week_number_string(date, "YYYY년 W주차")
        except Exception as e:
            print(f"⚠️ 주차 계산 중 오류: {e}")
            return f"{date.year}년 1주차"

    def _extract_trading_date_from_ai_result(self, ai_result_text: str, chart_type: str) -> str:
        """
        AI 분석 결과에서 거래일/거래기간 추출 (다양한 패턴 지원)
        
        Args:
            ai_result_text (str): AI 분석 결과 텍스트
            chart_type (str): 차트 타입 (일봉, 주봉, 월봉)
            
        Returns:
            str: 추출된 거래일/거래기간 (추출 실패 시 "N/A")
        """
        try:
            import re
            
            if chart_type == "일봉":
                # 일봉: "2025-09-03" 형식 추출
                date_pattern = r'(\d{4}-\d{2}-\d{2})'
                date_match = re.search(date_pattern, ai_result_text)
                if date_match:
                    extracted_date = date_match.group(1)
                    print(f"📅 일봉 거래일 추출: {extracted_date}")
                    return extracted_date
                    
            elif chart_type == "주봉":
                # 주봉: 다양한 패턴으로 주차 추출 시도
                extracted_week = self._extract_weekly_date_flexible(ai_result_text)
                if extracted_week != "N/A":
                    print(f"📅 주봉 주차 추출: {extracted_week}")
                    # 성공한 경우 캐시에 저장
                    self._weekly_trading_date_cache = extracted_week
                    return extracted_week
                else:
                    # 추출 실패 시 캐시된 거래일 사용
                    if self._weekly_trading_date_cache:
                        print(f"📅 주봉 주차 추출 실패 - 캐시된 거래일 사용: {self._weekly_trading_date_cache}")
                        return self._weekly_trading_date_cache
                    else:
                        print(f"⚠️ 주봉 주차 추출 실패 - 캐시도 없음")
                        return "N/A"
                    
            elif chart_type == "월봉":
                # 월봉: "2025년 8월" 형식 추출
                month_pattern = r'(\d{4}년 \d{1,2}월)'
                month_match = re.search(month_pattern, ai_result_text)
                if month_match:
                    extracted_month = month_match.group(1)
                    print(f"📅 월봉 거래월 추출: {extracted_month}")
                    return extracted_month
            
            print(f"⚠️ {chart_type} 거래일/기간 추출 실패")
            return "N/A"
            
        except Exception as e:
            print(f"❌ 거래일 추출 중 오류: {e}")
            return "N/A"
    
    def _extract_weekly_date_flexible(self, ai_result_text: str) -> str:
        """
        주봉 거래일 추출 (다양한 텍스트 패턴 지원)
        
        Args:
            ai_result_text (str): AI 분석 결과 텍스트
            
        Returns:
            str: 추출된 주차 (추출 실패 시 "N/A")
        """
        import re
        from datetime import datetime
        
        # 패턴 1: "2025년 35주차" 형식 (숫자 주차)
        pattern1 = r'(\d{4}년 \d{1,2}주차)'
        match1 = re.search(pattern1, ai_result_text)
        if match1:
            return match1.group(1)
        
        # 패턴 2: "2025년 9월 1주차" 형식 (월+주차)
        pattern2 = r'(\d{4}년 \d{1,2}월 \d{1,2}주차)'
        match2 = re.search(pattern2, ai_result_text)
        if match2:
            return match2.group(1)
        
        # 패턴 3: "2025년 8월 25일 ~ 2025년 9월 1일" 형식
        pattern3 = r'(\d{4}년 \d{1,2}월 \d{1,2}일 ~ \d{4}년 \d{1,2}월 \d{1,2}일)'
        match3 = re.search(pattern3, ai_result_text)
        if match3:
            return match3.group(1)
        
        # 패턴 4: "**1) 거래 요약:** 2025년 35주차(...)" 형식
        pattern4 = r'\*\*1\)\s*거래 요약:\*\*\s*(\d{4}년 \d{1,2}주차)'
        match4 = re.search(pattern4, ai_result_text)
        if match4:
            return match4.group(1)
        
        # 패턴 5: "1. **거래 요약:** 2025년 35주차(...)" 형식
        pattern5 = r'1\.\s*\*\*거래 요약:\*\*\s*(\d{4}년 \d{1,2}주차)'
        match5 = re.search(pattern5, ai_result_text)
        if match5:
            return match5.group(1)
        
        # 패턴 6: "2025년 9월 1일 기준" 형식에서 주차 계산 (마지막에 시도)
        pattern6 = r'(\d{4}년 \d{1,2}월 \d{1,2}일)'
        match6 = re.search(pattern6, ai_result_text)
        if match6:
            try:
                date_str = match6.group(1)
                # "2025년 9월 1일" -> "2025-09-01" 변환
                date_obj = datetime.strptime(date_str, "%Y년 %m월 %d일")
                # 해당 날짜의 주차 계산 (ISO 8601 표준)
                week_info = self._calculate_week_number(date_obj)
                return week_info
            except Exception as e:
                print(f"⚠️ 날짜 파싱 실패: {e}")
        
        return "N/A"
    
    def clear_weekly_trading_date_cache(self):
        """주봉 거래일 캐시 초기화 (분석 완료 시 호출)"""
        self._weekly_trading_date_cache = None
        print("🧹 주봉 거래일 캐시 초기화 완료")

    def _get_trading_info_from_db(self, stock_code: str, chart_type: str = "일봉", ai_result_text: str = "", trading_type: str = "거래대금") -> Dict[str, Any]:
        """
        DB에서 거래일, 거래대금, 거래율, 순위, 유통주식수, 거래량 정보 조회
        
        Args:
            stock_code (str): 종목코드 (6자리)
            chart_type (str): 차트 타입 (일봉, 주봉, 월봉)
            ai_result_text (str): AI 분석 결과 텍스트 (거래일 추출용)
            trading_type (str): 거래 타입 (거래대금, 거래율)
            
        Returns:
            Dict[str, Any]: 거래정보 딕셔너리
                - 거래일: 추출된 거래일/거래기간
                - 거래대금: volume * close (억원 단위)
                - 거래율: 거래율 값 (퍼센트, 소수점 2자리)
                - 순위: 거래대금/거래율 기준 순위 (N위 형식)
                - 유통주식수: 해당 종목의 유통주식수 (주 단위)
                - 거래량: 거래일 기준 누적 거래량 (주 단위)
                
        Note:
            - 거래일: AI 분석 결과에서 추출 (차트 타입별로 다른 형식)
            - 거래대금: volume * close (억원 단위로 표시)
            - 거래율: volume / outstanding_shares * 100 (퍼센트)
            - 순위 계산: 거래대금/거래율 기준 내림차순 정렬
            - 유통주식수: outstanding_shares 값 (주 단위)
            - 거래량: volume 값 (주 단위)
            - DB 연결 실패 시 모든 값이 "N/A"로 반환
        """
        try:
            print(f"🔗 database_config 모듈 import 시도...")
            from database_config import DatabaseManager
            print(f"✅ database_config 모듈 import 성공")
            db = DatabaseManager()
            print(f"✅ DatabaseManager 인스턴스 생성 성공")
            
            if not db.connect():
                print(f"⚠️ DB 연결 실패 - 거래정보 조회 불가")
                return {
                    "거래일": "N/A",
                    "거래대금": "N/A",
                    "거래율": "N/A",
                    "순위": "N/A"
                }
            
            # 1. AI 분석 결과에서 거래일/거래기간 추출
            extracted_date = "N/A"
            if ai_result_text and chart_type in ["일봉", "주봉", "월봉"]:
                extracted_date = self._extract_trading_date_from_ai_result(ai_result_text, chart_type)
                print(f"📅 AI 결과에서 추출된 거래일/기간: {extracted_date}")
            
            # 2. 자체적으로 거래대금/거래율 순위 조회 (VolumeRankingDataManager 의존성 제거)
            try:
                
                # 차트 타입과 거래 타입에 따라 적절한 함수 호출
                ranking_data = []
                target_date = None
                
                if chart_type == "일봉":
                    # 일봉의 경우 추출된 날짜 사용, 없으면 최신 날짜 사용
                    if extracted_date != "N/A":
                        try:
                            target_date = datetime.strptime(extracted_date, "%Y-%m-%d").strftime('%Y-%m-%d')
                        except ValueError:
                            print(f"⚠️ 잘못된 날짜 형식: {extracted_date}")
                            target_date = None
                    
                    if not target_date:
                        # 최신 거래일 조회
                        latest_date_query = "SELECT MAX(trade_date) as latest_date FROM daily_data WHERE stock_code = %s"
                        latest_date_result = db.fetch_one(latest_date_query, (stock_code,))
                        if latest_date_result and latest_date_result['latest_date']:
                            target_date = latest_date_result['latest_date'].strftime('%Y-%m-%d')
                    
                    if target_date:
                        # 개별 종목 데이터 조회 (VolumeRankingDataManager 대신 직접 조회)
                        target_stock_data = self._get_individual_stock_trading_data(stock_code, target_date, "daily")
                        print(f"📊 {stock_code} 개별 거래 데이터 조회 완료")
                        
                        # 전체 순위 데이터 조회 (거래타입 전달)
                        ranking_data = self.ranking_calculator.get_volume_ranking(target_date, "일봉", limit=50, trading_type=trading_type)
                        print(f"📊 일봉 전체 순위 데이터 조회 완료: {len(ranking_data)}개 종목 ({trading_type} 기준)")
                
                elif chart_type == "주봉":
                    # 주봉의 경우 주간 시작일 계산
                    week_start = None
                    week_display = "N/A"  # 최종 표시용 주차 정보
                    
                    if extracted_date != "N/A" and "~" in extracted_date:
                        # "2025년 8월 25일 ~ 2025년 9월 1일" 형식에서 시작일 추출
                        week_start_str = extracted_date.split("~")[0].strip()
                        # 추출된 기간을 그대로 사용 (주차 변경하지 않음)
                        week_display = extracted_date  # 추출된 기간 그대로 사용
                        try:
                            start_date = datetime.strptime(week_start_str, "%Y년 %m월 %d일")
                            # week_start를 YYYY-MM-DD 형식으로 변환
                            week_start = start_date.strftime('%Y-%m-%d')
                        except ValueError:
                            week_display = "N/A"
                            week_start = None
                    elif extracted_date != "N/A" and "주차" in extracted_date:
                        # "2025년 35주차" 형식 파싱 - 추출된 주차를 그대로 사용
                        week_display = extracted_date  # 추출된 주차 그대로 사용
                        week_start = self._parse_weekly_date(extracted_date)
                        if not week_start:
                            # 파싱 실패 시 N/A 반환 (주차 변경하지 않음)
                            print(f"⚠️ 주차 파싱 실패: {extracted_date}")
                            db.disconnect()
                            return {
                                "거래일": week_display,
                                "거래대금": "N/A",
                                "거래율": "N/A",
                                "순위": "N/A"
                            }
                        
                        # 파싱된 주차에 거래 데이터가 있는지 확인
                        if week_start:
                            # 해당 주차에 거래 데이터가 있는지 확인 (직접 DB 조회)
                            test_query = """
                                SELECT COUNT(*) as count 
                                FROM daily_data 
                                WHERE stock_code = %s 
                                AND trade_date BETWEEN %s AND DATE_ADD(%s, INTERVAL 6 DAY)
                            """
                            test_result = db.fetch_one(test_query, (stock_code, week_start, week_start))
                            if not test_result or test_result['count'] == 0:
                                print(f"⚠️ {week_start} 주차에 거래 데이터가 없음")
                                # 데이터가 없으면 N/A 반환
                                db.disconnect()
                                return {
                                    "거래일": week_display,
                                    "거래대금": "N/A",
                                    "거래율": "N/A",
                                    "순위": "N/A",
                                    "유통주식수": "N/A",
                                    "거래대금": "N/A"
                                }
                            else:
                                print(f"✅ {week_start} 주차에 거래 데이터 존재")
                    else:
                        # 추출된 날짜가 없으면 N/A 반환 (주차 변경하지 않음)
                        print(f"⚠️ 주봉 거래일을 추출할 수 없음: {extracted_date}")
                        db.disconnect()
                        return {
                            "거래일": "N/A",
                            "거래대금": "N/A",
                            "거래율": "N/A",
                            "순위": "N/A"
                        }
                    
                    target_date = week_start
                    if week_start:
                        # 개별 종목 데이터 조회
                        target_stock_data = self._get_individual_stock_trading_data(stock_code, week_start, "weekly")
                        print(f"📊 {stock_code} 주봉 개별 거래 데이터 조회 완료")
                        
                        # 전체 순위 데이터 조회 (거래타입 전달)
                        ranking_data = self.ranking_calculator.get_volume_ranking(week_start, "주봉", limit=50, trading_type=trading_type)
                        print(f"📊 주봉 전체 순위 데이터 조회 완료: {len(ranking_data)}개 종목 ({trading_type} 기준)")
                    else:
                        ranking_data = []
                
                elif chart_type == "월봉":
                    # 월봉의 경우 월간 시작일 계산
                    year_month = None  # 초기화
                    
                    if extracted_date != "N/A" and "~" in extracted_date:
                        # "2024-01-01~2024-01-31" 형식에서 시작일 추출
                        month_start = extracted_date.split("~")[0]
                        year_month = month_start[:7]  # "2024-01"
                        print(f"📅 월봉 기간 형식 파싱: {extracted_date} → {year_month}")
                    elif extracted_date != "N/A" and "월" in extracted_date:
                        # "2025년 9월" 형식 파싱 - 추출된 월을 그대로 사용
                        year_month = self._parse_monthly_date(extracted_date)
                        if not year_month:
                            # 파싱 실패 시 N/A 반환 (월 변경하지 않음)
                            print(f"⚠️ 월 파싱 실패: {extracted_date}")
                            db.disconnect()
                            return {
                                "거래일": extracted_date,
                                "거래대금": "N/A",
                                "거래율": "N/A",
                                "순위": "N/A"
                            }
                        print(f"📅 월봉 년월 형식 파싱: {extracted_date} → {year_month}")
                    else:
                        # 추출된 날짜가 없으면 N/A 반환 (월 변경하지 않음)
                        print(f"⚠️ 월봉 거래일을 추출할 수 없음: {extracted_date}")
                        db.disconnect()
                        return {
                            "거래일": "N/A",
                            "거래대금": "N/A",
                            "거래율": "N/A",
                            "순위": "N/A"
                        }
                    
                    target_date = year_month
                    print(f"📊 월봉 처리 - year_month: {year_month}, target_date: {target_date}")
                    
                    if year_month:
                        # 개별 종목 데이터 조회
                        target_stock_data = self._get_individual_stock_trading_data(stock_code, year_month, "monthly")
                        print(f"📊 {stock_code} 월봉 개별 거래 데이터 조회 완료")
                        
                        # 전체 순위 데이터 조회 (거래타입 전달)
                        ranking_data = self.ranking_calculator.get_volume_ranking(year_month, "월봉", limit=50, trading_type=trading_type)
                        print(f"📊 월봉 전체 순위 데이터 조회 완료: {len(ranking_data)}개 종목 ({trading_type} 기준)")
                    else:
                        print(f"⚠️ year_month가 None이어서 데이터 조회 불가")
                        ranking_data = []
                
                # 개별 종목 데이터 확인
                if not target_stock_data:
                    print(f"⚠️ {stock_code} 종목의 {chart_type} 데이터를 찾을 수 없음")
                    db.disconnect()
                    # 주봉의 경우 week_display 사용, 다른 경우는 기존 로직 유지
                    if chart_type == "주봉" and 'week_display' in locals():
                        display_date = week_display
                    else:
                        display_date = extracted_date if extracted_date != "N/A" else target_date
                    
                    return {
                        "거래일": display_date,
                        "거래대금": "N/A",
                        "거래율": "N/A",
                        "순위": "N/A",
                        "유통주식수": "N/A",
                        "거래대금": "N/A"
                    }
                
                
                # 순위 계산 (전체 순위 데이터에서 정확한 순위 찾기)
                print(f"🔍 순위 계산 시작: {stock_code} ({chart_type}, {trading_type})")
                ranking = 999  # 기본값 (순위를 찾을 수 없는 경우)
                
                if ranking_data and len(ranking_data) > 0:
                    # 전체 순위 데이터에서 해당 종목의 순위 찾기
                    for stock_data in ranking_data:
                        if stock_data.get('stock_code') == stock_code:
                            ranking = int(stock_data.get('ranking', 999))
                            print(f"📊 {trading_type} 기준 순위: {ranking}위 (전체 {len(ranking_data)}개 종목 중)")
                            break
                    else:
                        # 순위에 없으면 거래량이 0이거나 데이터 없음
                        ranking = 999
                        print(f"📊 {stock_code}이 순위에 없음 (거래대금 0 또는 데이터 없음)")
                else:
                    print(f"⚠️ 전체 순위 데이터가 없어 순위 계산 불가")
                
                # 거래대금 계산 (안전한 None 처리)
                volume_value = target_stock_data.get('volume', 0)
                if volume_value is None or volume_value == '':
                    volume_value = 0
                
                try:
                    volume = float(volume_value)
                except (ValueError, TypeError):
                    print(f"⚠️ {stock_code} volume 값 변환 실패: {volume_value}")
                    volume = 0
                
                # 유통주식수 계산
                outstanding_shares_value = target_stock_data.get('outstanding_shares', 0)
                if outstanding_shares_value is None or outstanding_shares_value == '':
                    outstanding_shares = 0
                else:
                    try:
                        outstanding_shares = float(outstanding_shares_value)
                    except (ValueError, TypeError):
                        print(f"⚠️ {stock_code} outstanding_shares 값 변환 실패: {outstanding_shares_value}")
                        outstanding_shares = 0
                
                # 거래대금 계산 (volume * close 또는 이미 계산된 total_amount 사용)
                total_amount_value = target_stock_data.get('total_amount', 0)
                if total_amount_value is None or total_amount_value == '':
                    # total_amount가 없으면 volume * close로 계산
                    close_price_value = target_stock_data.get('close_price', 0)
                    if close_price_value is None or close_price_value == '':
                        close_price = 0
                    else:
                        try:
                            close_price = float(close_price_value)
                        except (ValueError, TypeError):
                            print(f"⚠️ {stock_code} close_price 값 변환 실패: {close_price_value}")
                            close_price = 0
                    
                    # 거래대금 = 거래대금 * 종가
                    total_amount = volume * close_price if close_price > 0 else 0
                else:
                    # 이미 계산된 total_amount 사용
                    try:
                        total_amount = float(total_amount_value)
                    except (ValueError, TypeError):
                        print(f"⚠️ {stock_code} total_amount 값 변환 실패: {total_amount_value}")
                        total_amount = 0
                
                # 거래율 계산 (모든 차트 타입에서 항상 계산)
                turnover_rate_value = target_stock_data.get('turnover_rate', 0)
                if turnover_rate_value is None or turnover_rate_value == '':
                    turnover_rate = 0
                else:
                    try:
                        turnover_rate = float(turnover_rate_value)
                    except (ValueError, TypeError):
                        print(f"⚠️ {stock_code} turnover_rate 값 변환 실패: {turnover_rate_value}")
                        turnover_rate = 0
                
                # 순위는 이미 계산됨
                
                print(f"💰 {stock_code} {chart_type} 조회:")
                if total_amount > 0:
                    print(f"   거래대금(SUM(volume*close)): {total_amount:,.0f}원")
                else:
                    print(f"   거래대금(SUM(volume*close)): 0원 (데이터 없음)")
                print(f"   거래량: {volume:,.0f}주")
                print(f"   유통주식수: {outstanding_shares:,.0f}주")
                # 모든 차트 타입에서 거래율 표시
                print(f"   거래율: {turnover_rate:.2f}%")
                # 순위 표시 처리
                if ranking == 999:
                    ranking_display = "N/A"
                else:
                    ranking_display = f"{ranking}위"
                print(f"   순위: {ranking_display} (전체 {len(ranking_data)}개 종목 중)")
                
                db.disconnect()
                
                # 주봉의 경우 week_display 사용, 다른 경우는 기존 로직 유지
                if chart_type == "주봉" and 'week_display' in locals():
                    display_date = week_display
                else:
                    display_date = extracted_date if extracted_date != "N/A" else target_date
                
                # 억원 단위로 변환
                amount_in_hundred_millions = total_amount / 100000000
                
                # 모든 차트 타입에서 거래율 반환
                return {
                    "거래일": display_date,
                    "거래대금": f"{amount_in_hundred_millions:,.0f}억원",
                    "거래율": f"{turnover_rate:.2f}%",
                    "순위": ranking_display,
                    "유통주식수": f"{outstanding_shares:,.0f}주",
                    "거래량": f"{volume:,.0f}주"
                }
                
            except Exception as e:
                print(f"❌ 거래 데이터 조회 중 오류: {e}")
                # 기존 방식으로 fallback
                db.disconnect()
                return {
                    "거래일": extracted_date if extracted_date != "N/A" else "N/A",
                    "거래대금": "N/A",
                    "거래율": "N/A",
                    "순위": "N/A",
                    "유통주식수": "N/A",
                    "거래대금": "N/A"
                }
            
        except Exception as e:
            print(f"❌ 거래정보 조회 중 오류: {e}")
            try:
                db.disconnect()
            except:
                pass
            return {
                "거래일": "N/A",
                "거래대금": "N/A",
                "거래율": "N/A",
                "순위": "N/A",
                "유통주식수": "N/A",
                "거래대금": "N/A"
            }
    
    def _add_trading_info_to_result(self, result: Dict[str, Any], stock_code: str, result_type: str, chart_type: str = "일봉", ai_result_text: str = "", trading_type: str = "거래대금") -> Dict[str, Any]:
        """
        결과에 거래일, 거래대금, 거래율, 순위, 유통주식수, 거래량 정보 추가
        
        Args:
            result (Dict[str, Any]): 기존 결과 딕셔너리
            stock_code (str): 종목코드 (6자리)
            result_type (str): 결과 타입 구분값
                - "json파싱성공": AI 응답이 JSON으로 정상 파싱된 경우
                - "파싱실패": AI 응답이 JSON 형식이 아니거나 파싱 실패한 경우 (현재 주로 사용)
                - "fallback": 기본 fallback 결과 생성 시
            chart_type (str): 차트 타입 (일봉, 주봉, 월봉)
            ai_result_text (str): AI 분석 결과 텍스트 (거래일 추출용)
            trading_type (str): 거래 타입 (거래대금, 거래율)
                
        Returns:
            Dict[str, Any]: 거래정보가 추가된 결과 딕셔너리
            
        Note:
            - 종목정보 섹션에 거래일, 거래대금, 거래율, 순위, 유통주식수, 거래량, 타입 필드 추가
            - DB 조회 실패 시 모든 값이 "N/A"로 설정
            - result_type은 디버깅 및 유지보수 목적으로 사용
        """
        try:
            # 거래정보 조회
            print(f"🔍 거래정보 조회 시작: {stock_code} ({chart_type}, {trading_type})")
            trading_info = self._get_trading_info_from_db(stock_code, chart_type, ai_result_text, trading_type)
            print(f"📊 조회된 거래정보: {trading_info}")
            
            # 종목정보 섹션이 있는지 확인
            if "종목정보" in result:
                # 기존 종목정보에 추가
                result["종목정보"]["거래일"] = trading_info["거래일"]
                result["종목정보"]["거래대금"] = trading_info["거래대금"]
                result["종목정보"]["거래율"] = trading_info["거래율"]
                result["종목정보"]["순위"] = trading_info["순위"]
                result["종목정보"]["유통주식수"] = trading_info["유통주식수"]
                result["종목정보"]["거래대금"] = trading_info["거래대금"]
                result["종목정보"]["거래타입"] = trading_type
                result["종목정보"]["타입"] = result_type
                print(f"✅ 거래정보 추가 완료 ({result_type}, {trading_type}): {trading_info}")
            else:
                print(f"⚠️ 종목정보 섹션이 없어 거래정보를 추가할 수 없음")
                
        except Exception as e:
            print(f"❌ 거래정보 추가 중 오류: {e}")
            # 오류 발생 시에도 기본값 추가
            if "종목정보" in result:
                result["종목정보"]["거래일"] = "N/A"
                result["종목정보"]["거래대금"] = "N/A"
                result["종목정보"]["거래율"] = "N/A"
                result["종목정보"]["순위"] = "N/A"
                result["종목정보"]["유통주식수"] = "N/A"
                result["종목정보"]["거래대금"] = "N/A"
                result["종목정보"]["거래타입"] = trading_type
                result["종목정보"]["타입"] = result_type
        
        return result

    def _create_basic_fallback_result(self, stock_name: str, chart_type: str, ai_response: str, error_type: str, stock_code: str = "000000", trading_type: str = "거래대금") -> Dict[str, Any]:
        """기본 fallback 결과 생성"""
        result = {
            "종목정보": {
                "종목명": stock_name,
                "종목번호": stock_code,
                "분석일시": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "차트유형": chart_type,
                "거래타입": trading_type,
                "파싱상태": error_type
            },
            "AI분석결과": ai_response
        }
        
        # 거래정보 추가 (기본 fallback)
        # 기본 fallback 결과 생성 시 실행됨
        # 현재는 거의 사용되지 않지만 향후 확장성을 위해 유지
        # 거래일, 거래대금, 순위 정보를 DB에서 조회하여 추가
        result = self._add_trading_info_to_result(result, stock_code, "fallback", chart_type, ai_response, trading_type)
        
        return result




    def save_analysis_result(self, result: Dict[str, Any], output_path: str) -> bool:
        """
        분석 결과를 JSON 파일로 저장
        
        Args:
            result (Dict[str, Any]): 분석 결과
            output_path (str): 저장할 파일 경로
            
        Returns:
            bool: 저장 성공 여부
        """
        try:
            # 출력 디렉토리 생성
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # JSON 파일로 저장
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(result, f, ensure_ascii=False, indent=2)
            
            print(f"💾 JSON 분석 결과 저장 완료: {output_path}")
            return True
            
        except Exception as e:
            print(f"❌ JSON 결과 저장 중 오류: {e}")
            return False

    def create_consolidated_word_document(self, analysis_results: list, chart_image_paths: list, output_path: str, chart_type: str = "일봉") -> bool:
        """
        여러 종목의 분석 결과를 하나의 통합 Word 문서로 생성
        
        Args:
            analysis_results (list): 분석 결과 리스트 (각 요소는 Dict[str, Any])
            chart_image_paths (list): 차트 이미지 경로 리스트
            output_path (str): 저장할 Word 파일 경로
            chart_type (str): 차트 유형 (일봉/주봉/월봉)
            
        Returns:
            bool: 저장 성공 여부
        """
        try:
            # 출력 디렉토리 생성
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # Word 문서 생성
            doc = Document()
            
            # 한글 폰트 설정을 위한 스타일 설정
            from docx.oxml.ns import qn
            
            # 제목 설정 (통합 리포트)
            title = doc.add_heading(f'주식 차트 분석 통합 리포트 ({chart_type})', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 제목에 한글 폰트 적용
            for run in title.runs:
                run.font.name = '맑은 고딕'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                run.font.size = Pt(14)
            
            # 분석 요약 정보
            heading_summary = doc.add_heading('분석 요약', level=1)
            for run in heading_summary.runs:
                run.font.name = '맑은 고딕'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            
            # 요약 프롬프트로 통합 분석 요약 생성
            summary_result = self._generate_consolidated_summary(analysis_results, chart_type)
            if summary_result:
                p_summary = doc.add_paragraph(summary_result)
                for run in p_summary.runs:
                    run.font.name = '맑은 고딕'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            
            # 각 종목별 상세 분석
            for i, (result, image_path) in enumerate(zip(analysis_results, chart_image_paths)):
                if not result:
                    continue
                
                # 종목 구분선
                if i > 0:
                    doc.add_paragraph("=" * 80)
                
                # 종목 정보
                stock_info = result.get("종목정보", {})
                stock_name = stock_info.get("종목명", f"종목{i+1}")
                stock_code = stock_info.get("종목번호", "000000")
                
                heading_stock = doc.add_heading(f'종목 {i+1}: {stock_name} ({stock_code})', level=1)
                for run in heading_stock.runs:
                    run.font.name = '맑은 고딕'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                
                # 종목별 기본 정보
                p_info1 = doc.add_paragraph(f"종목명: {stock_name}")
                p_info2 = doc.add_paragraph(f"종목번호: {stock_code}")
                p_info3 = doc.add_paragraph(f"분석일시: {stock_info.get('분석일시', 'N/A')}")
                p_info4 = doc.add_paragraph(f"차트유형: {stock_info.get('차트유형', chart_type)}")
                
                for p in [p_info1, p_info2, p_info3, p_info4]:
                    for run in p.runs:
                        run.font.name = '맑은 고딕'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                
                # 차트 이미지 추가
                if os.path.exists(image_path):
                    heading_chart = doc.add_heading('차트 이미지', level=2)
                    for run in heading_chart.runs:
                        run.font.name = '맑은 고딕'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                    
                    doc.add_picture(image_path, width=Inches(6))
                    doc.add_paragraph()
                
                
                # 종합 분석 점수
                if "종합분석점수" in result:
                    heading_score = doc.add_heading('종합 분석 점수', level=2)
                    for run in heading_score.runs:
                        run.font.name = '맑은 고딕'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                    
                    score = result["종합분석점수"]
                    p_score1 = doc.add_paragraph(f"점수: {score.get('점수', 'N/A')}/100")
                    p_score2 = doc.add_paragraph(f"요약: {score.get('요약', 'N/A')}")
                    
                    # 점수 강조
                    p_score1.runs[0].bold = True
                    p_score1.runs[0].font.size = Pt(14)
                    
                    for p in [p_score1, p_score2]:
                        for run in p.runs:
                            run.font.name = '맑은 고딕'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                
                # 핵심 기술적 분석 지표 (간략화)
                if "핵심기술적지표" in result:
                    heading_tech = doc.add_heading('핵심 기술적 지표', level=2)
                    for run in heading_tech.runs:
                        run.font.name = '맑은 고딕'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                    
                    tech = result["핵심기술적지표"]
                    tech_summary = []
                    
                    # 주요 지표들만 요약
                    for key, value in tech.items():
                        if value and value != 'N/A':
                            tech_summary.append(f"{key}: {value}")
                    
                    if tech_summary:
                        p_tech = doc.add_paragraph(" • ".join(tech_summary[:5]))  # 최대 5개
                        for run in p_tech.runs:
                            run.font.name = '맑은 고딕'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                
                # 투자 아이디어 (간략화)
                if chart_type == "일봉" and "단기투자아이디어" in result:
                    heading_idea = doc.add_heading('단기 투자 아이디어', level=2)
                    for run in heading_idea.runs:
                        run.font.name = '맑은 고딕'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                    
                    idea = result["단기투자아이디어"]
                    if "매매시그널" in idea:
                        p_idea = doc.add_paragraph(f"매매 시그널: {idea.get('매매시그널', 'N/A')}")
                        p_idea.runs[0].bold = True
                        for run in p_idea.runs:
                            run.font.name = '맑은 고딕'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                
                elif chart_type == "주봉" and "중기투자아이디어" in result:
                    heading_idea = doc.add_heading('중기 투자 아이디어', level=2)
                    for run in heading_idea.runs:
                        run.font.name = '맑은 고딕'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                    
                    idea = result["중기투자아이디어"]
                    if "매매시그널" in idea:
                        p_idea = doc.add_paragraph(f"매매 시그널: {idea.get('매매시그널', 'N/A')}")
                        p_idea.runs[0].bold = True
                        for run in p_idea.runs:
                            run.font.name = '맑은 고딕'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                
                elif chart_type == "월봉" and "장기투자아이디어" in result:
                    heading_idea = doc.add_heading('장기 투자 아이디어', level=2)
                    for run in heading_idea.runs:
                        run.font.name = '맑은 고딕'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                    
                    idea = result["장기투자아이디어"]
                    if "투자전략" in idea:
                        p_idea = doc.add_paragraph(f"투자 전략: {idea.get('투자전략', 'N/A')}")
                        p_idea.runs[0].bold = True
                        for run in p_idea.runs:
                            run.font.name = '맑은 고딕'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            
            # 문서 저장
            doc.save(output_path)
            print(f"📄 통합 Word 문서 저장 완료: {output_path}")
            return True
            
        except Exception as e:
            print(f"❌ 통합 Word 문서 생성 중 오류: {e}")
            return False

    def _generate_consolidated_summary(self, analysis_results: list, chart_type: str) -> str:
        """
        여러 종목의 분석 결과를 요약 프롬프트로 통합 요약 생성
        
        Args:
            analysis_results (list): 분석 결과 리스트
            chart_type (str): 차트 유형
            
        Returns:
            str: 통합 요약 텍스트
        """
        try:
            if not analysis_results:
                return "분석 결과가 없습니다."
            
            # 요약용 데이터 준비
            summary_data = {
                "분석종목수": len(analysis_results),
                "차트유형": chart_type,
                "분석일시": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "종목별요약": []
            }
            
            # 각 종목별 핵심 정보 추출
            for result in analysis_results:
                if not result:
                    continue
                
                stock_info = result.get("종목정보", {})
                stock_name = stock_info.get("종목명", "알 수 없음")
                stock_code = stock_info.get("종목번호", "000000")
                summary_item = {
                    "종목명": stock_name,
                    "종목번호": stock_code
                }
                
                # 종합 분석 점수
                if "종합분석점수" in result:
                    score = result["종합분석점수"]
                    summary_item["점수"] = score.get("점수", "N/A")
                    summary_item["요약"] = score.get("요약", "N/A")
                
                # 핵심 기술적 지표
                if "핵심기술적지표" in result:
                    tech = result["핵심기술적지표"]
                    summary_item["기술적지표"] = tech
                
                # 투자 아이디어
                if chart_type == "일봉" and "단기투자아이디어" in result:
                    idea = result["단기투자아이디어"]
                    summary_item["투자아이디어"] = idea.get("매매시그널", "N/A")
                elif chart_type == "주봉" and "중기투자아이디어" in result:
                    idea = result["중기투자아이디어"]
                    summary_item["투자아이디어"] = idea.get("매매시그널", "N/A")
                elif chart_type == "월봉" and "장기투자아이디어" in result:
                    idea = result["장기투자아이디어"]
                    summary_item["투자아이디어"] = idea.get("투자전략", "N/A")
                
                summary_data["종목별요약"].append(summary_item)
            
            # 요약 프롬프트 생성
            summary_prompt = self.prompts.get_summary_prompt(chart_type)
            summary_prompt += f"\n\n다음 {len(analysis_results)}개 종목의 {chart_type} 차트 분석 결과를 요약해주세요:\n\n"
            
            # 요약 데이터를 JSON 형태로 추가
            summary_prompt += f"분석 데이터:\n{json.dumps(summary_data, ensure_ascii=False, indent=2)}\n\n"
            summary_prompt += "위 데이터를 바탕으로 전체적인 시장 동향과 주요 투자 포인트를 요약해주세요."
            
            # AI 요약 생성
            try:
                response = self.model.generate_content(summary_prompt)
                if response.text:
                    return response.text
                else:
                    return f"{len(analysis_results)}개 종목의 {chart_type} 차트 분석이 완료되었습니다."
            except Exception as e:
                print(f"⚠️ AI 요약 생성 실패: {e}")
                return f"{len(analysis_results)}개 종목의 {chart_type} 차트 분석이 완료되었습니다."
                
        except Exception as e:
            print(f"❌ 통합 요약 생성 중 오류: {e}")
            return f"{len(analysis_results)}개 종목의 {chart_type} 차트 분석이 완료되었습니다."


    def _load_additional_data_files(self, json_data_path: str, csv_data_path: str, text_summary_path: str) -> str:
        """
        추가 데이터 파일들을 로드하고 프롬프트용 텍스트로 변환
        
        Args:
            json_data_path (str): JSON 데이터 파일 경로
            csv_data_path (str): CSV 데이터 파일 경로
            text_summary_path (str): 텍스트 요약 파일 경로
            
        Returns:
            str: 프롬프트에 추가할 데이터 정보 텍스트
        """
        additional_info = ""
        
        # 1. JSON 데이터 파일 로드
        if json_data_path and os.path.exists(json_data_path):
            try:
                print(f"📊 JSON 데이터 파일 로드 중: {json_data_path}")
                with open(json_data_path, 'r', encoding='utf-8') as f:
                    json_data = json.load(f)
                
                # JSON 데이터를 구조화된 텍스트로 변환
                json_info = f"""
**JSON 구조화 데이터 정보:**
- 종목명: {json_data.get('metadata', {}).get('stock_name', 'N/A')}
- 종목코드: {json_data.get('metadata', {}).get('stock_code', 'N/A')}
- 데이터 기간: {json_data.get('metadata', {}).get('data_period', {}).get('start', 'N/A')} ~ {json_data.get('metadata', {}).get('data_period', {}).get('end', 'N/A')}
- 총 데이터 수: {json_data.get('metadata', {}).get('total_records', 'N/A')}개

**요약 정보:**
- 최근 종가: {json_data.get('summary', {}).get('latest_close', 'N/A'):,.0f}원
- 최근 거래량: {json_data.get('summary', {}).get('latest_volume', 'N/A'):,}주
- 가격 변동: {json_data.get('summary', {}).get('price_change', 'N/A'):+,.0f}원
- 변동률: {json_data.get('summary', {}).get('price_change_pct', 'N/A'):+.2f}%
- 최고가: {json_data.get('summary', {}).get('highest_price', 'N/A'):,.0f}원
- 최저가: {json_data.get('summary', {}).get('lowest_price', 'N/A'):,.0f}원
- 평균 거래량: {json_data.get('summary', {}).get('avg_volume', 'N/A'):,.0f}주

**기술적 지표 (최근값):**
"""
                
                # 기술적 지표 정보 추가
                tech_indicators = json_data.get('technical_indicators', {}).get('latest_values', {})
                for indicator, value in tech_indicators.items():
                    if value is not None:
                        if 'ma' in indicator.lower():
                            json_info += f"- {indicator.upper()}: {value:,.0f}원\n"
                        else:
                            json_info += f"- {indicator.upper()}: {value:.2f}\n"
                
                # 최근 차트 데이터 (최대 5개)
                chart_data = json_data.get('chart_data', [])
                if chart_data:
                    json_info += f"\n**최근 5개 거래일 데이터:**\n"
                    for i, data_point in enumerate(chart_data[-5:]):
                        json_info += f"- {data_point['date']}: 시가 {data_point['open']:,.0f}, 고가 {data_point['high']:,.0f}, 저가 {data_point['low']:,.0f}, 종가 {data_point['close']:,.0f}, 거래대금 {data_point['volume']:,}\n"
                
                additional_info += json_info
                print(f"✅ JSON 데이터 로드 완료")
                
            except Exception as e:
                print(f"❌ JSON 데이터 파일 로드 실패: {e}")
        
        # 2. CSV 데이터 파일 로드
        if csv_data_path and os.path.exists(csv_data_path):
            try:
                print(f"📊 CSV 데이터 파일 로드 중: {csv_data_path}")
                import pandas as pd
                csv_data = pd.read_csv(csv_data_path, encoding='utf-8-sig')
                
                csv_info = f"""
**CSV 데이터 정보:**
- 파일 경로: {csv_data_path}
- 데이터 수: {len(csv_data)}개
- 컬럼: {', '.join(csv_data.columns.tolist())}

**최근 5개 데이터:**
"""
                
                # 최근 5개 데이터 추가
                for i, row in csv_data.tail(5).iterrows():
                    csv_info += f"- {row.iloc[0]}: 시가 {row['Open']:,.0f}, 고가 {row['High']:,.0f}, 저가 {row['Low']:,.0f}, 종가 {row['Close']:,.0f}, 거래대금 {row['Volume']:,}\n"
                
                additional_info += csv_info
                print(f"✅ CSV 데이터 로드 완료")
                
            except Exception as e:
                print(f"❌ CSV 데이터 파일 로드 실패: {e}")
        
        # 3. 텍스트 요약 파일 로드
        if text_summary_path and os.path.exists(text_summary_path):
            try:
                print(f"📊 텍스트 요약 파일 로드 중: {text_summary_path}")
                with open(text_summary_path, 'r', encoding='utf-8') as f:
                    text_content = f.read()
                
                text_info = f"""
**텍스트 요약 정보:**
{text_content}
"""
                
                additional_info += text_info
                print(f"✅ 텍스트 요약 로드 완료")
                
            except Exception as e:
                print(f"❌ 텍스트 요약 파일 로드 실패: {e}")
        
        return additional_info

    def analyze_text_with_prompt(self, prompt: str) -> dict:
        """
        텍스트 기반 프롬프트 분석 (월봉 요약용)
        
        Args:
            prompt (str): 분석할 프롬프트
            
        Returns:
            dict: AI 분석 결과
        """
        try:
            print("🤖 텍스트 기반 AI 분석 시작...")
            
            # Gemini 모델 초기화
            if not self.model:
                print("❌ Gemini 모델이 초기화되지 않았습니다")
                return None
            
            # 프롬프트 전송 및 응답 받기
            response = self.model.generate_content(prompt)
            
            if response and response.text:
                print("✅ AI 분석 응답 수신 완료")
                
                # 응답 파싱
                analysis_result = self._parse_text_analysis_response(response.text)
                return analysis_result
            else:
                print("❌ AI 분석 응답이 비어있습니다")
                return None
                
        except Exception as e:
            print(f"❌ 텍스트 기반 AI 분석 중 오류: {e}")
            return None
    
    def _parse_text_analysis_response(self, response_text: str) -> dict:
        """
        텍스트 기반 AI 분석 응답 파싱
        
        Args:
            response_text (str): AI 응답 텍스트
            
        Returns:
            dict: 파싱된 분석 결과
        """
        try:
            # JSON 코드 블록 제거 (```json ... ``` 형태)
            cleaned_text = response_text.strip()
            
            # JSON 코드 블록 제거
            if cleaned_text.startswith('```json'):
                cleaned_text = cleaned_text[7:]
            elif cleaned_text.startswith('```'):
                cleaned_text = cleaned_text[3:]
            
            if cleaned_text.endswith('```'):
                cleaned_text = cleaned_text[:-3]
            
            cleaned_text = cleaned_text.strip()
            
            # 간단한 파싱 (JSON 형식이 아닌 경우를 대비)
            result = {
                "분석_결과": cleaned_text,  # 코드 블록이 제거된 텍스트 사용
                "생성_시간": datetime.now().isoformat(),
                "분석_유형": "월봉_요약_AI분석"
            }
            
            # 응답에서 주요 섹션 추출 시도
            lines = response_text.split('\n')
            current_section = ""
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                    
                # 섹션 헤더 감지
                if line.startswith('📈') or line.startswith('📉') or line.startswith('🌍') or line.startswith('💼'):
                    current_section = line
                    result[current_section] = []
                elif current_section and line.startswith('•'):
                    if current_section not in result:
                        result[current_section] = []
                    result[current_section].append(line[1:].strip())
            
            return result
            
        except Exception as e:
            print(f"⚠️ 응답 파싱 중 오류: {e}")
            # 기본 결과 반환
            return {
                "분석_결과": response_text,
                "생성_시간": datetime.now().isoformat(),
                "분석_유형": "월봉_요약_AI분석",
                "파싱_오류": str(e)
            }

    def analyze_chart_with_data_files(self, image_path: str, json_data_path: str = "", csv_data_path: str = "", 
                                     text_summary_path: str = "", stock_name: str = "", chart_type: str = "일봉",
                                     enable_summary_analysis: bool = False) -> Optional[Dict[str, Any]]:
        """
        차트 이미지와 데이터 파일들을 함께 AI로 분석하는 편의 메서드 (하이브리드 지원)
        
        Args:
            image_path (str): 차트 이미지 파일 경로
            json_data_path (str): JSON 데이터 파일 경로
            csv_data_path (str): CSV 데이터 파일 경로
            text_summary_path (str): 텍스트 요약 파일 경로
            stock_name (str): 종목명
            chart_type (str): 차트 유형 (일봉/주봉/월봉)
            enable_summary_analysis (bool): 요약 분석 활성화 여부
            
        Returns:
            Dict[str, Any]: 분석 결과 JSON (요약 분석 포함 시 summary_analysis 키 추가)
        """
        print(f"🚀 차트 이미지와 데이터 파일들을 함께 분석합니다...")
        print(f"📈 차트 이미지: {image_path}")
        print(f"📊 JSON 데이터: {json_data_path if json_data_path else '없음'}")
        print(f"📋 CSV 데이터: {csv_data_path if csv_data_path else '없음'}")
        print(f"📝 텍스트 요약: {text_summary_path if text_summary_path else '없음'}")
        
        return self.analyze_chart_image(
            image_path=image_path,
            stock_name=stock_name,
            chart_type=chart_type,
            json_data_path=json_data_path,
            csv_data_path=csv_data_path,
            text_summary_path=text_summary_path,
            enable_summary_analysis=enable_summary_analysis
        )

    def find_related_data_files(self, image_path: str) -> tuple:
        """
        차트 이미지 파일과 관련된 데이터 파일들을 자동으로 찾기
        
        Args:
            image_path (str): 차트 이미지 파일 경로
            
        Returns:
            tuple: (json_path, csv_path, text_path)
        """
        print(f"🔍 관련 데이터 파일들을 찾는 중: {image_path}")
        
        # 이미지 파일명에서 기본 정보 추출
        image_filename = os.path.basename(image_path)
        image_name_without_ext = os.path.splitext(image_filename)[0]
        
        # 파일명에서 종목명과 종목코드 추출
        parts = image_name_without_ext.split('_')
        if len(parts) >= 3:
            chart_type = parts[0]  # daily, weekly, monthly
            stock_name = parts[1]
            stock_code = parts[2]
            date_part = parts[3] if len(parts) > 3 else ""
        else:
            print(f"⚠️ 이미지 파일명 형식을 인식할 수 없습니다: {image_filename}")
            return "", "", ""
        
        # 관련 파일들 찾기
        json_path = ""
        csv_path = ""
        text_path = ""
        
        # 1. JSON 파일 찾기
        json_pattern = f"{chart_type}_{stock_name}_{stock_code}_{date_part}.json"
        json_dir = "chart_data_json"
        if os.path.exists(json_dir):
            for file in os.listdir(json_dir):
                if file.startswith(f"{chart_type}_{stock_name}_{stock_code}_{date_part}"):
                    json_path = os.path.join(json_dir, file)
                    break
        
        # 2. CSV 파일 찾기
        csv_pattern = f"{chart_type}_{stock_name}_{stock_code}_{date_part}.csv"
        csv_dir = "chart_data_csv"
        if os.path.exists(csv_dir):
            for file in os.listdir(csv_dir):
                if file.startswith(f"{chart_type}_{stock_name}_{stock_code}_{date_part}"):
                    csv_path = os.path.join(csv_dir, file)
                    break
        
        # 3. 텍스트 요약 파일 찾기
        text_pattern = f"{chart_type}_{stock_name}_{stock_code}_{date_part}_summary.txt"
        text_dir = "chart_data_text"
        if os.path.exists(text_dir):
            for file in os.listdir(text_dir):
                if file.startswith(f"{chart_type}_{stock_name}_{stock_code}_{date_part}_summary"):
                    text_path = os.path.join(text_dir, file)
                    break
        
        print(f"📊 찾은 관련 파일들:")
        print(f"   JSON: {json_path if json_path else '없음'}")
        print(f"   CSV: {csv_path if csv_path else '없음'}")
        print(f"   텍스트: {text_path if text_path else '없음'}")
        
        return json_path, csv_path, text_path

def main():
    """메인 함수"""
    print("�� AI 제미나이 차트 분석 프로그램 (개선된 버전)")
    print("="*60)
    
    # 데이터베이스 설정 로드
    try:
        from config import config
        db_config = config.get_database_config()
        print("✅ 데이터베이스 설정 로드 완료")
    except Exception as e:
        print(f"⚠️ 데이터베이스 설정 로드 실패: {e}")
        db_config = None
    
    # AI 분석기 초기화 (DB 기반)
    try:
        analyzer = AIChartAnalyzer(api_key=None, db_config=db_config)
        print("✅ AI 분석기 초기화 완료")
    except Exception as e:
        print(f"❌ AI 분석기 초기화 실패: {e}")
        return
    
    # 차트 폴더들 확인
    chart_folders = ["daily_charts", "weekly_charts", "monthly_charts"]
    available_folders = []
    
    for folder in chart_folders:
        if os.path.exists(folder):
            chart_files = [f for f in os.listdir(folder) if f.endswith('.png')]
            if chart_files:
                available_folders.append((folder, chart_files))
    
    if not available_folders:
        print("❌ 차트 이미지가 있는 폴더를 찾을 수 없습니다.")
        print("먼저 차트 생성 프로그램을 실행하여 차트를 생성해주세요.")
        return
    
    print("📁 발견된 차트 파일들:")
    file_index = 1
    file_mapping = {}
    
    for folder, files in available_folders:
        chart_type = folder.replace("_charts", "")
        print(f"\n📊 {chart_type} 차트:")
        for file in files:
            print(f"  {file_index}. {file}")
            file_mapping[file_index] = (folder, file, chart_type)
            file_index += 1
    
    # 분석할 파일 선택 (여러 파일 선택 가능)
    print(f"\n📊 분석할 차트를 선택하세요:")
    print(f"1. 단일 파일 분석")
    print(f"2. 여러 파일 통합 분석 (권장)")
    
    while True:
        try:
            analysis_mode = input("분석 모드 선택 (1 또는 2): ").strip()
            if analysis_mode in ['1', '2']:
                break
            else:
                print("❌ 1 또는 2를 입력해주세요.")
        except ValueError:
            print("❌ 숫자를 입력해주세요.")
    
    if analysis_mode == '1':
        # 단일 파일 분석 (기존 방식 - 주석 처리됨)
        print(f"\n⚠️ 단일 파일 분석은 현재 주석 처리되어 있습니다.")
        print(f"   통합 분석 모드를 사용하세요.")
        return
        
        # 주석 처리된 기존 단일 파일 분석 로직
        """
        while True:
            try:
                choice = input(f"\n📊 분석할 차트 번호를 선택하세요 (1-{len(file_mapping)}): ").strip()
                file_index = int(choice)
                
                if file_index in file_mapping:
                    folder, selected_file, chart_type = file_mapping[file_index]
                    break
                else:
                    print("❌ 올바른 번호를 입력해주세요.")
            except ValueError:
                print("❌ 숫자를 입력해주세요.")
        """
        
    else:
        # 통합 분석 모드
        print(f"\n🚀 통합 분석 모드를 시작합니다...")
        
        # 차트 유형별로 파일 그룹화
        chart_type_groups = {}
        for folder, files in available_folders:
            chart_type = folder.replace("_charts", "")
            if chart_type not in chart_type_groups:
                chart_type_groups[chart_type] = []
            
            for file in files:
                file_path = os.path.join(folder, file)
                chart_type_groups[chart_type].append(file_path)
        
        # 차트 유형 선택
        print(f"\n📊 분석할 차트 유형을 선택하세요:")
        chart_types = list(chart_type_groups.keys())
        for i, chart_type in enumerate(chart_types):
            file_count = len(chart_type_groups[chart_type])
            print(f"  {i+1}. {chart_type} 차트 ({file_count}개 파일)")
        
        while True:
            try:
                type_choice = input(f"차트 유형 선택 (1-{len(chart_types)}): ").strip()
                type_index = int(type_choice) - 1
                
                if 0 <= type_index < len(chart_types):
                    selected_chart_type = chart_types[type_index]
                    break
                else:
                    print("❌ 올바른 번호를 입력해주세요.")
            except ValueError:
                print("❌ 숫자를 입력해주세요.")
        
        selected_files = chart_type_groups[selected_chart_type]
        print(f"\n📊 {selected_chart_type} 차트 {len(selected_files)}개 파일을 분석합니다...")
        
        # 모든 파일 분석
        analysis_results = []
        chart_image_paths = []
        
        for i, file_path in enumerate(selected_files):
            print(f"\n🔍 분석 진행 중... ({i+1}/{len(selected_files)})")
            print(f"📁 파일: {os.path.basename(file_path)}")
            
            # 관련 데이터 파일들 자동 찾기
            json_path, csv_path, text_path = analyzer.find_related_data_files(file_path)
            
            # AI 분석 실행
            if json_path or csv_path or text_path:
                print(f"✅ 관련 데이터 파일들과 함께 분석합니다.")
                result = analyzer.analyze_chart_with_data_files(
                    image_path=file_path,
                    json_data_path=json_path,
                    csv_data_path=csv_path,
                    text_summary_path=text_path,
                    stock_name="",
                    chart_type=selected_chart_type
                )
            else:
                print(f"⚠️ 이미지만으로 분석합니다.")
                result = analyzer.analyze_chart_image(file_path, "", selected_chart_type)
            
            if result:
                # 개별 분석 결과에 JSON 파일 경로 추가
                stock_info = result.get("종목정보", {})
                stock_name = stock_info.get("종목명", "unknown")
                stock_code = stock_info.get("종목번호", "000000")
                
                # 개별 분석 JSON 파일 경로 생성
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                individual_json_filename = f"individual_analysis_{selected_chart_type}_{stock_name}_{stock_code}_{timestamp}.json"
                individual_json_path = os.path.join(output_dir, individual_json_filename)
                
                # 개별 분석 결과에 JSON 파일 경로 추가
                result["individual_analysis_file"] = individual_json_path
                
                analysis_results.append(result)
                chart_image_paths.append(file_path)
                print(f"✅ 분석 완료: {os.path.basename(file_path)}")
            else:
                print(f"❌ 분석 실패: {os.path.basename(file_path)}")
        
        if analysis_results:
            # 통합 문서 생성
            output_dir = "ai_analysis_results"
            if not os.path.exists(output_dir):
                os.makedirs(output_dir)
                print(f"📁 {output_dir} 폴더를 생성했습니다.")
            
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            
            # 통합 Word 문서 저장
            doc_filename = f"consolidated_analysis_{selected_chart_type}_{timestamp}.docx"
            doc_path = os.path.join(output_dir, doc_filename)
            
            # 개별 분석 JSON 파일들 저장
            print(f"💾 개별 분석 JSON 파일들 저장 중...")
            for result in analysis_results:
                individual_json_path = result.get("individual_analysis_file", "")
                if individual_json_path:
                    json_success = analyzer.save_analysis_result(result, individual_json_path)
                    if json_success:
                        print(f"✅ 개별 분석 JSON 저장: {os.path.basename(individual_json_path)}")
                    else:
                        print(f"❌ 개별 분석 JSON 저장 실패: {os.path.basename(individual_json_path)}")
            
            # 통합 Word 문서 생성
            doc_success = analyzer.create_consolidated_word_document(
                analysis_results, chart_image_paths, doc_path, selected_chart_type
            )
            
            if doc_success:
                print(f"\n✅ AI 차트 통합 분석이 완료되었습니다!")
                print(f"📄 통합 Word 문서 파일: {doc_path}")
                print(f"📊 분석 완료된 종목 수: {len(analysis_results)}개")
                
                # 주요 결과 요약
                print(f"\n📈 주요 분석 결과:")
                for i, result in enumerate(analysis_results):
                    stock_info = result.get("종목정보", {})
                    stock_name = stock_info.get("종목명", f"종목{i+1}")
                    stock_code = stock_info.get("종목번호", "000000")
                    
                    if "종합분석점수" in result:
                        score = result["종합분석점수"]
                        print(f"  {stock_name} ({stock_code}): {score.get('점수', 'N/A')}/100")
                    else:
                        print(f"  {stock_name} ({stock_code}): 분석 완료")
            else:
                print("❌ 통합 Word 문서 생성에 실패했습니다.")
        else:
            print("❌ 모든 파일 분석에 실패했습니다.")
        
        return
    
    # 주석 처리된 기존 단일 파일 분석 로직
    """
    # 파일 경로 설정
    image_path = os.path.join(folder, selected_file)
    
    print(f"\n🔍 분석 시작: {selected_file}")
    print(f"📁 파일: {image_path}")
    print(f"📊 차트 유형: {chart_type}")
    
    # 관련 데이터 파일들 자동 찾기
    print(f"\n🔍 관련 데이터 파일들을 찾는 중...")
    json_path, csv_path, text_path = analyzer.find_related_data_files(image_path)
    
    # 분석 모드 선택
    print(f"\n📊 분석 모드를 선택하세요:")
    print(f"1. 이미지만으로 분석 (기본)")
    print(f"2. 이미지 + 데이터 파일들과 함께 분석 (권장)")
    
    while True:
        try:
            mode_choice = input("선택 (1 또는 2): ").strip()
            if mode_choice in ['1', '2']:
                break
            else:
                print("❌ 1 또는 2를 입력해주세요.")
        except ValueError:
            print("❌ 숫자를 입력해주세요.")
    
    # AI 분석 실행
    if mode_choice == '1':
        print(f"\n📊 이미지만으로 분석을 진행합니다...")
        result = analyzer.analyze_chart_image(image_path, "", chart_type)
    else:
        print(f"\n📊 이미지와 데이터 파일들을 함께 분석합니다...")
        if json_path or csv_path or text_path:
            print(f"✅ 관련 데이터 파일들을 찾았습니다!")
            result = analyzer.analyze_chart_with_data_files(
                image_path=image_path,
                json_data_path=json_path,
                csv_data_path=csv_path,
                text_summary_path=text_path,
                stock_name="",
                chart_type=chart_type
            )
        else:
            print(f"⚠️ 관련 데이터 파일을 찾을 수 없어 이미지만으로 분석합니다.")
            result = analyzer.analyze_chart_image(image_path, "", chart_type)
    
    if result:
        # 결과 저장
        output_dir = "ai_analysis_results"
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
            print(f"📁 {output_dir} 폴더를 생성했습니다.")
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        # 종목정보 추출
        stock_info = result.get("종목정보", {})
        stock_name = stock_info.get("종목명", "unknown")
        stock_code = stock_info.get("종목번호", "000000")
        
        # Word 문서 저장
        doc_filename = f"analysis_{chart_type}_{stock_name}_{stock_code}_{timestamp}.docx"
        doc_path = os.path.join(output_dir, doc_filename)
        
        # JSON 파일 저장 (개별 분석 결과)
        json_filename = f"individual_analysis_{chart_type}_{stock_name}_{stock_code}_{timestamp}.json"
        json_path = os.path.join(output_dir, json_filename)
        json_success = analyzer.save_analysis_result(result, json_path)
        
        # Word 문서 생성 (하이브리드 방식)
        doc_success = analyzer.create_word_document_hybrid(result, image_path, doc_path, chart_type)
        
        if doc_success and json_success:
            print("\n✅ AI 차트 분석이 완료되었습니다!")
            print(f"📄 Word 문서 파일: {doc_path}")
            print(f"📊 JSON 분석 결과: {json_path}")
            
            # 주요 결과 출력
            if "종합분석점수" in result:
                score = result["종합분석점수"]
                print(f"\n📊 종합 분석 점수: {score.get('점수', 'N/A')}/100")
                print(f"📝 요약: {score.get('요약', 'N/A')}")
            
            # 투자 아이디어 출력
            if chart_type == "일봉" and "단기투자아이디어" in result:
                idea = result["단기투자아이디어"]
                print(f"\n📈 단기 투자 아이디어:")
                print(f"   추세 요약: {idea.get('추세요약', 'N/A')}")
                print(f"   매매 시그널: {idea.get('매매시그널', 'N/A')}")
            elif chart_type == "주봉" and "중기투자아이디어" in result:
                idea = result["중기투자아이디어"]
                print(f"\n📈 중기 투자 아이디어:")
                print(f"   추세 요약: {idea.get('추세요약', 'N/A')}")
                print(f"   매매 시그널: {idea.get('매매시그널', 'N/A')}")
            elif chart_type == "월봉" and "장기투자아이디어" in result:
                idea = result["장기투자아이디어"]
                print(f"\n📈 장기 투자 아이디어:")
                print(f"   사이클 요약: {idea.get('사이클요약', 'N/A')}")
                print(f"   투자 전략: {idea.get('투자전략', 'N/A')}")
        else:
            if not doc_success:
                print("❌ Word 문서 생성에 실패했습니다.")
    else:
        print("❌ AI 분석에 실패했습니다.")
    """
    
    # 주석 처리된 기존 단일 파일 분석 로직 완료
    print("⚠️ 기존 단일 파일 분석 로직은 주석 처리되었습니다.")
    print("   통합 분석 모드를 사용하세요.")

def analyze_single_chart_with_data(image_path: str, json_data_path: str = "", csv_data_path: str = "", 
                                  text_summary_path: str = "", chart_type: str = "일봉"):
    """
    단일 차트를 데이터 파일들과 함께 분석하는 편의 함수
    
    Args:
        image_path (str): 차트 이미지 파일 경로
        json_data_path (str): JSON 데이터 파일 경로
        csv_data_path (str): CSV 데이터 파일 경로
        text_summary_path (str): 텍스트 요약 파일 경로
        chart_type (str): 차트 유형 (일봉/주봉/월봉)
    """
    print("🤖 AI 제미나이 차트 분석 프로그램 (DB 기반 단일 파일 분석)")
    print("="*60)
    
    # 데이터베이스 설정 로드
    try:
        from config import config
        db_config = config.get_database_config()
        print("✅ 데이터베이스 설정 로드 완료")
    except Exception as e:
        print(f"⚠️ 데이터베이스 설정 로드 실패: {e}")
        db_config = None
    
    # AI 분석기 초기화 (DB 기반)
    try:
        analyzer = AIChartAnalyzer(api_key=None, db_config=db_config)
        print("✅ AI 분석기 초기화 완료")
    except Exception as e:
        print(f"❌ AI 분석기 초기화 실패: {e}")
        return None
    
    # 파일 존재 확인
    if not os.path.exists(image_path):
        print(f"❌ 차트 이미지 파일을 찾을 수 없습니다: {image_path}")
        return None
    
    print(f"🔍 분석 시작: {os.path.basename(image_path)}")
    print(f"📁 파일: {image_path}")
    print(f"📊 차트 유형: {chart_type}")
    
    # AI 분석 실행
    result = analyzer.analyze_chart_with_data_files(
        image_path=image_path,
        json_data_path=json_data_path,
        csv_data_path=csv_data_path,
        text_summary_path=text_summary_path,
        stock_name="",
        chart_type=chart_type
    )
    
    if result:
        # 결과 저장
        output_dir = "ai_analysis_results"
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        # 종목정보 추출
        stock_info = result.get("종목정보", {})
        stock_name = stock_info.get("종목명", "unknown")
        stock_code = stock_info.get("종목번호", "000000")
        
        # Word 문서 저장
        doc_filename = f"analysis_{chart_type}_{stock_name}_{stock_code}_{timestamp}.docx"
        doc_path = os.path.join(output_dir, doc_filename)
        
        # Word 문서 생성 (하이브리드 방식)
        doc_success = analyzer.create_word_document_hybrid(result, image_path, doc_path, chart_type)
        
        if doc_success:
            print("\n✅ AI 차트 분석이 완료되었습니다!")
            print(f"📄 Word 문서 파일: {doc_path}")
            return result
        else:
            print("❌ Word 문서 생성에 실패했습니다.")
            return None
    else:
        print("❌ AI 분석에 실패했습니다.")
        return None



class SummaryFileGenerator:
    """통합 요약 파일 생성기"""
    
    def __init__(self, db_config: dict = None):
        """
        통합 요약 파일 생성기 초기화
        
        Args:
            db_config (dict): 데이터베이스 설정
        """
        self.db_config = db_config
        
        # 종목명 매퍼 초기화 (DB 기반)
        try:
            self.stock_mapper = StockNameMapper(db_config)
            print("✅ 종목명 매퍼 초기화 완료")
        except Exception as e:
            print(f"⚠️ 종목명 매퍼 초기화 실패: {e}")
            self.stock_mapper = None
        
        # 프롬프트 관리자 초기화
        if db_config:
            try:
                self.prompts = ChartAnalysisPrompts(db_config)
                print("✅ 요약 파일 생성기 초기화 완료")
            except Exception as e:
                print(f"⚠️ 프롬프트 관리자 초기화 실패: {e}")
                self.prompts = ChartAnalysisPrompts(None)
        else:
            self.prompts = ChartAnalysisPrompts(None)
        
        # AI 분석기 초기화 (요약 분석용)
        try:
            self.analyzer = AIChartAnalyzer(api_key=None, db_config=db_config)
            print("✅ AI 분석기 초기화 완료")
        except Exception as e:
            print(f"⚠️ AI 분석기 초기화 실패: {e}")
            self.analyzer = None
    
    def scan_analysis_results(self, results_dir: str = "ai_analysis_results") -> dict:
        """
        분석 결과 폴더를 스캔하여 차트 유형별로 그룹화
        
        Args:
            results_dir (str): 분석 결과 폴더 경로
            
        Returns:
            dict: 차트 유형별로 그룹화된 파일 정보
        """
        try:
            print(f"📁 분석 결과 폴더 스캔 중: {results_dir}")
            
            if not os.path.exists(results_dir):
                print(f"❌ 분석 결과 폴더가 존재하지 않습니다: {results_dir}")
                return {}
            
            grouped_files = {
                "daily": [],    # 일봉
                "weekly": [],   # 주봉  
                "monthly": []   # 월봉
            }
            
            # JSON 파일들만 스캔
            json_files = [f for f in os.listdir(results_dir) if f.endswith('.json')]
            print(f"📊 발견된 JSON 파일 수: {len(json_files)}개")
            
            # 종목별 최신 파일만 선택하기 위한 딕셔너리
            latest_files = {
                "daily": {},
                "weekly": {},
                "monthly": {}
            }
            
            for json_file in json_files:
                file_path = os.path.join(results_dir, json_file)
                
                # 파일명에서 차트 유형과 종목코드 추출
                if json_file.startswith('analysis_daily_'):
                    # analysis_daily_종목코드_날짜시간.json 형태에서 종목코드 추출
                    parts = json_file.split('_')
                    if len(parts) >= 3:
                        stock_code = parts[2]
                        # 같은 종목코드의 파일이 이미 있다면 더 최신 파일로 교체
                        if stock_code not in latest_files["daily"] or json_file > latest_files["daily"][stock_code]["filename"]:
                            latest_files["daily"][stock_code] = {"path": file_path, "filename": json_file}
                elif json_file.startswith('analysis_weekly_'):
                    parts = json_file.split('_')
                    if len(parts) >= 3:
                        stock_code = parts[2]
                        if stock_code not in latest_files["weekly"] or json_file > latest_files["weekly"][stock_code]["filename"]:
                            latest_files["weekly"][stock_code] = {"path": file_path, "filename": json_file}
                elif json_file.startswith('analysis_monthly_'):
                    parts = json_file.split('_')
                    if len(parts) >= 3:
                        stock_code = parts[2]
                        if stock_code not in latest_files["monthly"] or json_file > latest_files["monthly"][stock_code]["filename"]:
                            latest_files["monthly"][stock_code] = {"path": file_path, "filename": json_file}
            
            # 최신 파일들만 grouped_files에 추가
            for chart_type in ["daily", "weekly", "monthly"]:
                for stock_code, file_info in latest_files[chart_type].items():
                    grouped_files[chart_type].append(file_info["path"])
            
            print(f"📈 차트 유형별 파일 수:")
            print(f"   - 일봉: {len(grouped_files['daily'])}개")
            print(f"   - 주봉: {len(grouped_files['weekly'])}개") 
            print(f"   - 월봉: {len(grouped_files['monthly'])}개")
            
            return grouped_files
            
        except Exception as e:
            print(f"❌ 분석 결과 스캔 중 오류: {e}")
            return {}
    
    def load_analysis_results(self, file_paths: list) -> list:
        """
        분석 결과 JSON 파일들을 로드
        
        Args:
            file_paths (list): JSON 파일 경로 리스트
            
        Returns:
            list: 로드된 분석 결과 리스트
        """
        try:
            print(f"📄 분석 결과 로딩 중: {len(file_paths)}개 파일")
            
            loaded_results = []
            
            for file_path in file_paths:
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        analysis_data = json.load(f)
                        
                    # 파일 정보 추가
                    analysis_data["_file_info"] = {
                        "file_path": file_path,
                        "file_name": os.path.basename(file_path),
                        "loaded_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    }
                    
                    loaded_results.append(analysis_data)
                    
                except Exception as e:
                    print(f"⚠️ 파일 로드 실패: {file_path} - {e}")
                    continue
            
            print(f"✅ 분석 결과 로딩 완료: {len(loaded_results)}개 성공")
            return loaded_results
            
        except Exception as e:
            print(f"❌ 분석 결과 로딩 중 오류: {e}")
            return []
    
    def generate_consolidated_summary(self, analysis_results: list, chart_type: str) -> dict:
        """
        여러 종목의 분석 결과를 요약 프롬프트로 통합 요약
        
        Args:
            analysis_results (list): 분석 결과 리스트
            chart_type (str): 차트 유형 (daily/weekly/monthly)
            
        Returns:
            dict: 통합 요약 결과
        """
        try:
            print(f"🤖 {chart_type} 통합 요약 생성 중...")
            
            if not analysis_results:
                print("❌ 분석 결과가 없습니다")
                return None
            
            # 차트 유형이 이미 한글인 경우 그대로 사용, 영문인 경우 변환
            if chart_type in ["일봉", "주봉", "월봉"]:
                chart_type_kr = chart_type
            else:
                chart_type_kr = {
                    "daily": "일봉",
                    "weekly": "주봉", 
                    "monthly": "월봉"
                }.get(chart_type, "일봉")
            
            # 요약 프롬프트 가져오기
            summary_prompt = self.prompts.get_summary_prompt(chart_type_kr)
            if not summary_prompt:
                print(f"❌ {chart_type_kr} 요약 프롬프트를 찾을 수 없습니다")
                return None
            
            # 거래일 추출 (첫 번째 유효한 거래일 사용)
            trading_date = self._extract_trading_date_from_results(analysis_results)
            
            # 거래타입 추출 (첫 번째 유효한 거래타입 사용)
            trading_type = self._extract_trading_type_from_results(analysis_results)
            
            # 분석 결과 요약 데이터 준비
            summary_data = {
                "분석종목수": len(analysis_results),
                "차트유형": chart_type_kr,
                "분석일시": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "종목별요약": []
            }
            
            # 각 종목별 핵심 정보 추출 (하이브리드 방식: 파일명에서 기본 정보만 추출)
            for result in analysis_results:
                try:
                    # 파일 정보에서 종목 정보 추출
                    file_info = result.get("_file_info", {})
                    file_name = file_info.get("file_name", "")
                    
                    # 파일명에서 종목 정보 추출
                    stock_name, stock_code = self.stock_mapper.extract_stock_info_from_filename(file_name)
                    
                    summary_item = {
                        "종목명": stock_name,
                        "종목번호": stock_code,
                        "분석일시": file_info.get("loaded_at", "N/A"),
                        "파일명": file_name
                    }
                    
                    # 개별 종목 JSON에서 거래 정보 추출하여 추가
                    stock_info = result.get("종목정보", {})
                    if stock_info:
                        # 거래일 정보 추가
                        if "거래일" in stock_info:
                            summary_item["거래일"] = stock_info["거래일"]
                        
                        # 거래대금 정보 추가
                        if "거래대금" in stock_info:
                            summary_item["거래대금"] = stock_info["거래대금"]
                        
                        # 거래율 정보 추가
                        if "거래율" in stock_info:
                            summary_item["거래율"] = stock_info["거래율"]
                        
                        # 순위 정보 추가
                        if "순위" in stock_info:
                            summary_item["순위"] = stock_info["순위"]
                        
                        # 유통주식수 정보 추가
                        if "유통주식수" in stock_info:
                            summary_item["유통주식수"] = stock_info["유통주식수"]
                        
                        # 거래대금 정보 추가
                        if "거래대금" in stock_info:
                            summary_item["거래대금"] = stock_info["거래대금"]
                    
                    # 하이브리드 방식: 상세 정보는 개별 워드 파일에서 확인하도록 안내
                    summary_item["주요정보"] = {
                        "상세분석": "개별 워드 파일에서 확인 가능",
                        "파일경로": file_info.get("file_path", "N/A"),
                        "개별분석JSON": result.get("individual_analysis_file", "N/A")
                    }
                    
                    # 투자 아이디어도 개별 워드 파일에서 확인하도록 안내
                    summary_item["투자아이디어"] = "개별 워드 파일에서 확인 가능"
                    
                    summary_data["종목별요약"].append(summary_item)
                    
                except Exception as e:
                    print(f"⚠️ 종목별 요약 추출 실패: {e}")
                    continue
            
            # AI 요약 분석 실행 (하이브리드 방식: 개별 분석 텍스트를 직접 요약)
            if self.analyzer:
                # 개별 분석 결과에서 AI 응답 텍스트 수집
                individual_analysis_texts = []
                for result in analysis_results:
                    try:
                        # AI 응답 텍스트 추출
                        ai_response = ""
                        if isinstance(result, dict):
                            if "original_ai_response" in result:
                                ai_response = result["original_ai_response"]
                            elif "ai_response" in result:
                                ai_response = result["ai_response"]
                            else:
                                # 구조화된 데이터를 텍스트로 변환
                                ai_response = self._convert_structured_data_to_text(result)
                        elif isinstance(result, str):
                            ai_response = result
                        else:
                            ai_response = str(result)
                        
                        if ai_response:
                            # 파일 정보 추가
                            file_info = result.get("_file_info", {})
                            file_name = file_info.get("file_name", "unknown")
                            stock_name, stock_code = self.stock_mapper.extract_stock_info_from_filename(file_name)
                            
                            individual_analysis_texts.append(f"=== {stock_name} ({stock_code}) ===\n{ai_response}\n")
                            
                    except Exception as e:
                        print(f"⚠️ 개별 분석 텍스트 추출 실패: {e}")
                        continue
                
                # 요약 프롬프트에 개별 분석 텍스트 추가
                if individual_analysis_texts:
                    all_analysis_text = "\n".join(individual_analysis_texts)
                    formatted_prompt = summary_prompt + f"\n\n개별 분석 결과들:\n{all_analysis_text}\n\n"
                    formatted_prompt += "위 개별 분석 결과들을 바탕으로 전체적인 시장 동향과 주요 투자 포인트를 요약해주세요."
                else:
                    # 개별 분석 텍스트가 없는 경우 기본 데이터 사용
                    formatted_prompt = summary_prompt + f"\n\n분석 데이터:\n{json.dumps(summary_data, ensure_ascii=False, indent=2)}\n\n"
                    formatted_prompt += "위 데이터를 바탕으로 전체적인 시장 동향과 주요 투자 포인트를 요약해주세요."
                
                # AI 텍스트 분석 실행
                ai_result = self.analyzer.analyze_text_with_prompt(formatted_prompt)
                
                if ai_result:
                    # 통합 요약 결과 구성
                    consolidated_result = {
                        "summary_meta": {
                            "chart_type": chart_type_kr,
                            "total_stocks": len(analysis_results),
                            "trading_date": trading_date,
                            "trading_type": trading_type,
                            "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                            "summary_method": "AI_요약_프롬프트"
                        },
                        "market_summary": ai_result,
                        "stock_details": summary_data["종목별요약"],
                        "raw_analysis_count": len(analysis_results),
                    }
                    
                    print(f"✅ {chart_type_kr} 통합 요약 생성 완료")
                    return consolidated_result
                else:
                    print(f"⚠️ AI 요약 분석 실패, 기본 요약 사용")
            
            # AI 분석 실패 시 기본 요약
            consolidated_result = {
                "summary_meta": {
                    "chart_type": chart_type_kr,
                    "total_stocks": len(analysis_results),
                    "trading_date": trading_date,
                    "trading_type": trading_type,
                    "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "summary_method": "기본_요약"
                },
                "market_summary": {
                    "분석_결과": f"{chart_type_kr} 차트 {len(analysis_results)}개 종목 분석 완료",
                    "생성_시간": datetime.now().isoformat(),
                    "분석_유형": f"{chart_type_kr}_통합요약"
                },
                "stock_details": summary_data["종목별요약"],
                "raw_analysis_count": len(analysis_results)
            }
            
            print(f"✅ {chart_type_kr} 기본 요약 생성 완료")
            return consolidated_result
            
        except Exception as e:
            print(f"❌ 통합 요약 생성 중 오류: {e}")
            return None
    
    def _extract_trading_date_from_results(self, analysis_results: list) -> str:
        """
        개별 분석 결과에서 거래일 추출 (첫 번째 유효한 거래일 사용)
        
        Args:
            analysis_results (list): 분석 결과 리스트
            
        Returns:
            str: 거래일 (YYYY-MM-DD 형식) 또는 "N/A"
        """
        try:
            for result in analysis_results:
                try:
                    # 종목정보에서 거래일 추출
                    stock_info = result.get("종목정보", {})
                    trading_date = stock_info.get("거래일", "")
                    
                    if trading_date and trading_date != "N/A":
                        print(f"✅ 거래일 추출 성공: {trading_date}")
                        return trading_date
                except Exception as e:
                    print(f"⚠️ 거래일 추출 실패: {e}")
                    continue
            
            print("⚠️ 유효한 거래일을 찾을 수 없습니다")
            return "N/A"
            
        except Exception as e:
            print(f"❌ 거래일 추출 중 오류: {e}")
            return "N/A"
    
    def _extract_trading_type_from_results(self, analysis_results: list) -> str:
        """
        개별 분석 결과에서 거래타입 추출 (첫 번째 유효한 거래타입 사용)
        
        Args:
            analysis_results (list): 분석 결과 리스트
            
        Returns:
            str: 거래타입 (거래대금/거래율) 또는 ""
        """
        try:
            for result in analysis_results:
                try:
                    # 종목정보에서 거래타입 추출
                    stock_info = result.get("종목정보", {})
                    trading_type = stock_info.get("거래타입", "")
                    
                    if trading_type and trading_type.strip():
                        print(f"✅ 거래타입 추출 성공: {trading_type}")
                        return trading_type
                except Exception as e:
                    print(f"⚠️ 거래타입 추출 실패: {e}")
                    continue
            
            print("⚠️ 유효한 거래타입을 찾을 수 없습니다")
            return ""
            
        except Exception as e:
            print(f"❌ 거래타입 추출 중 오류: {e}")
            return ""
    
    def save_summary_files(self, consolidated_result: dict, chart_type: str, output_dir: str = "ai_analysis_results") -> tuple:
        """
        통합 요약 결과를 JSON 및 DOCX 파일로 저장
        
        Args:
            consolidated_result (dict): 통합 요약 결과
            chart_type (str): 차트 유형
            output_dir (str): 출력 디렉토리
            
        Returns:
            tuple: (json_path, docx_path, success)
        """
        try:
            print(f"💾 요약 파일 저장 중...")
            
            if not consolidated_result:
                print("❌ 저장할 요약 결과가 없습니다")
                return None, None, False
            
            # 출력 디렉토리 생성
            os.makedirs(output_dir, exist_ok=True)
            
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            chart_type_kr = consolidated_result.get("summary_meta", {}).get("chart_type", chart_type)
            
            # JSON 파일 저장
            json_filename = f"summary_{chart_type}_{timestamp}.json"
            json_path = os.path.join(output_dir, json_filename)
            
            with open(json_path, 'w', encoding='utf-8') as f:
                json.dump(consolidated_result, f, ensure_ascii=False, indent=2)
            
            print(f"✅ JSON 요약 파일 저장 완료: {json_path}")
            
            # 원본 AI 응답을 TXT 파일로 저장 (가공하지 않은 순수 응답)
            txt_filename = f"summary_{chart_type}_{timestamp}_raw.txt"
            txt_path = os.path.join(output_dir, txt_filename)
            
            # AI 원본 응답 추출
            market_summary = consolidated_result.get("market_summary", {})
            raw_ai_response = market_summary.get("분석_결과", "")
            
            if raw_ai_response:
                with open(txt_path, 'w', encoding='utf-8') as f:
                    f.write("=" * 60 + "\n")
                    f.write(f"AI 요약 분석 원본 응답 ({chart_type_kr})\n")
                    f.write("=" * 60 + "\n")
                    f.write(f"생성일시: {consolidated_result.get('summary_meta', {}).get('generated_at', 'N/A')}\n")
                    f.write(f"분석종목수: {consolidated_result.get('summary_meta', {}).get('total_stocks', 0)}개\n")
                    f.write("=" * 60 + "\n\n")
                    f.write(raw_ai_response)
                    f.write("\n\n" + "=" * 60 + "\n")
                    f.write("※ 이 파일은 AI(재미나이)의 원본 응답을 가공하지 않고 그대로 저장한 것입니다.\n")
                
                print(f"✅ 원본 AI 응답 TXT 파일 저장 완료: {txt_path}")
            else:
                print("⚠️ 원본 AI 응답이 없어 TXT 파일을 생성하지 않습니다.")
            
            # DOCX 파일 저장
            docx_filename = f"summary_{chart_type}_{timestamp}.docx"
            docx_path = os.path.join(output_dir, docx_filename)
            
            docx_success = self._create_summary_docx(consolidated_result, docx_path)
            
            if docx_success:
                print(f"✅ DOCX 요약 파일 저장 완료: {docx_path}")
                return json_path, docx_path, True
            else:
                print(f"⚠️ DOCX 파일 저장 실패")
                return json_path, None, True
                
        except Exception as e:
            print(f"❌ 요약 파일 저장 중 오류: {e}")
            return None, None, False
    
    def _add_ranking_table_to_docx(self, doc, consolidated_result: dict, trading_type: str) -> None:
        """
        상위 50개 종목 순위표를 DOCX에 추가
        
        Args:
            doc: Word 문서 객체
            consolidated_result (dict): 통합 요약 결과
            trading_type (str): 거래 타입 (거래대금/거래율)
        """
        try:
            # stock_details에서 데이터 추출
            stock_details = consolidated_result.get("stock_details", [])
            if not stock_details:
                print("⚠️ stock_details 데이터가 없습니다.")
                return
            
            # 거래대금 또는 거래율 기준으로 정렬
            sort_key = "거래대금" if trading_type == "거래대금" else "거래율"
            
            # 정렬 가능한 데이터만 필터링하고 정렬
            sortable_stocks = []
            for stock in stock_details:
                if sort_key in stock and stock[sort_key] != "N/A":
                    try:
                        # 거래대금: "149억원" -> 149
                        # 거래율: "72.47%" -> 72.47
                        value_str = stock[sort_key]
                        if sort_key == "거래대금":
                            # "149억원" -> 149
                            value = float(value_str.replace("억원", "").replace(",", ""))
                        else:  # 거래율
                            # "72.47%" -> 72.47
                            value = float(value_str.replace("%", "").replace(",", ""))
                        
                        sortable_stocks.append((value, stock))
                    except (ValueError, AttributeError):
                        continue
            
            # 내림차순 정렬 (높은 값부터)
            sortable_stocks.sort(key=lambda x: x[0], reverse=True)
            
            # 상위 50개만 선택
            top_50_stocks = [stock for _, stock in sortable_stocks[:50]]
            
            if not top_50_stocks:
                print("⚠️ 정렬 가능한 종목 데이터가 없습니다.")
                return
            
            # 표 제목 추가
            table_title = doc.add_heading(f'📈 {trading_type} 기준 상위 50개 종목 순위', level=1)
            for run in table_title.runs:
                run.font.name = '맑은 고딕'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            
            # 표 생성 (6열: 순위|거래율/거래대금|종목명|순위|거래율/거래대금|종목명)
            # 25행 (50개 종목을 2열로 나누어 표시)
            table = doc.add_table(rows=26, cols=6)  # 헤더 1행 + 데이터 25행
            table.style = 'Table Grid'
            
            # 헤더 설정
            header_cells = table.rows[0].cells
            header_texts = ['순위', '거래율' if trading_type == "거래율" else '거래대금', '종목명', 
                           '순위', '거래율' if trading_type == "거래율" else '거래대금', '종목명']
            
            for i, header_text in enumerate(header_texts):
                header_cells[i].text = header_text
                # 헤더 셀 스타일링
                for paragraph in header_cells[i].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = '맑은 고딕'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                        run.font.bold = True
            
            # 칸 크기 조정 (순위 칸 50% 줄이고, 거래율/종목명 칸 늘리기)
            # 전체 테이블 너비를 100%로 설정
            table.width = Inches(7.0)  # 전체 테이블 너비
            
            # 각 칸의 너비 설정 (Inches 단위)
            # 순위: 0.5인치 (50% 줄임), 거래율/거래대금: 1.2인치, 종목명: 1.8인치
            column_widths = [0.5, 1.2, 1.8, 0.5, 1.2, 1.8]  # 6개 칸의 너비
            
            for i, width in enumerate(column_widths):
                table.columns[i].width = Inches(width)
            
            # 데이터 행 채우기 (25행, 각 행에 2개 종목)
            for row_idx in range(1, 26):  # 1~25행
                data_cells = table.rows[row_idx].cells
                
                # 왼쪽 열 (1-25위)
                left_stock_idx = row_idx - 1
                if left_stock_idx < len(top_50_stocks):
                    left_stock = top_50_stocks[left_stock_idx]
                    data_cells[0].text = str(left_stock_idx + 1)  # 순위
                    data_cells[1].text = left_stock.get(sort_key, "N/A")  # 거래율/거래대금
                    data_cells[2].text = left_stock.get("종목명", "N/A")  # 종목명
                
                # 오른쪽 열 (26-50위)
                right_stock_idx = left_stock_idx + 25
                if right_stock_idx < len(top_50_stocks):
                    right_stock = top_50_stocks[right_stock_idx]
                    data_cells[3].text = str(right_stock_idx + 1)  # 순위
                    data_cells[4].text = right_stock.get(sort_key, "N/A")  # 거래율/거래대금
                    data_cells[5].text = right_stock.get("종목명", "N/A")  # 종목명
            
            # 테이블 전체에 한글 폰트 적용
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.name = '맑은 고딕'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            
            print(f"✅ {trading_type} 기준 상위 50개 종목 순위표 추가 완료")
            
        except Exception as e:
            print(f"❌ 순위표 추가 중 오류: {e}")
    
    def _create_summary_docx(self, consolidated_result: dict, output_path: str) -> bool:
        """
        통합 요약 결과를 DOCX 파일로 생성
        
        Args:
            consolidated_result (dict): 통합 요약 결과
            output_path (str): 출력 파일 경로
            
        Returns:
            bool: 생성 성공 여부
        """
        try:
            # Word 문서 생성
            doc = Document()
            
            # 한글 폰트 설정
            from docx.oxml.ns import qn
            
            # 메타 정보 추출
            meta = consolidated_result.get("summary_meta", {})
            chart_type = meta.get("chart_type", "차트")
            total_stocks = meta.get("total_stocks", 0)
            generated_at = meta.get("generated_at", "N/A")
            
            # 제목 설정
            title = doc.add_heading(f'{chart_type} 통합 분석 요약 보고서', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 한글 폰트 적용
            for run in title.runs:
                run.font.name = '맑은 고딕'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                run.font.size = Pt(14)
            
            # 요약 개요 테이블 생성
            doc.add_heading('📊 분석 개요', level=1)
            
            # 분석 개요 테이블 생성
            table = doc.add_table(rows=2, cols=3)
            table.style = 'Table Grid'
            
            # 헤더 행 설정
            header_cells = table.rows[0].cells
            header_cells[0].text = '차트 유형'
            header_cells[1].text = '분석 종목 수'
            header_cells[2].text = '생성일시'
            
            # 데이터 행 설정
            data_cells = table.rows[1].cells
            data_cells[0].text = chart_type
            data_cells[1].text = f'{total_stocks}개'
            # 날짜 형식을 YY-MM-DD (ddd) HH 형태로 변경
            now = datetime.now()
            weekday_korean = ['월', '화', '수', '목', '금', '토', '일']
            formatted_date = f"{now.strftime('%y.%m.%d')}({weekday_korean[now.weekday()]}) {now.strftime('%H')}시"
            data_cells[2].text = formatted_date
            
            # 테이블 전체에 한글 폰트 적용
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = '맑은 고딕'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            
            # 분석 개요 설명 추가 (모든 차트 유형에 적용)
            # summary_meta에서 거래일, 차트 유형, 거래타입 추출
            summary_meta = consolidated_result.get("summary_meta", {}) if consolidated_result else {}
            trading_date = summary_meta.get("trading_date", "N/A")
            chart_type_from_meta = summary_meta.get("chart_type", chart_type)
            trading_type_from_meta = summary_meta.get("trading_type", "거래대금")
            
            # 거래일 형식 변환 (YYYY-MM-DD → M월 D일)
            formatted_trading_date = "N/A"
            if trading_date != "N/A":
                try:
                    date_obj = datetime.strptime(trading_date, "%Y-%m-%d")
                    formatted_trading_date = f"{date_obj.month}월 {date_obj.day}일"
                except:
                    formatted_trading_date = trading_date
            
            overview_desc = f"{formatted_trading_date} {trading_type_from_meta} 기준 상위 50개 종목의 {chart_type_from_meta} 차트를 분석한 결과, 특이사항을 나타낸 종목은 아래와 같습니다. 핵심내용만 요약하여 제공해드리고, 자세한 내용은 첨부파일의 개별 종목별 차트분석 결과를 참고하시기 바랍니다."
            
            para_desc = doc.add_paragraph(overview_desc)
            for run in para_desc.runs:
                run.font.name = '맑은 고딕'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            
            # 상위 50개 종목 순위표 추가
            self._add_ranking_table_to_docx(doc, consolidated_result, trading_type_from_meta)
            
            # 마크다운 파싱 함수 정의
            def _parse_markdown_to_word(text: str, doc) -> None:
                """
                마크다운 텍스트를 Word 문서 형식으로 파싱하여 추가
                
                Args:
                    text (str): 마크다운 형태의 텍스트
                    doc: Word 문서 객체
                """
                try:
                    import re
                    
                    # **내용** 형태를 볼드체로 변환하는 정규식
                    bold_pattern = r'\*\*(.*?)\*\*'
                    
                    # 줄바꿈으로 분리
                    lines = text.split('\n')
                    
                    for line in lines:
                        line = line.strip()
                        if not line:
                            continue
                        
                        # **내용** 패턴을 찾아서 볼드체로 변환
                        if re.search(bold_pattern, line):
                            # 볼드체가 포함된 줄 처리
                            para = doc.add_paragraph()
                            
                            # **내용** 패턴으로 분할
                            parts = re.split(bold_pattern, line)
                            
                            for i, part in enumerate(parts):
                                if i % 2 == 0:
                                    # 일반 텍스트
                                    if part:
                                        run = para.add_run(part)
                                        run.font.name = '맑은 고딕'
                                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                                else:
                                    # 볼드체 텍스트
                                    if part:
                                        run = para.add_run(part)
                                        run.font.name = '맑은 고딕'
                                        run.bold = True
                                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                        else:
                            # 볼드체가 없는 일반 줄 처리
                            if line.startswith('📊') or line.startswith('📈') or line.startswith('📉') or line.startswith('💡') or line.startswith('🔍') or line.startswith('⚠️') or line.startswith('✅'):
                                para = doc.add_heading(line, level=2)
                            elif line.startswith('•') or line.startswith('-') or line.startswith('*'):
                                para = doc.add_paragraph(line)
                            else:
                                para = doc.add_paragraph(line)
                            
                            # 한글 폰트 적용
                            for run in para.runs:
                                run.font.name = '맑은 고딕'
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                                
                except Exception as e:
                    print(f"⚠️ 마크다운 파싱 중 오류: {e}")
                    # 오류 발생 시 일반 텍스트로 처리
                    para = doc.add_paragraph(text)
                    for run in para.runs:
                        run.font.name = '맑은 고딕'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            
            # AI 시장 요약
            market_summary = consolidated_result.get("market_summary", {})
            if market_summary:
                doc.add_heading('🤖 시장 분석 요약', level=1)
                
                analysis_result = market_summary.get("분석_결과", "")
                if analysis_result:
                    # JSON 코드 블록 제거 (Word 문서용)
                    cleaned_analysis = analysis_result.strip()
                    if cleaned_analysis.startswith('```json'):
                        cleaned_analysis = cleaned_analysis[7:]
                    elif cleaned_analysis.startswith('```'):
                        cleaned_analysis = cleaned_analysis[3:]
                    if cleaned_analysis.endswith('```'):
                        cleaned_analysis = cleaned_analysis[:-3]
                    cleaned_analysis = cleaned_analysis.strip()
                    
                    # JSON 형태의 응답인지 확인하고 파싱 시도
                    if cleaned_analysis.startswith('{') and cleaned_analysis.endswith('}'):
                        try:
                            # JSON 파싱 시도
                            parsed_data = json.loads(cleaned_analysis)
                            
                            # 시장동향요약 추가
                            if "시장동향요약" in parsed_data:
                                doc.add_heading('📈 시장 동향 요약', level=2)
                                market_trend = doc.add_paragraph(parsed_data["시장동향요약"])
                                for run in market_trend.runs:
                                    run.font.name = '맑은 고딕'
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                            
                            # 주요투자포인트 추가
                            if "주요투자포인트" in parsed_data:
                                doc.add_heading('💡 주요 투자 포인트', level=2)
                                investment_points = parsed_data["주요투자포인트"]
                                
                                if isinstance(investment_points, list):
                                    for i, point in enumerate(investment_points, 1):
                                        paras_to_format = []
                                        
                                        if isinstance(point, dict):
                                            # 딕셔너리 형태인 경우
                                            if "종목" in point and "포인트" in point:
                                                point_title = doc.add_paragraph(f"{i}. {point['종목']}")
                                                point_title.style = 'List Number'
                                                point_content = doc.add_paragraph(f"   {point['포인트']}")
                                                paras_to_format = [point_title, point_content]
                                            else:
                                                # 다른 딕셔너리 구조인 경우
                                                point_text = f"{i}. " + str(point)
                                                point_para = doc.add_paragraph(point_text)
                                                paras_to_format = [point_para]
                                        else:
                                            # 문자열 형태인 경우
                                            point_para = doc.add_paragraph(f"{i}. {str(point)}")
                                            paras_to_format = [point_para]
                                        
                                        # 폰트 설정
                                        for para in paras_to_format:
                                            for run in para.runs:
                                                run.font.name = '맑은 고딕'
                                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                                elif isinstance(investment_points, str):
                                    # 문자열인 경우 그대로 추가
                                    point_para = doc.add_paragraph(investment_points)
                                    for run in point_para.runs:
                                        run.font.name = '맑은 고딕'
                                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                            
                            # 전반적시장분위기 추가
                            if "전반적시장분위기" in parsed_data:
                                doc.add_heading('🌍 전반적 시장 분위기', level=2)
                                market_mood = doc.add_paragraph(parsed_data["전반적시장분위기"])
                                for run in market_mood.runs:
                                    run.font.name = '맑은 고딕'
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                                    
                        except json.JSONDecodeError:
                            # JSON 파싱 실패 시 마크다운 파싱 적용
                            _parse_markdown_to_word(cleaned_analysis, doc)
                    else:
                        # JSON이 아닌 경우 마크다운 파싱 적용
                        _parse_markdown_to_word(cleaned_analysis, doc)
                else:
                    doc.add_paragraph("AI 분석 결과가 없습니다.")
            
            # 한글 폰트 적용 (모든 단락)
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    run.font.name = '맑은 고딕'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            
            # 문서 저장
            doc.save(output_path)
            print(f"📄 DOCX 요약 파일 생성 완료: {output_path}")
            return True
            
        except Exception as e:
            print(f"❌ DOCX 파일 생성 중 오류: {e}")
            return False
    
    def generate_all_summaries(self, results_dir: str = "ai_analysis_results") -> dict:
        """
        모든 차트 유형에 대해 통합 요약 파일 생성
        
        Args:
            results_dir (str): 분석 결과 폴더 경로
            
        Returns:
            dict: 생성된 요약 파일 정보
        """
        try:
            print(f"🚀 전체 통합 요약 파일 생성 시작")
            print("="*60)
            
            # 1. 분석 결과 스캔
            grouped_files = self.scan_analysis_results(results_dir)
            if not grouped_files:
                print("❌ 분석할 파일이 없습니다")
                return {}
            
            generated_summaries = {}
            
            # 2. 각 차트 유형별로 요약 생성
            for chart_type, file_paths in grouped_files.items():
                if not file_paths:
                    print(f"⚠️ {chart_type} 파일이 없어 건너뜁니다")
                    continue
                
                print(f"\n📊 {chart_type} 요약 생성 중...")
                
                # 분석 결과 로드
                analysis_results = self.load_analysis_results(file_paths)
                if not analysis_results:
                    print(f"❌ {chart_type} 분석 결과 로드 실패")
                    continue
                
                # 통합 요약 생성
                consolidated_result = self.generate_consolidated_summary(analysis_results, chart_type)
                if not consolidated_result:
                    print(f"❌ {chart_type} 통합 요약 생성 실패")
                    continue
                
                # 요약 파일 저장
                json_path, docx_path, success = self.save_summary_files(consolidated_result, chart_type, results_dir)
                
                if success:
                    generated_summaries[chart_type] = {
                        "json_path": json_path,
                        "docx_path": docx_path,
                        "total_stocks": len(analysis_results),
                        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    }
                    print(f"✅ {chart_type} 요약 완료")
                else:
                    print(f"❌ {chart_type} 요약 파일 저장 실패")
            
            print(f"\n🎉 전체 통합 요약 파일 생성 완료!")
            print(f"📊 생성된 요약:")
            for chart_type, info in generated_summaries.items():
                print(f"   - {chart_type}: {info['total_stocks']}개 종목")
                print(f"     JSON: {info['json_path']}")
                print(f"     DOCX: {info['docx_path']}")
            
            return generated_summaries
            
        except Exception as e:
            print(f"❌ 전체 요약 생성 중 오류: {e}")
            return {}

    def create_consolidated_summary_from_files(self, chart_type: str, file_paths: list) -> Optional[Dict[str, str]]:
        """
        특정 파일들로부터 요약 분석 실행
        
        Args:
            chart_type (str): 차트 유형 ("daily", "weekly", "monthly")
            file_paths (list): 분석할 JSON 파일 경로들
            
        Returns:
            Optional[Dict[str, str]]: 생성된 파일 경로들 {"json_path": "...", "docx_path": "..."}
        """
        try:
            print(f"🔍 {chart_type} 특정 파일들로 요약 분석 중... (파일 수: {len(file_paths)})")
            
            if not file_paths:
                print(f"⚠️ 분석할 파일이 없습니다.")
                return None
            
            # 1. 파일 경로들을 차트 유형별로 그룹화
            chart_results = {chart_type: file_paths}
            
            # 2. 분석 결과 로딩
            print(f"📄 분석 결과 로딩 중: {len(file_paths)}개 파일")
            analysis_data = self.load_analysis_results(file_paths)
            if not analysis_data:
                print(f"❌ 분석 결과 로딩 실패")
                return None
            
            print(f"✅ 분석 결과 로딩 완료: {len(analysis_data)}개 성공")
            
            # 3. 통합 요약 생성
            print(f"🤖 {chart_type} 통합 요약 생성 중...")
            summary_result = self.generate_consolidated_summary(analysis_data, chart_type)
            
            if summary_result:
                print(f"✅ {chart_type} 통합 요약 생성 완료")
                
                # 4. 파일 저장
                print(f"💾 요약 파일 저장 중...")
                json_path, docx_path, success = self.save_summary_files(summary_result, chart_type)
                
                if success:
                    print(f"✅ 요약 파일 저장 완료")
                    return {
                        "json_path": json_path,
                        "docx_path": docx_path
                    }
                else:
                    print(f"❌ 요약 파일 저장 실패")
                    return None
            else:
                print(f"❌ {chart_type} 통합 요약 생성 실패")
                return None
                
        except Exception as e:
            print(f"❌ {chart_type} 요약 분석 중 오류: {e}")
            import traceback
            traceback.print_exc()
            return None

    def create_consolidated_summary_by_type(self, chart_type: str) -> Optional[Dict[str, str]]:
        """
        특정 차트 유형의 분석 결과들을 통합하여 요약 분석 실행
        
        Args:
            chart_type (str): 차트 유형 ("daily", "weekly", "monthly")
            
        Returns:
            Optional[Dict[str, str]]: 생성된 파일 경로들 {"json_path": "...", "docx_path": "..."}
        """
        try:
            print(f"🔍 {chart_type} 분석 결과 스캔 중...")
            
            # 1. 해당 차트 유형의 분석 결과 스캔
            analysis_results = self.scan_analysis_results()
            if not analysis_results or chart_type not in analysis_results:
                print(f"⚠️ {chart_type} 분석 결과가 없습니다.")
                return None
            
            chart_results = analysis_results[chart_type]
            if not chart_results:
                print(f"⚠️ {chart_type} 분석 파일이 없습니다.")
                return None
            
            print(f"📊 {chart_type} 분석 파일 {len(chart_results)}개 발견")
            
            # 2. 분석 결과 로드
            loaded_results = self.load_analysis_results(chart_results)
            if not loaded_results:
                print(f"⚠️ {chart_type} 분석 결과 로드 실패")
                return None
            
            # 3. 차트 유형을 한글로 변환
            chart_type_korean = {
                "daily": "일봉",
                "weekly": "주봉", 
                "monthly": "월봉"
            }.get(chart_type, chart_type)
            
            # 4. AI 요약 분석 실행 (차트 타입을 한글로 전달)
            summary_result = self.generate_consolidated_summary(loaded_results, chart_type_korean)
            if not summary_result:
                print(f"⚠️ {chart_type} AI 요약 분석 실패")
                return None
            
            # 5. 요약 파일 저장
            json_path, docx_path, success = self.save_summary_files(summary_result, chart_type)
            
            if success:
                return {
                    "json_path": json_path,
                    "docx_path": docx_path
                }
            else:
                return None
            
        except Exception as e:
            print(f"❌ {chart_type} 통합 요약 분석 중 오류: {e}")
            return None

def create_summary_files(results_dir: str = "ai_analysis_results"):
    """
    통합 요약 파일 생성 편의 함수
    
    Args:
        results_dir (str): 분석 결과 폴더 경로
    """
    print("🚀 통합 요약 파일 생성기 시작")
    print("="*60)
    
    # 데이터베이스 설정 로드
    try:
        from config import config
        db_config = config.get_database_config()
        print("✅ 데이터베이스 설정 로드 완료")
    except Exception as e:
        print(f"⚠️ 데이터베이스 설정 로드 실패: {e}")
        db_config = None
    
    # 요약 파일 생성기 초기화
    try:
        generator = SummaryFileGenerator(db_config)
        
        # 모든 요약 파일 생성
        results = generator.generate_all_summaries(results_dir)
        
        if results:
            print(f"\n✅ 통합 요약 파일 생성 완료!")
            return results
        else:
            print(f"\n❌ 통합 요약 파일 생성 실패")
            return None
            
    except Exception as e:
        print(f"❌ 요약 파일 생성기 초기화 실패: {e}")
        return None

def create_consolidated_analysis(analysis_results: list, chart_type: str) -> dict:
    """
    개별 분석 결과들을 통합하여 consolidated analysis 생성
    
    Args:
        analysis_results (list): 개별 분석 결과 리스트
        chart_type (str): 차트 유형
        
    Returns:
        dict: 통합 분석 결과
    """
    try:
        print(f"🔗 {chart_type} 통합 분석 생성 중...")
        
        consolidated_result = {
            "metadata": {
                "chart_type": chart_type,
                "total_stocks": len(analysis_results),
                "created_at": datetime.now().isoformat(),
                "analysis_version": "1.0",
                "file_type": "consolidated_analysis"
            },
            "analysis_summary": {
                "successful_analyses": 0,
                "failed_analyses": 0,
                "average_score": 0
            },
            "individual_results": []
        }
        
        total_score = 0
        successful_count = 0
        
        for result in analysis_results:
            # 개별 분석 결과에서 성공 여부 확인
            success = False
            if isinstance(result, dict):
                # 종합분석점수가 있으면 성공으로 간주
                if "종합분석점수" in result:
                    success = True
                # 또는 success 키가 있으면 그 값 사용
                elif "success" in result:
                    success = result["success"]
            
            if success:
                successful_count += 1
                
                # 분석 점수 추출
                score = 0
                if "종합분석점수" in result:
                    score_info = result["종합분석점수"]
                    if isinstance(score_info, dict):
                        score = score_info.get("점수", 0)
                    else:
                        score = score_info
                elif "analysis_score" in result:
                    score = result["analysis_score"]
                
                total_score += score
                
                # 종목 정보 추출
                stock_info = result.get("종목정보", {})
                stock_code = stock_info.get("종목번호", result.get("stock_code", "000000"))
                stock_name = stock_info.get("종목명", result.get("stock_name", "unknown"))
                
                # AI 분석 데이터가 있는 경우 포함
                individual_result = {
                    "stock_code": stock_code,
                    "stock_name": stock_name,
                    "analysis_score": score,
                    "success": True,
                    "timestamp": result.get("timestamp", datetime.now().isoformat()),
                    "individual_analysis_file": result.get("individual_analysis_file", "")
                }
                
                if "ai_analysis_data" in result:
                    individual_result["ai_analysis_data"] = result["ai_analysis_data"]
                
                consolidated_result["individual_results"].append(individual_result)
            else:
                # 실패한 경우
                stock_info = result.get("종목정보", {}) if isinstance(result, dict) else {}
                stock_code = stock_info.get("종목번호", result.get("stock_code", "000000")) if isinstance(result, dict) else "000000"
                stock_name = stock_info.get("종목명", result.get("stock_name", "unknown")) if isinstance(result, dict) else "unknown"
                
                consolidated_result["individual_results"].append({
                    "stock_code": stock_code,
                    "stock_name": stock_name,
                    "analysis_score": 0,
                    "success": False,
                    "error": result.get("error", "분석 실패") if isinstance(result, dict) else "분석 실패",
                    "timestamp": result.get("timestamp", datetime.now().isoformat()) if isinstance(result, dict) else datetime.now().isoformat(),
                    "individual_analysis_file": result.get("individual_analysis_file", "") if isinstance(result, dict) else ""
                })
        
        # 요약 통계 업데이트
        consolidated_result["analysis_summary"]["successful_analyses"] = successful_count
        consolidated_result["analysis_summary"]["failed_analyses"] = len(analysis_results) - successful_count
        consolidated_result["analysis_summary"]["average_score"] = total_score / successful_count if successful_count > 0 else 0
        
        print(f"✅ {chart_type} 통합 분석 생성 완료: {len(analysis_results)}개 종목")
        return consolidated_result
        
    except Exception as e:
        print(f"❌ {chart_type} 통합 분석 생성 실패: {e}")
        return None

def create_consolidated_word_document(consolidated_result: dict, chart_type: str, output_path: str) -> bool:
    """
    통합 분석 결과를 바탕으로 통합 Word 문서 생성
    
    Args:
        consolidated_result (dict): 통합 분석 결과
        chart_type (str): 차트 유형
        output_path (str): 출력 파일 경로
        
    Returns:
        bool: 생성 성공 여부
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        
        doc = Document()
        
        # 제목
        title = doc.add_heading(f'{chart_type} 통합 분석 보고서', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.name = '맑은 고딕'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            run.font.size = Pt(14)
        
        # 분석 개요
        doc.add_heading('📊 분석 개요', level=1)
        metadata = consolidated_result.get("metadata", {})
        summary = consolidated_result.get("analysis_summary", {})
        
        doc.add_paragraph(f'• 차트 유형: {metadata.get("chart_type", "N/A")}')
        doc.add_paragraph(f'• 총 분석 종목 수: {metadata.get("total_stocks", 0)}개')
        doc.add_paragraph(f'• 성공: {summary.get("successful_analyses", 0)}개')
        doc.add_paragraph(f'• 실패: {summary.get("failed_analyses", 0)}개')
        doc.add_paragraph(f'• 평균 점수: {summary.get("average_score", 0):.1f}점')
        doc.add_paragraph(f'• 생성일시: {metadata.get("created_at", "N/A")}')
        
        # 개별 종목 분석 결과
        doc.add_heading('📈 개별 종목 분석 결과', level=1)
        individual_results = consolidated_result.get("individual_results", [])
        
        for i, result in enumerate(individual_results, 1):
            if result.get("success", False):
                doc.add_heading(f'{i}. {result.get("stock_name", "N/A")} ({result.get("stock_code", "N/A")})', level=2)
                doc.add_paragraph(f'분석 점수: {result.get("analysis_score", 0)}점')
                
                # AI 분석 데이터가 있는 경우 요약 정보 추가
                ai_data = result.get("ai_analysis_data", {})
                if ai_data:
                    score_info = ai_data.get("종합분석점수", {})
                    if score_info:
                        doc.add_paragraph(f'종합 점수: {score_info.get("점수", "N/A")}/100')
                        doc.add_paragraph(f'분석 요약: {score_info.get("요약", "N/A")}')
        
        # 한글 폰트 적용
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = '맑은 고딕'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        
        doc.save(output_path)
        print(f"✅ 통합 Word 문서 생성 완료: {output_path}")
        return True
        
    except Exception as e:
        print(f"❌ 통합 Word 문서 생성 실패: {e}")
        return False

def create_summary_document(consolidated_result: dict, chart_type: str, output_path: str) -> bool:
    """
    통합 분석 결과를 바탕으로 요약 문서 생성
    
    Args:
        consolidated_result (dict): 통합 분석 결과
        chart_type (str): 차트 유형
        output_path (str): 출력 파일 경로
        
    Returns:
        bool: 생성 성공 여부
    """
    
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        
        doc = Document()
        
        # 제목
        title = doc.add_heading(f'{chart_type} 분석 요약 보고서', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.name = '맑은 고딕'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            run.font.size = Pt(14)
        
        # 분석 개요
        doc.add_heading('📊 분석 개요', level=1)
        metadata = consolidated_result.get("metadata", {})
        summary = consolidated_result.get("analysis_summary", {})
        
        doc.add_paragraph(f'• 차트 유형: {metadata.get("chart_type", "N/A")}')
        doc.add_paragraph(f'• 총 분석 종목 수: {metadata.get("total_stocks", 0)}개')
        doc.add_paragraph(f'• 성공: {summary.get("successful_analyses", 0)}개')
        doc.add_paragraph(f'• 실패: {summary.get("failed_analyses", 0)}개')
        doc.add_paragraph(f'• 평균 점수: {summary.get("average_score", 0):.1f}점')
        doc.add_paragraph(f'• 생성일시: {metadata.get("created_at", "N/A")}')
        
        # 성공한 종목 요약
        doc.add_heading('✅ 성공한 종목 요약', level=1)
        successful_results = [r for r in consolidated_result.get("individual_results", []) if r.get("success", False)]
        
        if successful_results:
            for i, result in enumerate(successful_results[:10], 1):  # 상위 10개만
                doc.add_paragraph(f'{i}. {result.get("stock_name", "N/A")} ({result.get("stock_code", "N/A")}) - {result.get("analysis_score", 0)}점')
            
            if len(successful_results) > 10:
                doc.add_paragraph(f'... 외 {len(successful_results) - 10}개')
        else:
            doc.add_paragraph('성공한 분석이 없습니다.')
        
        # 실패한 종목 요약
        failed_results = [r for r in consolidated_result.get("individual_results", []) if not r.get("success", False)]
        if failed_results:
            doc.add_heading('❌ 실패한 종목 요약', level=1)
            for i, result in enumerate(failed_results[:5], 1):  # 상위 5개만
                doc.add_paragraph(f'{i}. {result.get("stock_name", "N/A")} ({result.get("stock_code", "N/A")}) - {result.get("error", "N/A")}')
            
            if len(failed_results) > 5:
                doc.add_paragraph(f'... 외 {len(failed_results) - 5}개')
        
        # 한글 폰트 적용
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = '맑은 고딕'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        
        doc.save(output_path)
        print(f"✅ 요약 문서 생성 완료: {output_path}")
        return True
        
    except Exception as e:
        print(f"❌ 요약 문서 생성 실패: {e}")
        return False

if __name__ == "__main__":
    main() 